from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "家具智能客服项目理解与面试手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1E293B"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B7791F"
RED = "9B1C1C"
GREEN = "236B3C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("家具智能客服项目理解与面试手册")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("PROJECT UNDERSTANDING GUIDE")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("家具智能客服项目\n理解与面试手册")
    set_run_font(run, size=28, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("不看代码，也能从根本上讲清楚项目的业务、架构、演进与设计取舍")
    set_run_font(run, size=13, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(callout, [9120], indent_dxa=120)
    set_repeat_table_header(callout.rows[0])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "核心目标：建立一张完整的项目心理地图，让你知道每个模块为什么存在、"
        "一条请求如何流动、系统如何保证准确性与可靠性，以及面试时如何诚实、清晰地表达。"
    )
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run("适用方向：Agent 开发 / AI 应用工程 / RAG 应用")
    set_run_font(run, size=10, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("版本：基于当前项目工作区梳理")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill=CALLOUT, color=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9120], indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}：")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)


def add_comparison_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_interview_question(
    doc: Document,
    number: int,
    question: str,
    answer: str,
    focus: str,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {question}")
    set_run_font(r, size=11.5, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("考察点：")
    set_run_font(r, size=9.5, color=GOLD, bold=True)
    r = p.add_run(focus)
    set_run_font(r, size=9.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("满分答案：")
    set_run_font(r, size=10.5, color=GREEN, bold=True)
    r = p.add_run(answer)
    set_run_font(r, size=10.5, color=INK)


def add_technology_selection_review(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("技术选型专项：为什么选择 ReAct、LangChain、LangGraph 与多 Agent", level=1)
    add_body(
        doc,
        "这一章专门训练“为什么选”而不是“它是什么”。面试官问技术选型，通常不是想听框架官网上的优点，"
        "而是想判断你是否真正经历过问题、比较过方案、理解收益与代价，并且知道怎样证明改造有效。",
    )
    add_callout(
        doc,
        "先澄清一个容易说错的词",
        "本项目中的 ReactAgent 指的是 ReAct/工具调用式 Agent，不是 React 前端框架。"
        "项目当前前端使用 Streamlit。面试时如果对方只说“React”，可以先确认他问的是前端 React，"
        "还是 Reasoning + Acting 的 ReAct 思路，避免答非所问。",
        fill="FFF8E8",
        color=GOLD,
    )
    add_callout(
        doc,
        "技术选型回答公式",
        "原始问题 → 候选方案 → 选择标准 → 最终选择 → 实际变化 → 新增代价 → 验证方法。"
        "回答时先讲业务困难，再讲框架。不要用“流行、先进、大家都在用”作为主要理由。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )

    doc.add_heading("1. 四项技术各自解决什么问题", level=2)
    add_comparison_table(
        doc,
        ["技术", "主要解决的问题", "不会自动解决的问题", "主要代价"],
        [
            [
                "ReAct / 工具循环",
                "让模型根据当前问题动态选择工具，并根据工具结果继续判断",
                "不会让知识天然正确，也不会替代证据校验",
                "调用次数、延迟、费用和行为不确定性增加",
            ],
            [
                "LangChain",
                "统一模型、消息、Prompt、工具、文档和解析器等接口",
                "不会自动提高答案准确率，也不等于完整业务架构",
                "抽象层增加，版本变化时需要适配",
            ],
            [
                "LangGraph",
                "显式组织状态、节点、分支、并行、汇合与降级",
                "不会替代业务规则、数据库事务和监控",
                "节点与状态设计不当会造成过度工程化",
            ],
            [
                "多 Agent",
                "隔离不同领域的提示词和工具，并处理少量跨领域复合问题",
                "不会天然比单 Agent 更聪明，也不会自动消除冲突",
                "路由、合成、并行调用、成本和排障复杂度增加",
            ],
        ],
        [1450, 2850, 2850, 2210],
    )

    doc.add_heading("2. 技术选型面试问题与满分答案", level=2)
    questions = [
        (
            "如果面试官问“你为什么会想到引入这些框架”，你会怎样回答？",
            "我不会从框架名称开始讲，而会从项目逐步暴露的问题开始讲。最初只需要让模型回答问题，"
            "后来增加了知识库、外部数据、历史记忆、任务状态和持久化，单次模型调用已经无法清晰表达完整流程。"
            "因此我先用工具调用解决动态能力选择，再用 LangChain 统一模型和工具接口，用 LangGraph 管理跨节点状态与流程。"
            "多 Agent 是在发现领域提示词、工具权限和复合问题开始相互干扰后增加的可选方案。"
            "这不是一次性堆出技术栈，而是问题复杂度增长后逐层引入。",
            "能否先讲问题演进，再讲技术名称；是否把技术选型说成逐步决策。",
        ),
        (
            "为什么不直接把用户问题发给大模型，而要做成 Agent？",
            "直接调用大模型适合闲聊和通用问答，但家具资料、设备故障、用户报告和会话历史都属于模型上下文之外的数据。"
            "Agent 的价值是让模型在需要时选择知识库、报告或外部数据工具，再基于返回结果组织答案。"
            "我选择 Agent 不是为了让模型显得更智能，而是为了把模型的语言能力和系统内可验证的数据能力连接起来。"
            "对于可以由固定规则完成的步骤，我仍然使用确定性流程，不让模型随意决策。",
            "是否理解 Agent 的核心是受控工具使用，而不是模型自由发挥。",
        ),
        (
            "为什么选择 ReAct 思路？当时想解决什么问题？",
            "客服问题的处理步骤并不完全固定。用户可能只问家具保养，也可能在同一句话里要求查询知识、结合自己的使用报告再给建议。"
            "如果事先为每一种组合编写固定流程，分支会快速膨胀。ReAct 的思路允许模型先判断是否需要工具，"
            "看到工具结果后再决定继续调用还是输出答案，适合这种少量、受控、但组合不固定的工具使用场景。"
            "项目中我把这种自主性限制在回答节点内部，外层持久化和任务更新仍由确定流程控制。",
            "是否能说清 ReAct 适合动态工具组合，以及为什么要限制自主范围。",
        ),
        (
            "ReAct 和普通 Tool Calling 有什么区别？你的项目更准确的说法是什么？",
            "Tool Calling 是模型输出结构化工具调用请求的能力；ReAct 更强调“判断—行动—观察结果—继续判断”的循环。"
            "本项目的回答节点会绑定工具，执行模型提出的调用，把工具结果放回消息，再让模型继续判断，"
            "因此更准确地说是一个基于 Tool Calling 实现的有限 ReAct 式循环。"
            "我不会把模型内部不可见的思维过程当成系统能力来宣传，真正可观察的是工具名、参数、返回结果、调用次数和最终答案。",
            "是否区分底层能力与运行范式，是否避免宣称可读取模型思维链。",
        ),
        (
            "为什么不给 ReAct 无限次调用工具？",
            "无限循环会造成费用失控、延迟过高，甚至因为错误工具结果反复调用。"
            "项目中回答节点设置了有限迭代次数，达到上限后要求模型基于已有结果给出最终回答或明确说明信息不足。"
            "生产环境还应增加单工具超时、总耗时预算、最大工具调用数和重复调用检测。"
            "有限循环的本质是把 Agent 自主性放进工程边界，而不是完全交给模型。",
            "是否理解 Agent 必须有停止条件、预算和故障边界。",
        ),
        (
            "引入 ReAct 后，项目到底有什么提升？",
            "它最直接的提升不是模型知识变多，而是减少了手写工具组合分支，让同一个回答节点能处理不同的工具使用顺序。"
            "例如模型可以只查知识库，也可以先查用户报告再补充领域知识。结构上，新增工具的接入成本也更低。"
            "但如果没有做前后对照评测，我不会说准确率显著提升。要证明效果，需要比较任务完成率、正确工具选择率、"
            "平均工具调用次数、P95 延迟、单请求成本和拒答准确率。",
            "是否区分结构收益和效果收益，是否知道用数据而不是感受证明提升。",
        ),
        (
            "什么情况下你不会使用 ReAct？",
            "如果步骤固定、风险高或要求强一致，例如保存消息、扣减库存、更新订单状态，我不会让模型自由决定执行顺序。"
            "这类流程应使用明确的业务状态机、权限校验和数据库事务。"
            "简单的单工具查询也不一定需要循环，规则能稳定决定时直接调用更快、更便宜。"
            "我的原则是：只有决策确实需要结合自然语言和动态上下文时才交给 Agent。",
            "是否知道 ReAct 的适用边界，而不是所有步骤都交给模型。",
        ),
        (
            "为什么选择 LangChain？",
            "项目需要接入聊天模型、Embedding、Prompt、消息对象、工具、文档对象、输出解析和向量库。"
            "如果全部直接调用不同 SDK，会出现参数格式、消息结构和异常处理各不相同的问题。"
            "LangChain 提供了一套相对统一的接口，使模型替换、工具声明和 RAG 组件组合更方便。"
            "我选择它主要是为了降低集成成本和统一边界，而不是因为使用 LangChain 就能自动获得更好的答案。",
            "是否理解 LangChain 是组件与接口层，而不是智能本身。",
        ),
        (
            "为什么不直接调用大模型厂商 SDK？",
            "项目很小时直接调用 SDK 更简单，依赖也更少。这个项目同时使用模型消息、工具调用、文档、Embedding 和向量检索，"
            "统一抽象开始产生价值，所以采用 LangChain。"
            "但我不会把所有业务逻辑都写进框架链条，任务路由、事务、幂等和 Outbox 仍保留为普通业务模块。"
            "如果未来只有单一模型和少量接口，或者框架升级成本超过收益，也可以退回原生 SDK。",
            "是否做过规模判断，是否知道框架不应侵入全部业务逻辑。",
        ),
        (
            "引入 LangChain 后，项目有哪些可验证的提升？",
            "可以验证的主要是开发和维护效率，例如模型适配代码是否减少、工具接口是否统一、测试替身是否更容易注入、"
            "RAG 文档结构是否能够在组件之间复用。答案质量是否提升则不能只归因于 LangChain，"
            "因为质量更多取决于 Prompt、检索、数据和模型。面试中我会把“工程集成效率提升”和“模型效果提升”分开回答。",
            "是否能准确归因，不把所有改善都归功于框架。",
        ),
        (
            "LangChain 带来了什么代价？",
            "代价包括额外抽象、调用链更深、版本升级可能修改接口，以及异常有时需要穿过框架层才能定位。"
            "因此我会锁定依赖版本，在模型、工具和向量库外再保留自己的薄接口，并避免把核心业务对象完全绑定到框架类型。"
            "这样既能利用框架生态，又能在必要时替换局部实现。",
            "是否理解依赖治理、版本风险和框架隔离。",
        ),
        (
            "既然有 LangChain，为什么还要 LangGraph？",
            "LangChain 主要解决组件调用和组合，LangGraph 主要解决有状态流程编排。"
            "这个项目不仅要生成回答，还要并行做任务路由与历史召回，在汇合后组装上下文，"
            "再进入回答、持久化、任务更新和记忆刷新，并对不同节点设置不同的失败策略。"
            "这些步骤已经具有状态、分支、并行和汇合关系，用显式图比把流程散落在一个大函数里更容易理解、测试和观测。",
            "是否能区分组件层和流程层，是否能联系项目的菱形工作流。",
        ),
        (
            "为什么不用普通函数和 if/else 写工作流？",
            "普通函数完全可以完成第一版，而且简单流程应该优先这样做。"
            "当流程只有三四步时，引入图框架可能过度设计；但当项目出现并行分支、两种执行模式、节点级降级、共享状态和耗时统计后，"
            "if/else 会逐渐分散在多个位置，难以确认所有路径是否都经过持久化。"
            "LangGraph 的收益是把节点、边和状态集中表达。我的选择标准不是步骤数量本身，而是流程是否已经存在复杂分支和恢复需求。",
            "是否承认简单方案的价值，并说明复杂度达到什么程度才升级。",
        ),
        (
            "为什么不用 Celery、消息队列或者传统工作流引擎替代 LangGraph？",
            "它们解决的层次不同。LangGraph 管理一次请求内部的模型与工具状态流转；Celery 或消息队列更适合跨进程、可延迟执行的后台任务，"
            "例如本项目的索引更新 Worker。传统工作流引擎适合更长生命周期、跨服务和强审计流程。"
            "如果把一次对话中的每个小节点都拆成队列任务，延迟和运维成本会明显增加；"
            "如果把长期异步任务只放在 LangGraph 内存状态中，又缺少可靠投递。因此项目分别使用 LangGraph 和 Outbox Worker。",
            "是否能分清请求内编排和跨进程异步任务。",
        ),
        (
            "LangGraph 给项目带来的提升是什么？",
            "最明确的提升是可解释性、可测试性和故障隔离。任务路由与历史召回可以并行，"
            "每个节点能够记录耗时与错误，辅助节点失败时可返回降级状态，持久化失败时则可以明确中断。"
            "两种执行模式也可以共享前后的公共节点。"
            "它不会自动让答案更准确；答案质量仍需要用 RAG 与任务评测验证。LangGraph 改善的是系统组织方式和运行可观测性。",
            "是否准确描述编排收益，并避免把框架和答案质量混为一谈。",
        ),
        (
            "LangGraph 会不会是过度设计？你怎样判断？",
            "如果系统只是一次模型调用加一次知识库查询，我会认为是过度设计。"
            "判断是否值得要看是否有共享状态、多个分支、并行节点、失败策略、重试恢复和模式切换。"
            "这个项目已经具备这些需求，所以图编排有实际价值。"
            "不过我仍然会控制节点粒度：一个节点应该代表可独立测试、可观测或可降级的业务步骤，"
            "不能为了画图把每个小函数都拆成节点。",
            "是否能给出采用和不采用的条件，并理解合理节点粒度。",
        ),
        (
            "为什么要引入多 Agent，而不是一直使用单 Agent？",
            "当家具推荐、设备故障、使用报告和通用问答共享同一个 Agent 时，系统提示词会越来越长，"
            "工具权限和领域规则也可能相互干扰。多 Agent 可以让每个专家只关注自己的领域知识和工具，"
            "复合问题再由少量专家并行处理并合成。"
            "但多 Agent 不是默认答案。只有领域边界清楚、不同领域确实需要不同提示词或工具、"
            "而且单 Agent 评测出现明显瓶颈时，我才认为值得引入。",
            "是否从领域隔离和工具权限解释，而不是说多个 Agent 一定更聪明。",
        ),
        (
            "为什么不用一个 Agent 加很多工具？",
            "单 Agent 加多工具是更简单的基线，也是项目当前默认主链路。"
            "它的优势是调用次数少、上下文统一、延迟和排障成本更低。"
            "多 Agent 的优势只在于领域提示词隔离、工具集合缩小和复合问题并行。"
            "因此正确做法是先用单 Agent 建立基线，再用评测判断工具误选、领域混淆或复杂问题覆盖是否足以抵消多 Agent 的额外成本。",
            "是否把单 Agent 当成基线，并能进行真实取舍。",
        ),
        (
            "为什么最多只路由两个 Agent？",
            "用户的一次客服问题通常只有一个主领域，少数问题跨两个领域。"
            "如果允许四个专家全部运行，模型调用、延迟、重复内容和冲突都会快速增加。"
            "限制两个 Agent 是在覆盖复合问题和控制成本之间取平衡。"
            "路由器还需要去重，并让 GeneralAgent 与领域专家互斥，避免通用回答和专业回答重复竞争。",
            "是否理解并行度不是越高越好，以及去重和互斥设计。",
        ),
        (
            "引入多 Agent 后，项目真的提升了吗？",
            "目前项目已经实现并测试了多 Agent 路由、并行执行和答案合成，但默认生产路径仍是单 Agent 工具调用。"
            "因此我不会说多 Agent 已经让线上准确率显著提升。"
            "现在能确认的是它提供了领域隔离和复合问题处理能力；是否应该默认启用，需要通过单 Agent 与多 Agent 的 A/B 评测比较"
            "任务成功率、路由准确率、答案冲突率、P95 延迟和单请求成本。只有收益稳定超过代价才会切换默认路径。",
            "是否诚实说明当前运行边界，并知道怎样验证多 Agent 的真实价值。",
        ),
        (
            "多 Agent 路由错了怎么办？",
            "第一层用高精度规则处理明显关键词，低置信度时再让模型兜底；第二层限制候选专家数量并保留 GeneralAgent 回退；"
            "第三层让专家结果带上置信度、证据、覆盖点和未解决点，合成器不能盲目拼接。"
            "线上还应记录路由结果与用户反馈，构建误路由样本集。"
            "如果某类错误频繁出现，应先修正规则或训练路由评测集，而不是单纯加长 Prompt。",
            "是否具备路由纠错、回退与数据闭环意识。",
        ),
        (
            "两个 Agent 的答案冲突时怎么处理？",
            "合成器应优先比较证据质量和领域职责，而不是平均两种说法。"
            "同一事实如果证据冲突，应明确告诉用户存在不一致并请求更多信息，不能强行生成确定结论。"
            "可验证事实应回查数据库或知识库原文；建议性内容可以并列说明适用条件。"
            "同时把冲突类型记录下来，用于改进知识数据、路由规则和专家边界。",
            "是否理解答案合成不是文本拼接，而是证据与职责裁决。",
        ),
        (
            "什么情况下你会关闭多 Agent？",
            "如果大多数请求都属于单一领域，多 Agent 对任务成功率提升很小，却明显增加延迟和费用，我会关闭默认多 Agent。"
            "当路由置信度低、专家结果冲突率高，或者某个专家服务不稳定时，也可以降级为单 Agent 或 GeneralAgent。"
            "架构的目标是稳定解决用户问题，不是维持复杂形式。能够关闭一个昂贵但无收益的能力，也是一种工程判断。",
            "是否有成本和降级意识，是否愿意根据数据简化系统。",
        ),
        (
            "请完整讲一下这套技术选型的演进路线。",
            "我会把它概括为五步：第一步，直接模型调用验证客服场景；第二步，加入 RAG 和外部数据工具，"
            "用有限 ReAct 式循环处理动态工具选择；第三步，使用 LangChain 统一模型、消息、工具和文档接口；"
            "第四步，随着任务路由、历史召回、上下文组装、持久化和记忆更新增多，用 LangGraph 把状态流转显式化；"
            "第五步，在单 Agent 出现领域提示词和工具集合膨胀后，实现多 Agent 作为可评测、可切换的增强路径。"
            "每一步都应该由前一步的真实瓶颈触发。",
            "是否能把技术栈讲成有因果关系的演进，而不是同时堆砌。",
        ),
        (
            "如果让你重新做一次，你还会一开始就使用全部技术吗？",
            "不会。我会先做最小闭环：单 Agent、少量工具、基础 RAG 和可靠持久化，先建立评测集。"
            "只有当工具组合变得动态时引入 ReAct；当模型与检索组件增多时使用 LangChain；"
            "当流程出现并行、分支和降级时引入 LangGraph；当单 Agent 在领域隔离上有可量化瓶颈时再启用多 Agent。"
            "重新设计时我会更早加入链路追踪、成本统计和离线评测，因为没有这些数据，很难判断复杂架构是否真的有价值。",
            "是否具备从简单方案开始、用证据驱动升级的能力。",
        ),
        (
            "面试官说“你的项目像在堆框架”，你怎么回应？",
            "这个质疑是合理的，所以我会把每个框架对应到一个无法忽略的具体问题，并说明可删除条件。"
            "LangChain 用于统一模型和工具接口；LangGraph 用于并行、分支和节点级降级；"
            "ReAct 只负责回答节点中的动态工具选择；多 Agent 当前是可切换能力而不是默认主链路。"
            "如果评测证明某个组件没有改善质量、维护性或成本，我愿意删除它。"
            "技术选型的标准不是简历关键词数量，而是整体复杂度是否低于它解决的问题复杂度。",
            "是否能够接受质疑、给出边界，并体现删减技术的判断力。",
        ),
        (
            "你会用哪些指标证明这些技术选择有效？",
            "效果指标包括任务完成率、路由准确率、正确工具选择率、检索证据命中率、拒答准确率和用户反馈；"
            "性能指标包括首字延迟、P50/P95 总延迟、平均工具调用数、模型 Token 与单请求成本；"
            "工程指标包括节点失败率、降级成功率、重复请求比例、测试覆盖和故障定位时间。"
            "评测时应固定问题集、模型版本和知识库版本，对单 Agent、多 Agent、固定链路和 Agentic 链路做对照，"
            "否则无法把变化归因到某项技术。",
            "是否拥有完整评测观念，能同时衡量质量、性能、成本与工程性。",
        ),
    ]
    for number, (question, answer, focus) in enumerate(questions, start=1):
        add_interview_question(doc, number, question, answer, focus)

    doc.add_heading("3. 一分钟选型总结口径", level=2)
    add_callout(
        doc,
        "推荐表达",
        "我的技术选型是随着问题复杂度逐步演进的。ReAct 解决回答阶段动态选择工具的问题，LangChain 统一模型、消息、"
        "工具和检索组件接口，LangGraph 把任务路由、历史召回、上下文组装、回答和持久化组织成可观测的状态图。"
        "多 Agent 用来隔离家具、设备、报告等领域能力，但会增加延迟、成本和冲突，所以当前默认仍使用单 Agent 工具调用，"
        "多 Agent 作为可切换方案保留。每项技术是否继续使用，都应该通过任务成功率、延迟、成本和维护复杂度来验证。",
        fill="EAF7EF",
        color=GREEN,
    )


def add_enterprise_hot_questions(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("大厂高频补充专项：MCP、A2A、安全、可观测性与上下文工程", level=1)
    add_body(
        doc,
        "这一章补的是最近大厂 Agent 岗位里明显升温、但你前面手册还没有系统覆盖的题。"
        "它们共同反映一个趋势：面试官不再只问你会不会搭 Agent，而是开始追问你是否理解协议边界、运行时安全、"
        "链路观测、跨 Agent 协作和生产化治理。"
    )
    add_callout(
        doc,
        "为什么这些题会高频出现",
        "因为 2026 年很多大厂 Agent 岗位 JD 已经把 MCP、LangGraph、RAG、评测、规划和生产化治理写进岗位要求。"
        "同时技术社区的高频面经也开始从 ReAct、Tool Calling 继续下钻到 MCP、A2A、Prompt Injection、可观测性和安全确认。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("1. 先建立一张新考点地图", level=2)
    add_comparison_table(
        doc,
        ["考点", "面试官真正想听什么", "你应该如何回答"],
        [
            ["MCP", "你是否理解协议层和工具层的区别", "先讲它解决的是工具生态治理，不要只说“能连工具”"],
            ["A2A", "你是否理解 Agent 之间如何协作", "先区分 Agent 间协作和 Agent 调工具不是一回事"],
            ["Prompt Injection", "你是否有安全边界意识", "讲清数据流和控制流隔离、权限最小化、人工确认"],
            ["可观测性", "你是否真的有生产化思维", "说明至少记录问题、工具、参数、结果、延迟、成本和失败点"],
            ["上下文工程", "你是否理解 Agent 成败不只在提示词", "讲清什么信息该放上下文、什么必须结构化、什么不该喂给模型"],
            ["人工确认", "你是否知道哪些动作不能全自动", "强调高风险、不可逆、对外操作必须加确认节点"],
        ],
        [1600, 3600, 3920],
    )

    doc.add_heading("2. 大厂高频补充题与满分答案", level=2)
    questions = [
        (
            "MCP 是什么？为什么 2026 年大厂开始频繁问它？",
            "MCP 可以理解成面向 Agent 工具生态的标准协议。它解决的不只是“模型如何发起一次工具调用”，而是工具怎么被发现、怎么描述、怎么接入、怎么传输、怎么审计、怎么规模化管理。"
            "大厂开始频繁问，是因为 Agent 从 Demo 走向平台化之后，最大的问题不再是能不能调一个工具，而是几十上百个工具如何统一接入、复用和治理。"
            "所以如果面试官问 MCP，我会先回答它解决的是工具生态管理，再补充它常见的角色拆分、能力分类和传输方式，而不会只停留在‘比 Function Calling 更高级’这种空话。",
            "是否理解 MCP 的核心价值在于工具生态治理，而不只是一次调用成功。"
        ),
        (
            "MCP 和 Function Calling 的区别是什么？",
            "Function Calling 是模型输出层的一种约定，它解决的是‘模型如何把要调用的工具名和参数结构化表达出来’。"
            "MCP 是更外层的协议与生态约定，它关心的是工具如何被发现、描述、连接、权限化和审计。"
            "更通俗地说，Function Calling 解决一次怎么叫，MCP 解决长期怎么管。"
            "两者不是互相替代关系，很多 MCP 工具最终仍然会通过模型的结构化调用能力触发，但 MCP 额外补上了工程化接入、标准化描述和运行治理这一层。",
            "是否能把输出层能力和系统协议层能力分开。"
        ),
        (
            "如果工具已经能通过 Function Calling 调用了，为什么还要关注 MCP？",
            "因为当工具数量很少时，手写注册完全可行；一旦工具数量上来，问题就变成了版本管理、权限边界、描述一致性、可审计性和跨项目复用。"
            "MCP 的价值不在于让单个工具更聪明，而在于让工具接入从一次性代码工作变成可持续管理的标准化能力。"
            "如果我是做单项目小闭环，我未必急着引入 MCP；但如果目标是平台化、多人协作、跨团队复用或未来工具数持续增长，那就应该尽早理解 MCP 这层抽象。",
            "是否理解何时需要协议化，何时不必过度引入。"
        ),
        (
            "A2A 是什么？它和 MCP 是什么关系？",
            "A2A 可以理解成 Agent 与 Agent 之间协作的协议或模式，关注的是一个 Agent 如何把任务委托给另一个 Agent，或者如何交换状态、结果和协作上下文。"
            "MCP 面向的是 Agent 与工具、资源、提示模板的连接；A2A 面向的是多个 Agent 之间的协作。"
            "所以两者不在同一层。一个更像‘怎么接能力’，另一个更像‘怎么让多个智能体协同干活’。"
            "面试时我会明确说：不是做了多 Agent 就自动等于做了 A2A，也不是有了 MCP 就天然支持复杂多 Agent 协同，它们解决的是不同问题。",
            "是否能正确区分 Agent 调工具和 Agent 调 Agent。"
        ),
        (
            "你的项目目前没有 MCP 或 A2A，面试官追问时怎么回答才诚实又不显得弱？",
            "我会直接说明项目当前核心能力建立在 LangGraph、工具调用、RAG 和记忆系统上，重点解决的是单 Agent 到可切换多 Agent 的业务闭环，当前没有真正引入 MCP 或标准化 A2A 协议。"
            "但我理解它们分别适用于工具生态治理和跨 Agent 协作治理。"
            "如果项目未来要从单业务系统升级到 Agent 平台，我会优先考虑 MCP；如果未来真的出现稳定的跨 Agent 分工链路，再评估 A2A 或类似协作协议。"
            "这种回答的关键不是假装自己已经做过，而是表明你知道它们适合什么阶段、解决什么痛点。",
            "是否能够准确说明现状，同时展示扩展判断力。"
        ),
        (
            "Prompt Injection 到底是什么？为什么 Agent 比普通聊天更危险？",
            "Prompt Injection 的本质是外部输入试图篡改系统控制逻辑。普通聊天最多是让回答跑偏，但 Agent 一旦能调工具、读文件、发请求、执行动作，风险就会从‘答错’升级成‘做错’。"
            "所以 Agent 比普通聊天更危险，因为它不仅生成语言，还可能触达真实系统。"
            "我会把这件事解释成控制流和数据流混在一起时，用户输入就有机会污染系统指令；而当 Agent 具备外部执行能力时，这种污染就不只是文本风险，而是操作风险。",
            "是否理解 Agent 安全问题的爆炸半径。"
        ),
        (
            "如何从工程上降低 Prompt Injection 风险？",
            "第一层是控制流和数据流分离，用户输入、网页内容、工具返回结果都只能作为数据，不应该直接升级成系统命令。"
            "第二层是最小权限原则，Agent 不该默认拥有删除、付款、写生产库这类高风险能力。"
            "第三层是高风险动作强制人工确认，尤其是不可逆操作。"
            "第四层是工具结果不直接当指令执行，而是作为待解释的数据，再经过规则校验或安全节点。"
            "第五层是保留调用审计和异常告警，发生注入时能回溯是哪个输入、哪个工具、哪个环节出了问题。",
            "是否能从架构、权限和运行治理三个层面给出防御思路。"
        ),
        (
            "为什么说‘工具返回的是数据，不是命令’？",
            "这是 Agent 安全设计里特别关键的一句话。工具返回值本质上是外部世界给系统的一段数据，它可能不完整、被污染，甚至是恶意构造的。"
            "如果系统把工具返回值直接当成新的系统指令继续执行，相当于把外部输入抬升成控制流的一部分。"
            "正确做法是把工具返回视为证据或上下文，由模型或规则在受控边界内解释，再决定下一步。"
            "也就是说，工具能提供信息，但不应该越过边界指挥系统。",
            "是否真正理解数据和控制的边界。"
        ),
        (
            "哪些 Agent 操作必须人工确认？为什么？",
            "凡是高风险、不可逆、对外产生真实影响的动作，都应该加入人工确认节点。"
            "例如发送外部消息、执行支付、修改生产数据、删除文件、推送代码、触发高额付费 API。"
            "原因不是我们不相信模型，而是这些动作一旦出错，损失很难回滚。"
            "面试时我会强调不是所有操作都要人工审批，而是低风险高频操作自动化，高风险低频操作强制确认，这样才兼顾效率和安全。",
            "是否具备风险分级和人机协同思维。"
        ),
        (
            "Agent 系统最少应该做哪些可观测性记录？",
            "至少要记录用户原始问题、会话和请求标识、模型选了哪个工具、生成了什么参数、工具是否成功、工具耗时、工具返回了什么、最终回答是什么、整条链路耗时多少、用了多少 token、用户是否反馈有问题。"
            "如果是 RAG，还应该记录召回了哪些片段、重排结果如何、最终采用了哪些证据。"
            "因为没有这些记录，就很难分辨是模型选错工具、参数抽错、工具返回错，还是工具没问题但模型总结错。",
            "是否理解可观测性的目标是定位责任，而不是只留几条日志。"
        ),
        (
            "为什么说 Agent 的可观测性比传统后端更重要？",
            "传统后端里，HTTP 200 往往已经说明主流程成功了；但 Agent 场景里，HTTP 200 只代表系统返回了一段话，不代表这段话是对的，也不代表工具选得对、证据用得对。"
            "Agent 的错误经常不是崩溃式错误，而是静默质量错误。"
            "所以必须把模型输入、工具链路、检索证据、最终回答和用户反馈串成一条可回放的轨迹。"
            "没有这层观测，系统就只能靠猜，无法持续优化。",
            "是否知道 Agent 里的失败很多是‘看起来成功’。"
        ),
        (
            "上下文工程和提示工程有什么区别？",
            "提示工程更关注怎么写一句有效的指令；上下文工程更关注在模型真正推理前，到底给它什么信息、这些信息怎么组织、哪些应该结构化、哪些应该裁剪、哪些根本不该进入上下文。"
            "在生产系统里，Agent 成败往往不只取决于提示词文案，而更取决于上下文质量。"
            "比如任务状态、用户画像、最近对话、检索证据、工具返回、系统规则，这些都属于上下文工程的范围。"
            "所以我会把提示工程视为上下文工程的一部分，而不是全部。",
            "是否理解从‘写 Prompt’到‘设计上下文输入面’的升级。"
        ),
        (
            "如果上下文窗口有限，你优先保留什么？删除什么？",
            "我会优先保留对当前决策真正有影响的信息：当前任务目标、最近一两轮关键对话、明确的用户约束、最终被采信的证据片段、必须遵守的系统规则。"
            "我会优先删除重复信息、历史闲聊、未被采信的冗余检索片段、可以结构化存储后按需读取的大段原文。"
            "核心原则不是‘尽量多塞’，而是‘只保留会改变当前决策的信息’。"
            "因为过长上下文不仅贵，还会稀释重点、增加工具误选和注意力漂移。",
            "是否具备上下文预算意识和信息取舍能力。"
        ),
        (
            "规划和执行要不要拆成两个模型或两个阶段？",
            "这取决于任务复杂度。简单客服问答通常不值得把规划和执行彻底拆开，因为拆开会增加延迟、复杂度和状态同步成本。"
            "但当任务变成长链路、多工具、多步骤依赖时，先规划后执行会更稳定，也更容易检查每一步是否跑偏。"
            "我会说当前项目里的客服主链路更适合有限工具循环，而不是重规划型 Agent；如果以后扩展到长任务编排、报告生成或跨系统流程，才更值得考虑规划与执行分层。",
            "是否能根据任务形态选择架构，而不是套模板。"
        ),
        (
            "如果面试官问你‘怎么证明你的 Agent 系统真的比简单工作流强’，你怎么答？",
            "我会先建立基线，而不是凭感觉说更强。"
            "做法是准备同一批问题，对比固定工作流、单 Agent 工具调用、可选多 Agent 路径三种方案的任务完成率、工具误选率、首答延迟、平均 token 成本和人工复核正确率。"
            "如果 Agent 方案只带来了复杂度，却没有明显改善复杂问题覆盖或维护成本，那我不会强行说它更先进。"
            "真正有说服力的不是框架名字，而是对照评测结果。",
            "是否有基线思维，是否避免为了 Agent 而 Agent。"
        ),
    ]

    for number, (question, answer, focus) in enumerate(questions, start=1):
        add_interview_question(doc, number, question, answer, focus)

    doc.add_heading("3. 这一章最值得背下来的三句话", level=2)
    for item in [
        "Function Calling 解决一次怎么调，MCP 解决长期怎么管。",
        "工具返回的是数据，不是命令；高风险动作必须有权限边界和人工确认。",
        "Agent 成败很多时候不在提示词，而在上下文工程、链路观测和回归评测。",
    ]:
        add_bullet(doc, item)


def add_interview_bank(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("专项附录：按知识点整理的面试题总汇", level=1)
    add_body(
        doc,
        "这一部分把原来分散在正文里的项目问答、技术选型问答和大厂高频补充题统一收口到一个地方。"
        "你复习时可以把前面章节当成项目理解材料，把这里当成真正的面试题库来刷。"
    )
    add_callout(
        doc,
        "复习方式",
        "先遮住答案，用 1 至 3 分钟口述。回答时优先采用“问题背景 → 方案 → 为什么 → 边界与代价 → 如何验证”的结构。"
        "满分答案是表达框架，不建议逐字背诵。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_callout(
        doc,
        "四个高危表述",
        "第一，FOR UPDATE SKIP LOCKED 不是无锁，而是利用行锁跳过已锁记录；第二，Outbox 通常提供至少一次投递，"
        "不是天然 exactly-once；第三，辅助节点可以降级，但核心消息持久化失败通常必须中断；第四，项目实现多 Agent "
        "能力不等于默认链路正在运行多 Agent。高级面试官很可能专门抓这四点。",
        fill="FFF1F2",
        color=RED,
    )

    sections: list[tuple[str, list[tuple[str, str, str]]]] = [
        (
            "A. 项目全局、真实性与技术取舍",
            [
                (
                    "请用两分钟介绍这个项目，但不要堆框架名称。",
                    "这是一个面向家具和扫地机器人场景的智能客服系统，核心目标是让回答既有领域依据，又能延续用户正在处理的问题。"
                    "一次请求进入后，系统会识别当前任务、召回必要历史、组装上下文，再由模型选择知识库或外部数据工具。"
                    "知识侧通过混合 RAG 控制答案依据；数据侧用 MySQL 保存事实、Redis 做缓存、Chroma 做语义索引；"
                    "异步索引通过 Outbox 和 Worker 实现最终一致。项目重点不是单次生成效果，而是可追踪、可重试和可降级的完整链路。",
                    "能否从业务价值讲到技术闭环，而不是复述简历关键词。",
                ),
                (
                    "这个项目最难的三个问题是什么？",
                    "第一是上下文可信度：用户最新陈述、任务状态、长期记忆和语义历史可能冲突，需要明确优先级并验证历史来源。"
                    "第二是检索可信度：召回到相似内容不等于足以回答，所以需要混合召回、重排和证据门控。"
                    "第三是跨存储一致性：MySQL 与 Chroma 无法共享本地事务，因此采用事务性 Outbox、幂等消费和重试恢复。"
                    "如果谈工程落地，我还会补充延迟和成本控制是下一阶段重点。",
                    "是否真正理解项目难点，而非把功能列表当难点。",
                ),
                (
                    "为什么选择 LangGraph，而不是普通函数或一个 ReAct Agent？",
                    "如果只有一次模型调用和少量工具，普通函数更简单。这个项目同时存在任务识别、历史召回、上下文组装、"
                    "查询改写、工具循环、持久化和记忆更新，而且部分步骤可并行、部分步骤可降级。LangGraph 的价值是把状态、"
                    "节点依赖和分支显式化，便于测试、观测和扩展。代价是状态设计和并发合并更复杂，所以不能为了使用框架而使用框架。",
                    "框架选择是否由问题复杂度驱动。",
                ),
                (
                    "简历写全链路节点容错，是否意味着任何节点失败都能正常返回？",
                    "不能这样绝对表述。任务识别、语义召回、重排序或记忆刷新属于增强能力，可以使用默认路由、空历史或较弱排序降级。"
                    "但核心回答为空、用户身份校验失败、消息事务持久化失败等情况不应该伪装成功，否则会破坏事实一致性。"
                    "更准确的说法是：辅助节点支持隔离和降级，核心正确性节点采用失败快返。",
                    "能否区分可用性与正确性，识别简历中的过度承诺。",
                ),
                (
                    "如果让你重新做一遍，你会删掉什么复杂度？",
                    "我会先用单 Agent 加明确工具完成主链路，把多 Agent 作为可切换实验能力，而不是默认引入。"
                    "只有当评测证明不同领域提示词和并行专家显著提高质量时才启用多 Agent。"
                    "我还会更早建立离线评测、调用链追踪和延迟预算，避免先堆模块再寻找收益。",
                    "是否具备简化系统和用数据做架构决策的能力。",
                ),
            ],
        ),
        (
            "B. LangGraph、多 Agent 与答案合成",
            [
                (
                    "你说的菱形图工作流具体是什么？为什么是菱形？",
                    "入口后有两个相互独立的前置节点：任务路由和历史召回，它们可以并行执行；两者完成后汇合到上下文组装节点，"
                    "形成第一个菱形。回答持久化完成后，任务状态更新和记忆刷新又可以并行，形成第二个分叉汇合结构。"
                    "这种结构的收益是降低串行等待，前提是共享状态字段有明确的合并规则。",
                    "是否能把图结构说成具体依赖，而不是只会说 DAG。",
                ),
                (
                    "Legacy 与 Agentic 两种模式分别解决什么问题？",
                    "Legacy 模式强调固定、可预测的执行链，适合已有单 Agent 或专家路由方案的兼容与回退；Agentic 模式让模型"
                    "根据问题自主选择工具并循环观察结果，灵活性更高。两种模式应复用相同的输入状态、持久化和后处理契约，"
                    "切换只影响回答生成节点。实际生产中还应通过配置或灰度实验控制，而不是运行时随意变化。",
                    "是否理解模式切换的边界和接口兼容。",
                ),
                (
                    "多 Agent 为什么最多只路由两个？",
                    "这是一种质量、延迟和成本之间的约束。候选专家越多，调用费用和尾延迟越高，答案冲突也更难处理。"
                    "大部分复合问题通常由一个主领域和一个辅助领域覆盖，因此限制两个可以控制复杂度。这个数字不应凭感觉确定，"
                    "应通过复合问题数据集观察覆盖率、延迟和合成质量后校准。",
                    "是否能够说明限制的工程依据与验证方法。",
                ),
                (
                    "多个专家并行返回冲突答案，合成器应该怎么办？",
                    "合成器不能简单拼接。首先应保留每个结果的领域、证据、置信度和拒答原因；其次按证据质量和领域归属确定主结论。"
                    "对于可以并存的内容分点合并，对于事实冲突则明确指出冲突并优先选择可验证证据更强的一方。"
                    "如果冲突涉及安全或关键事实，应该拒绝武断结论并请求补充信息，而不是让另一次 LLM 调用随意裁决。",
                    "是否有证据驱动的冲突解决意识。",
                ),
                (
                    "并行调用一定比串行调用快吗？",
                    "不一定。并行可以把总耗时从多个调用之和降低到接近最慢调用，但会增加线程或协程调度、连接池压力和下游限流风险。"
                    "如果第二个调用依赖第一个结果，或者 QPS 已接近模型服务限额，并行反而可能恶化 P95。"
                    "需要设置超时、并发上限、取消策略，并通过实际延迟分布验证。",
                    "是否理解并发的收益、条件和资源代价。",
                ),
                (
                    "如何防止 Agent 无限调用工具？",
                    "要同时设置最大迭代次数、最大工具调用数、单工具超时和全请求截止时间。工具结果需要以结构化状态回传，"
                    "检测重复的工具名与参数，必要时触发熔断或强制进入最终回答。还应限制工具权限和参数范围，"
                    "避免模型通过循环调用造成成本失控或副作用重复。",
                    "是否具备 Agent 运行时治理意识。",
                ),
            ],
        ),
        (
            "C. 两级路由与置信度设计",
            [
                (
                    "为什么先用规则路由，再让 LLM 兜底？",
                    "规则对高频、边界清晰的关键词场景速度快、成本低、结果可解释；LLM 更适合模糊表达、复合意图和长上下文判断。"
                    "先规则后 LLM 可以让简单请求快速通过，只把低置信度或冲突请求交给模型。"
                    "关键是设置明确的兜底条件和评测集，否则两层路由可能互相掩盖错误。",
                    "是否理解混合路由的成本与准确性权衡。",
                ),
                (
                    "关键词加权打分如何避免“沙发故障”被误路由到设备 Agent？",
                    "不能只看到“故障”就路由设备，需要把对象词和意图词分开建模。设备问题至少应命中设备实体，"
                    "或者在不存在家具实体时才允许通用故障词提升设备分数。还可以加入互斥规则、短语优先级和负向特征。"
                    "最终要用混淆矩阵观察 Furniture 与 Device 的误路由情况。",
                    "是否能处理复合关键词与类别歧义。",
                ),
                (
                    "LLM 返回的 confidence 真的是概率吗？",
                    "通常不是经过校准的概率，只是模型自报分数，不能直接解释为 90% 正确。可以把它当排序特征，"
                    "但阈值应通过验证集校准，并检查可靠性曲线、分桶准确率或使用温度缩放等方法。"
                    "如果没有校准数据，我会明确称它为启发式置信分，而不是统计概率。",
                    "是否理解模型置信度和概率校准。",
                ),
                (
                    "路由系统应该用什么指标评估？",
                    "单标签场景看每个专家的 precision、recall、F1 和混淆矩阵；多标签复合问题还要看 micro/macro F1、"
                    "exact match、每请求平均路由数和专家覆盖率。工程指标包括路由耗时、LLM 兜底比例、错误路由带来的下游成本，"
                    "以及最终答案质量，而不能只看路由标签本身。",
                    "是否能建立端到端而非单模块评测。",
                ),
                (
                    "规则和 LLM 路由结果冲突时听谁的？",
                    "高置信度且有明确实体命中的规则可以直接通过；低置信度规则才进入 LLM。若两者都给出高置信但冲突，"
                    "应记录冲突并采用预先定义的优先级，例如安全领域或强实体规则优先，也可以选择两个专家并行验证。"
                    "不能临时凭代码顺序决定，需要通过离线误差分析制定策略。",
                    "是否具备冲突治理和可观测性设计。",
                ),
            ],
        ),
        (
            "D. 混合 RAG、重排序与证据可信度",
            [
                (
                    "请完整讲一遍你的 RAG 链路。",
                    "先对查询做归一化，提取对象、意图、材料、数字、型号和错误码等约束；随后同时执行向量召回和 BM25 关键词召回。"
                    "两路结果按文档片段 ID 去重，通过 RRF 融合排名，再对前若干候选使用 CrossEncoder 重排序。"
                    "最后检查对象、意图、必需术语、分数和版本有效性，证据足够才进入生成，否则返回拒答。",
                    "能否清晰说明每一层解决的不同问题。",
                ),
                (
                    "向量检索和 BM25 各自可能失败在哪些问题上？",
                    "向量检索可能把语义相似但对象不同的文本拉近，对型号、错误码和精确数字也可能不敏感。"
                    "BM25 依赖字面重合，遇到同义改写或自然语言表达变化容易漏召回。"
                    "因此双路召回强调互补，但融合后仍需要元数据过滤和重排序防止错误候选进入生成。",
                    "是否理解混合检索不是简单叠加。",
                ),
                (
                    "RRF 是什么？为什么不直接相加向量分和 BM25 分？",
                    "RRF 根据候选在各路结果中的排名计算分数，常见形式是各路 1 除以 k 加排名后求和。"
                    "向量距离和 BM25 分数不在同一尺度，直接相加需要复杂归一化，而 RRF 对不同分数分布更稳健。"
                    "它的缺点是忽略原始分数间隔，必要时可以引入通道权重或学习排序。",
                    "是否真正理解融合算法及其局限。",
                ),
                (
                    "Bi-Encoder 和 CrossEncoder 的区别是什么？",
                    "Bi-Encoder 分别编码查询和文档，文档向量可以预计算，适合大规模快速召回；CrossEncoder 将查询和文档成对输入模型，"
                    "能够建模更细粒度的交互，排序更准但计算昂贵。因此典型架构是 Bi-Encoder 召回较多候选，"
                    "CrossEncoder 只重排前十到几十条。",
                    "是否理解向量召回与重排模型的计算差异。",
                ),
                (
                    "证据可信度判定应该看哪些因素？",
                    "至少包括检索相关性、对象和意图覆盖、用户问题中的硬约束是否出现、资料版本是否有效、"
                    "多个证据是否互相矛盾，以及来源是否属于允许领域。对于错误码、剂量、温度和维修步骤等高风险信息，"
                    "应采用更严格阈值。不能只用一个相似度分数决定是否回答。",
                    "是否具备多维证据评估思维。",
                ),
                (
                    "怎么评估 RAG，而不是凭感觉看几个回答？",
                    "检索侧准备带相关片段标注的查询集，评估 Recall@K、MRR 或 NDCG，并分家具、设备、错误码和跨类别负样本统计。"
                    "生成侧评估 faithfulness、答案相关性、关键点覆盖和拒答准确率，重要样本由人工复核。"
                    "系统侧还要跟踪 P50/P95 延迟、失败率、空召回率和每次请求成本。",
                    "是否能够建立检索、生成和系统三层评测。",
                ),
                (
                    "如果知识库里没有答案，为什么不让模型使用常识补充？",
                    "项目定位是专业客服，错误维修步骤的成本高于拒答。模型常识缺少版本、产品和材料边界，"
                    "容易把一个对象的经验迁移到另一个对象。证据不足时应明确说明缺少什么，并建议补充型号、查看说明书或联系人工客服。"
                    "如果产品允许常识回答，也应把它标记为独立模式而不是伪装成知识库结论。",
                    "是否理解拒答策略的业务价值。",
                ),
                (
                    "如何防止知识库中的 Prompt Injection？",
                    "检索文档只能作为不可信数据，系统提示必须明确文档内容不是指令。入库时可以扫描可疑指令模式，"
                    "生成时用结构化边界包裹证据，并禁止证据修改系统规则或调用权限。高风险工具还需要独立授权和参数校验，"
                    "不能因为检索文本写着“调用某工具”就执行。",
                    "是否具备 RAG 安全意识。",
                ),
            ],
        ),
        (
            "E. 任务路由、幂等、记忆与历史召回",
            [
                (
                    "为什么要单独设计任务路由，聊天历史不够吗？",
                    "聊天历史记录的是原始对话，任务状态表达的是当前正在解决的问题，例如目标、约束、已尝试动作、结果和下一步。"
                    "长对话中用户可能切换多个问题，单纯拼接历史会把旧问题带进当前回答。任务路由可以决定新建、继续、恢复或不绑定任务，"
                    "让上下文只包含相关状态。",
                    "是否理解对话记忆与业务状态的区别。",
                ),
                (
                    "多阶段任务决策树为什么先过滤闲聊？",
                    "闲聊不应创建或污染业务任务。如果先做语义匹配，“谢谢”“你好”可能因为最近上下文而错误绑定到旧任务。"
                    "因此先识别高确定性的闲聊，再判断延续词、实体、意图和语义相似度，可以降低状态污染。",
                    "是否理解决策顺序会影响状态正确性。",
                ),
                (
                    "request_id 幂等到底保证了什么？",
                    "它保证同一业务请求因重试执行多次时，不会重复插入一轮消息、重复创建任务事件或重复触发同一副作用。"
                    "request_id 必须与数据库唯一约束配合，单靠应用层先查再写仍然有并发竞态。"
                    "它不保证不同 request_id 的重复语义请求去重，也不等于全系统 exactly-once。",
                    "是否理解幂等键的作用范围。",
                ),
                (
                    "乐观锁为什么适合任务状态？冲突后怎么办？",
                    "任务更新冲突通常不频繁，因此没有必要长时间持有数据库锁。读取任务时得到 state_version，更新时要求版本仍然一致；"
                    "如果受影响行数为零，说明有并发修改。此时重新读取最新状态，判断补丁是否还能安全应用，再有限次数重试；"
                    "不能无限覆盖，否则会丢失另一个请求的更新。",
                    "是否理解乐观并发控制和冲突处理。",
                ),
                (
                    "三层记忆中的信息发生冲突，应该相信谁？",
                    "优先级通常是用户本轮明确陈述最高，其次是 MySQL 中可定位的精确历史和当前任务事实，"
                    "再到最近对话、长期画像和语义召回。长期记忆需要来源、时间和置信度，用户明确否认后应更新或撤销旧记忆。"
                    "语义召回只能作为线索，不能覆盖当前输入。",
                    "是否能够处理记忆污染和事实更新。",
                ),
                (
                    "精确召回和语义召回为什么要分两条通道？",
                    "“我第一轮说了什么”要求确定性，应直接按会话和消息顺序查询 MySQL；“之前聊过类似问题吗”适合向量语义搜索。"
                    "如果所有历史问题都走向量检索，精确问题可能返回相似但错误的消息。语义候选还应回 MySQL 验证用户归属和原始内容。",
                    "是否理解确定性查询与近似检索的边界。",
                ),
            ],
        ),
        (
            "F. Outbox、Worker 与分布式一致性",
            [
                (
                    "事务性 Outbox 解决了什么问题？",
                    "它解决 MySQL 业务数据提交成功，但向量索引任务因进程崩溃或 Chroma 故障而永久丢失的问题。"
                    "消息和 Outbox 记录在同一个 MySQL 事务中提交，之后 Worker 异步处理。"
                    "它保证任务不会在业务事务之外凭空丢失，但不保证外部系统立即成功。",
                    "是否理解双写问题和 Outbox 的原子边界。",
                ),
                (
                    "Outbox 能保证 exactly-once 吗？",
                    "通常不能。Worker 可能已成功写入 Chroma，但在标记 Outbox 完成前崩溃，任务会再次执行，因此实际是至少一次投递。"
                    "要获得正确结果，消费者必须幂等，例如使用稳定向量 ID 做 upsert，完成状态更新也要带条件。"
                    "所谓 exactly-once 往往是业务效果幂等，而不是物理执行一次。",
                    "是否识别至少一次投递和 exactly-once 的区别。",
                ),
                (
                    "FOR UPDATE SKIP LOCKED 为什么不是无锁消费？",
                    "它仍然使用数据库行级悲观锁。SKIP LOCKED 的作用是其他 Worker 遇到已锁记录时不等待，而是跳过并领取下一批，"
                    "从而提高并发吞吐。更准确的描述是“基于行锁的非阻塞并发领取”，而不是无锁。",
                    "是否能主动纠正简历中的技术用词。",
                ),
                (
                    "Worker 处理到一半崩溃，任务如何恢复？",
                    "领取任务时将状态改为 processing 并记录 locked_at。监控或 Worker 启动时扫描超过阈值仍处于 processing 的任务，"
                    "将其恢复为 pending。由于任务可能已经执行过一部分，恢复后的处理必须是幂等的，"
                    "并且阈值要大于正常最长处理时间，避免误判仍在执行的任务。",
                    "是否理解租约、僵死任务和幂等恢复。",
                ),
                (
                    "指数退避为什么需要抖动？",
                    "纯指数退避会让同时失败的一批任务在相同时间再次重试，形成惊群。加入随机抖动可以打散重试时间，"
                    "减轻下游恢复瞬间的压力。还应设置最大间隔、最大重试次数和 dead-letter 状态，"
                    "并区分可重试错误与永久错误。",
                    "是否具备生产级重试策略知识。",
                ),
                (
                    "如何监控 Outbox Worker？",
                    "至少监控 pending 数量、最老任务年龄、processing 超时数、重试分布、dead 数量、单批处理耗时和成功率。"
                    "还要关联业务 request_id 与任务 ID，方便从一次用户请求追踪到索引结果。"
                    "告警应关注积压时间和增长趋势，而不只是瞬时任务数。",
                    "是否具备可观测性和运维思维。",
                ),
            ],
        ),
        (
            "G. 生产场景题与系统设计题",
            [
                (
                    "场景：Chroma 全部不可用，但 MySQL 正常，你怎么保证服务？",
                    "写入侧继续把核心消息和 Outbox 任务提交到 MySQL，让索引任务积压等待恢复。读取侧如果业务允许，可以退化到 BM25；"
                    "如果没有足够证据则拒答，不能让模型伪造知识。健康状态应显示 degraded，并对积压年龄告警。"
                    "Chroma 恢复后 Worker 按批次重放并校验索引数量。",
                    "是否能同时处理读降级、写可靠性和恢复流程。",
                ),
                (
                    "场景：Redis 不可用，会发生什么？",
                    "Redis 只是缓存，读取近期历史和用户记忆时回退 MySQL，写缓存失败只记录告警，不影响事实落库。"
                    "需要防止 Redis 故障导致所有请求同时打向 MySQL，可以设置连接超时、熔断、限流和本地短缓存。"
                    "恢复后按需回填缓存，而不是把 Redis 当权威数据。",
                    "是否理解缓存故障与缓存击穿。",
                ),
                (
                    "场景：LLM 调用 30 秒仍未返回，你怎么处理？",
                    "设置模型单次超时和全请求 deadline，超时后取消未完成的专家或工具任务。"
                    "如果已有高质量结构化结果，可以生成受限降级回答；否则返回可重试错误。"
                    "还应记录模型、阶段、token 数和 request_id，用于分析 P95 延迟，并考虑更小模型兜底或异步任务模式。",
                    "是否具备超时预算和尾延迟治理能力。",
                ),
                (
                    "场景：两个专家一个建议自行维修，一个建议立即停用联系售后，怎么合成？",
                    "先比较证据来源、对象和风险等级。涉及人身、设备损坏或保修风险时采用安全优先原则，"
                    "保留停止使用和联系售后的结论，并明确说明证据冲突，不应平均两种建议。"
                    "同时记录冲突样本进入评测集，修正领域提示词或路由。",
                    "是否能处理高风险冲突而非简单文本融合。",
                ),
                (
                    "场景：QPS 增长十倍，最先可能出现哪些瓶颈？",
                    "可能包括模型并发限额、每请求重复构建 BM25、数据库连接池、CrossEncoder CPU/GPU、线程池和 Chroma 查询。"
                    "我会先通过链路追踪定位各阶段 P95，再缓存 BM25、批量 embedding、限制并发、拆分 Worker、扩展连接池并做负载测试。"
                    "不能在没有数据时直接扩大所有组件。",
                    "是否能以测量驱动容量优化。",
                ),
                (
                    "场景：用户要求删除自己的全部数据，系统要删哪些地方？",
                    "需要删除或匿名化 MySQL 中的用户、会话、消息、任务、记忆和反馈，并清理 Redis 缓存。"
                    "Chroma 中知识索引不受影响，但该用户的历史对话向量必须按 user_uuid 删除；相关 Outbox 任务也要取消或变成删除任务。"
                    "还要考虑日志、备份保留策略和删除审计，避免只删主表却保留可识别副本。",
                    "是否具备隐私、数据生命周期和跨存储删除意识。",
                ),
            ],
        ),
        (
            "H. AI 基础与通用 Agent 基础",
            [
                (
                    "Temperature 和 top_p 分别控制什么？客服系统怎么设置？",
                    "Temperature 调整 logits 分布的平滑程度，top_p 只从累计概率达到阈值的候选集合中采样。"
                    "两者都影响随机性，通常不建议同时大幅调整。专业客服和结构化路由更重视稳定性，"
                    "一般使用较低 temperature；创意表达可以适当提高，但事实仍应由检索证据约束。",
                    "是否理解生成采样参数与业务目标。",
                ),
                (
                    "余弦相似度是什么？它高就一定相关吗？",
                    "余弦相似度比较两个向量方向，常用于语义向量检索。高相似只说明 embedding 空间中接近，"
                    "不保证对象、时间、型号和意图都一致，也不等于事实正确。"
                    "因此需要元数据过滤、关键词约束、重排序和业务证据判断。",
                    "是否理解向量相似度的含义和局限。",
                ),
                (
                    "RAG 和微调怎么选？",
                    "需要更新私有知识、来源可追踪和快速删除时优先 RAG；需要稳定改变模型风格、格式或任务行为时可考虑微调。"
                    "微调不会天然让模型可靠记住频繁变化的知识，也无法替代实时数据查询。"
                    "很多系统是 RAG 提供事实，提示词或微调改善行为。",
                    "是否理解知识注入与行为调整的差异。",
                ),
                (
                    "工具调用和让模型输出一段 JSON 有什么区别？",
                    "工具调用通常由模型服务提供结构化的工具名称和参数协议，并能与工具结果消息形成完整对话状态；"
                    "普通 JSON 只是文本约定，可能格式错误或夹带额外内容。无论哪种方式，服务端都必须做 schema 校验、权限检查和参数限制，"
                    "不能相信模型输出天然安全。",
                    "是否理解结构化输出、执行边界与安全校验。",
                ),
                (
                    "真正的 token 流式和把完整答案逐字发送有什么区别？",
                    "真正流式是在模型生成 token 时立即向用户发送，能够降低首 token 延迟；逐字发送完整答案只是改善视觉效果，"
                    "首字仍要等待模型、工具和后处理全部完成。真正流式还需要考虑回答未持久化、客户端断开、工具调用阶段无 token、"
                    "以及生成失败后的补偿策略。",
                    "是否理解流式体验背后的执行语义。",
                ),
                (
                    "Agent 与固定工作流的边界是什么？",
                    "固定工作流适合步骤稳定、合规要求高和副作用明确的业务；Agent 适合需要根据自然语言动态选择信息和工具的部分。"
                    "生产系统通常采用外层确定性工作流加内层受限 Agent，而不是让模型控制所有步骤。"
                    "数据库提交、权限、金额和删除等关键操作应由确定性代码控制。",
                    "是否具备受控 Agent 的系统观。",
                ),
            ],
        ),
        (
            "I. 评测、可观测性与工程成熟度",
            [
                (
                    "如何证明这个项目真的抑制了幻觉？",
                    "先构建包含可回答、证据不足、跨类别干扰、错误码和冲突资料的标注集。比较无 RAG、单路 RAG 和完整链路的"
                    "事实一致性、拒答准确率、关键点覆盖率与人工错误率。还要对错误进行分类，"
                    "例如检索漏召回、错误候选、生成越界和错误拒答。没有对照实验不能只凭主观案例声称有效。",
                    "是否能用实验支持简历结论。",
                ),
                (
                    "你会为整条 Agent 链路记录哪些 trace？",
                    "每次请求记录 request_id、用户和会话的脱敏标识、各节点开始结束时间、路由结果、工具名称、"
                    "检索候选数量与分数、模型和 token 使用量、降级原因、持久化结果和最终状态。"
                    "敏感正文应按策略脱敏或采样，trace 需要能从 API 追踪到 Outbox Worker。",
                    "是否具备端到端可观测性设计。",
                ),
                (
                    "离线评测很好，但线上质量如何监控？",
                    "线上关注用户点踩、转人工率、重复追问率、拒答率、工具失败率和任务完成率，并按路由类别和模型版本切分。"
                    "通过抽样人工复核和回放形成新的困难样本，加入离线集。上线新策略时做灰度或 A/B，"
                    "同时设置质量、延迟和成本护栏。",
                    "是否理解离线集与线上反馈闭环。",
                ),
                (
                    "如何控制 Agent 成本？",
                    "先减少不必要调用：规则命中时不调用路由 LLM，无需历史时不做向量召回，简单问题不启用多 Agent。"
                    "其次控制上下文长度、候选数、工具迭代数并使用小模型处理改写和抽取。"
                    "最后按请求记录 token、模型和工具成本，优化单位有效回答成本，而不是只看单价。",
                    "是否有系统性的成本治理思路。",
                ),
                (
                    "如果面试官问这个项目是不是 AI 写的，你怎么回答？",
                    "我会诚实说明 AI 辅助完成了大量初版实现，但不会把生成代码等同于项目完成。"
                    "我负责通过运行测试、梳理调用链、核对真实执行模式、发现部署和性能问题，并逐步亲自修改关键模块。"
                    "我能够说明每个设计解决的问题、边界和替代方案，也会明确哪些优化尚未完成。",
                    "技术诚信、项目所有权和真实学习能力。",
                ),
            ],
        ),
        (
            "J. 大模型微调、训练与模型选型",
            [
                (
                    "在你的这个项目里，优先应该做微调，还是先做 RAG 和工程优化？为什么？",
                    "我会明确说当前阶段优先级是 RAG、路由、上下文工程、评测和工具链治理，而不是直接做微调。"
                    "因为这个项目的核心问题主要来自私有知识接入、工具选择、历史记忆和证据可信度，这些问题微调并不能直接解决。"
                    "微调更适合稳定输出风格、结构化格式、领域术语表达或特定决策边界，但前提是问题已经被定位为模型参数层问题，而不是检索、上下文或流程层问题。"
                    "如果基础链路都没稳定，先微调往往只是把系统问题藏起来。",
                    "是否理解微调不是大多数 Agent 项目的第一优先级。"
                ),
                (
                    "什么情况下你会考虑给这个项目做微调？",
                    "当我已经通过评测确认，检索证据是对的、上下文也是对的、工具调用链路也稳定，但模型在固定类型任务上仍然持续出现相似偏差时，我才会认真考虑微调。"
                    "例如固定的客服语气风格、统一的售后话术模板、稳定的结构化抽取格式、特定品类的意图分类，或者规则很稳定但 Prompt 很难长期压住漂移时。"
                    "换句话说，微调适合解决稳定重复的问题，不适合替代缺失的外部知识和工程治理。",
                    "是否知道微调的适用边界。"
                ),
                (
                    "预训练、指令微调、继续预训练和强化学习分别在干什么？",
                    "预训练是在海量通用语料上学习语言和世界知识；继续预训练是在更窄领域语料上补领域分布；指令微调是让模型更好地遵循任务指令和输出格式；强化学习更偏向按目标反馈优化行为偏好。"
                    "面试时我不会把它们混成一个概念。"
                    "对于业务项目来说，最常见的是使用成熟基础模型，再通过 Prompt、RAG、少量指令数据、评测和工程策略做适配，而不是从头训练大模型。",
                    "是否分得清不同训练阶段的目标。"
                ),
                (
                    "LoRA、全量微调、Prefix/Adapter 这类参数高效微调方式有什么区别？",
                    "全量微调会更新几乎所有参数，效果上限高但成本、显存和部署复杂度也最高。"
                    "LoRA 这类参数高效微调只训练少量低秩增量参数，训练和部署成本更低，适合大多数业务场景。"
                    "Prefix、Adapter 也是类似思路，本质都是尽量少改主模型，用更小代价换取可控适配。"
                    "如果面试官问选型，我会说除非有非常强的资源和收益证明，否则业务团队通常先从参数高效微调开始。",
                    "是否理解微调方案的成本差异。"
                ),
                (
                    "SFT 数据应该怎么构造？最怕什么问题？",
                    "SFT 数据首先要和目标任务强相关，标签风格要稳定，输入输出边界要明确。"
                    "最怕的问题是脏数据、互相矛盾的数据、为了追求数量而牺牲一致性，以及把检索错误、工具错误、上下文缺失这种系统问题错误地塞进微调数据里。"
                    "如果训练数据本身不干净，模型只会学到不稳定行为。"
                    "所以我会优先做数据抽样审查、去重、模板一致性检查和失败样本分类，再决定哪些值得进入训练集。",
                    "是否具备训练数据质量意识。"
                ),
                (
                    "微调完成后应该怎么评估，而不是只看 loss 下降？",
                    "loss 下降只能说明训练集拟合变好了，不能说明业务效果一定更好。"
                    "我会保留独立验证集和测试集，比较微调前后的任务完成率、结构化输出准确率、拒答表现、工具参数抽取准确率、风格一致性和线上成本。"
                    "如果是客服项目，还要看是否引入新的幻觉、是否压低了泛化能力。"
                    "最终是否上线，应该由离线评测、灰度表现和人工复核共同决定。",
                    "是否有完整的训练后评测观念。"
                ),
                (
                    "为什么很多业务项目微调后反而效果变差？",
                    "常见原因有四类。第一，数据量不够但分布太窄，导致模型过拟合。"
                    "第二，训练数据质量不稳定，把错误行为也学进去了。"
                    "第三，业务真正缺的是实时知识和系统边界，而不是参数能力。"
                    "第四，评测方式不对，只看少量样例或只看训练损失，没有做完整对照。"
                    "所以我会把微调当成假设验证，而不是默认正确的升级路径。",
                    "是否理解微调失败通常不是模型本身的问题。"
                ),
                (
                    "如果面试官问你有没有做过大模型训练，你应该怎么诚实回答？",
                    "如果没有真正做过预训练或完整微调流水线，我会明确说没有把简历项目说成自己训练过基础模型。"
                    "更诚实的回答是：我理解训练与微调的基本目标、数据准备、评测方法和适用边界，但当前项目主要工作在 Agent 编排、RAG、记忆、路由和工程化治理上。"
                    "如果我做过少量 SFT 或格式微调实验，我会如实说明实验范围、数据规模和结果，不夸大成训练了一个行业模型。",
                    "是否具备技术诚信。"
                ),
                (
                    "为什么说很多 Agent 项目更像系统工程，而不是训练工程？",
                    "因为它们的大部分效果差异来自检索质量、上下文组织、工具设计、权限边界、链路观测、错误恢复和评测闭环，而不是重新训练参数。"
                    "在这类项目里，模型更像大脑的一部分，但系统真正可用，靠的是外部知识、业务约束和工程护栏。"
                    "所以我会把训练能力看成加分项，但不会把它误说成所有 Agent 项目的核心矛盾。",
                    "是否理解 Agent 项目和模型训练项目的主战场不同。"
                ),
            ],
        ),
        (
            "K. Agent 沙箱、权限边界与危险操作治理",
            [
                (
                    "什么是 Agent 沙箱？为什么它在生产环境里很重要？",
                    "Agent 沙箱可以理解成给智能体划出来的一块受限制运行空间。"
                    "它的核心目的不是让 Agent 更聪明，而是让 Agent 即使出错，也只能在受控范围内出错。"
                    "一旦 Agent 具备读写文件、执行命令、调用外部 API 或操作内部系统的能力，没有沙箱就等于把模型错误直接放大成系统风险。"
                    "所以沙箱本质上是安全边界和爆炸半径控制，而不是性能优化。",
                    "是否理解沙箱的作用是限制破坏范围。"
                ),
                (
                    "沙箱一般要限制哪些东西？",
                    "通常至少限制四类能力。第一是文件系统访问范围，例如只允许访问指定工作目录。"
                    "第二是网络访问范围，例如只允许访问白名单域名或特定内网服务。"
                    "第三是命令执行范围，例如只开放少量安全命令，而不是任意 shell。"
                    "第四是凭证与身份范围，例如不同 Agent 只能拿到最小必要的 token 和账号权限。"
                    "如果这些边界没有拆清，沙箱就只是形式存在。",
                    "是否能把沙箱具体化成文件、网络、命令、身份四层边界。"
                ),
                (
                    "为什么说‘最小权限原则’是 Agent 系统的基本原则？",
                    "因为 Agent 的错误经常不是不会回答，而是会做错事。"
                    "最小权限原则的意思是：一个 Agent 只拿到完成当前任务所必需的最少能力。"
                    "比如客服 Agent 不应该默认拥有删库、发公告、批量改用户资料的权限；只读检索 Agent 不该能写生产系统。"
                    "权限越小，错误越容易被局部化，安全审计也越清楚。",
                    "是否理解权限边界和风险面之间的关系。"
                ),
                (
                    "沙箱和 Prompt 约束有什么区别？",
                    "Prompt 约束只是告诉模型‘不要这样做’，沙箱约束是真正让它‘做不到这样做’。"
                    "前者是行为提醒，后者是系统护栏。"
                    "生产系统不能只依赖 Prompt，因为 Prompt 可能被用户输入、工具返回或模型漂移绕开。"
                    "所以我的态度是：Prompt 可以帮助提高行为概率，但高风险边界必须靠权限、校验和沙箱落地。",
                    "是否理解软约束和硬约束的区别。"
                ),
                (
                    "如果 Agent 能执行 shell，你会如何设计权限边界？",
                    "我不会直接开放完整 shell。"
                    "更安全的做法是预定义一组允许的命令模板、参数范围和工作目录，必要时再通过中间层包装成结构化工具。"
                    "例如可以允许只读搜索、运行测试、读取日志，但禁止任意删除、系统级安装、跨目录移动和未授权网络访问。"
                    "如果必须执行有副作用命令，还要增加人工确认和完整审计。",
                    "是否具备把命令执行能力收束成安全工具的意识。"
                ),
                (
                    "为什么‘能执行命令’不等于‘应该让模型直接写命令’？",
                    "因为自然语言生成的命令包含太多不确定性，既可能写错路径，也可能越过边界。"
                    "让模型直接输出 shell 文本，本质上是在把开放式生成变成高风险执行。"
                    "更好的设计是把常见能力做成参数化工具，例如 search_logs(service, keyword) 或 run_test(target)。"
                    "这样系统能对参数做校验，而不是事后猜测命令有没有问题。",
                    "是否理解开放式生成和结构化执行之间的安全差异。"
                ),
                (
                    "哪些操作必须经过人工确认，而不能只靠沙箱？",
                    "沙箱负责限制范围，但高风险动作即使在受限范围内也可能有严重后果，所以还需要人工确认。"
                    "例如删除用户数据、推送代码、发外部消息、修改生产配置、执行付费动作、批量导出敏感信息。"
                    "换句话说，沙箱解决的是‘最多能伤到哪里’，人工确认解决的是‘这次到底该不该做’。",
                    "是否能区分边界控制和业务授权。"
                ),
                (
                    "如果用户在提示词里要求 Agent 越权访问文件或网络，你会怎么防？",
                    "第一，不让用户输入直接决定权限。"
                    "第二，把权限校验放在工具执行层，而不是只放在模型侧。"
                    "第三，工具收到超范围路径、域名或动作时直接拒绝，并返回可审计错误。"
                    "第四，把这类越权尝试记录进安全日志和评测集。"
                    "本质上要做到：用户可以提出请求，但能不能执行由系统权限模型决定。",
                    "是否理解权限判断必须在执行层落地。"
                ),
                (
                    "沙箱会不会影响 Agent 能力？你怎么回答这个取舍？",
                    "会，沙箱一定会减少自由度。"
                    "但在生产里，这不是缺点，而是必要代价。"
                    "真正的问题不是要不要限制，而是限制得是否贴合业务。"
                    "如果边界过窄，Agent 完不成任务；如果边界过宽，系统风险失控。"
                    "所以应该根据任务类型设计分级权限，而不是给所有 Agent 同一把万能钥匙。",
                    "是否具备能力与安全之间的权衡意识。"
                ),
                (
                    "如果你来设计一个面向企业的 Agent 权限体系，你会怎么分层？",
                    "我会至少分成四层。第一层是角色层，比如客服 Agent、检索 Agent、运维 Agent、审批 Agent。"
                    "第二层是资源层，控制它能访问哪些文件、表、服务和 API。"
                    "第三层是动作层，区分只读、可写、可删除、可对外发送、可触发付费等操作。"
                    "第四层是审批层，高风险动作必须走人工确认或双重校验。"
                    "这样权限既可解释，也方便审计和扩展。",
                    "是否具备系统化权限模型思维。"
                ),
            ],
        ),
    ]

    question_number = 1
    for title, questions in sections:
        doc.add_heading(title, level=2)
        for question, answer, focus in questions:
            add_interview_question(
                doc,
                question_number,
                question,
                answer,
                focus,
            )
            question_number += 1

    doc.add_heading("面试官可能继续追问的证明材料", level=2)
    add_body(
        doc,
        "只会回答概念还不够。高级面试官可能要求候选人拿出能够验证设计的材料。即使面试时不展示，也应提前准备以下内容："
    )
    for item in [
        "一张当前真实运行链路图，并标明默认模式和可选模式。",
        "一组路由混淆矩阵以及典型误路由案例。",
        "RAG 的检索与拒答评测结果，包括至少一组改造前后对照。",
        "一次重复 request_id、任务版本冲突和 Worker 重试的演示。",
        "Outbox 积压、死信任务和僵死任务恢复的监控示例。",
        "项目 P50/P95 延迟拆分与每请求 token 成本。",
        "至少一个由自己亲手完成并能解释前后差异的优化。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("快速评分标准", level=2)
    add_comparison_table(
        doc,
        ["表现", "典型特征", "面试判断"],
        [
            ["较弱", "只复述简历名词，无法解释失败场景、指标和边界", "项目所有权不足"],
            ["合格", "能够讲清主流程和常见概念，但缺少量化验证与生产经验", "具备初中级实现能力"],
            ["良好", "能解释取舍、降级、幂等、一致性和评测，并主动指出问题", "具备独立负责模块能力"],
            ["优秀", "能纠正简历过度表述，给出容量、成本、安全和演进方案，并用数据验证", "接近高级 Agent 工程师思维"],
        ],
        [1300, 5200, 2860],
    )


def add_beginner_core_concepts(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("零基础重点专题：七个核心机制从入门到真正理解", level=1)
    add_body(
        doc,
        "这一章专门写给尚未接触过后端并发、Agent 编排和分布式任务的人。阅读时不要急着记英文名词，"
        "先理解每个机制解决的现实问题，再理解它在本项目中的位置。"
    )
    add_callout(
        doc,
        "先看全局",
        "用户发来一句话后，路由负责决定“这是什么事、交给谁”；LangGraph 负责安排“先做什么、后做什么”；"
        "多 Agent 负责在需要时让不同专家分工；Redis 负责快速读取常用信息；request_id 幂等和乐观锁负责避免重复与并发覆盖；"
        "Outbox 负责把不必立即完成的索引工作可靠地交给后台 Worker。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_comparison_table(
        doc,
        ["机制", "它回答的问题", "本项目中的作用"],
        [
            ["路由", "这次问题属于哪一类？正在继续哪个任务？", "选择任务、工具或专家"],
            ["LangGraph", "这些处理步骤以什么顺序执行？", "组织整条请求流程"],
            ["多 Agent", "是否需要多个不同角色共同处理？", "家具、设备、报表等专家分工"],
            ["Redis", "怎样快速拿到最近经常使用的数据？", "缓存近期对话和记忆"],
            ["request_id 幂等", "同一请求重复到达怎么办？", "避免重复消息和重复任务"],
            ["乐观锁", "两个请求同时改同一任务怎么办？", "发现并发覆盖并重试"],
            ["Outbox", "业务保存成功后，后台任务怎样保证不丢？", "可靠投递向量索引任务"],
        ],
        [1900, 3600, 3860],
    )

    doc.add_heading("专题一：LangGraph 到底是什么", level=2)
    doc.add_heading("1. 先纠正名称", level=3)
    add_body(
        doc,
        "正确名称是 LangGraph，不是 LangGroup。LangGraph 可以拆成两个词：Lang 表示语言模型应用，Graph 表示图。"
        "它不是一个新的大模型，也不负责训练模型；它是用来组织大模型应用处理流程的框架。"
    )

    doc.add_heading("2. 为什么普通函数逐渐不够用", level=3)
    add_body(
        doc,
        "最简单的聊天系统只有一步：把用户问题交给模型，然后返回答案。这时一个函数就足够。"
        "但本项目回答前后还需要识别任务、查询历史、拼接上下文、改写问题、调用工具、保存消息、更新任务和刷新记忆。"
        "如果这些逻辑全部塞在一个大函数中，代码会出现大量条件判断，某一步失败也很难判断应该终止还是继续。"
    )
    add_callout(
        doc,
        "生活类比",
        "LangGraph 像医院的就诊流程图。挂号、分诊、检查、医生诊断、缴费和取药都是不同节点；"
        "患者携带的病历信息就是状态；流程图决定每一步完成后去哪里。",
    )

    doc.add_heading("3. 图、节点、边和状态", level=3)
    add_comparison_table(
        doc,
        ["概念", "通俗解释", "项目例子"],
        [
            ["节点 Node", "一个相对独立的处理步骤", "任务路由、历史召回、回答生成"],
            ["边 Edge", "一个节点完成后去哪个节点", "召回完成后进入上下文组装"],
            ["状态 State", "整条流程共同携带的资料包", "问题、用户 ID、任务、历史、最终答案"],
            ["分支", "根据条件选择不同路径", "Agentic 模式或专家模式"],
            ["汇合", "等待多个前置步骤结束后继续", "任务路由和历史召回汇入上下文组装"],
        ],
        [1750, 3900, 3710],
    )
    add_body(
        doc,
        "状态非常关键。你可以把它想象成一只文件袋：开始时只有用户问题和几个 ID；任务节点往文件袋放入当前任务；"
        "历史节点放入召回结果；回答节点放入最终答案；持久化节点再从文件袋读取这些信息。"
    )

    doc.add_heading("4. 什么是菱形工作流", level=3)
    for step in [
        "请求开始后，任务路由和历史召回同时启动，因为它们互不依赖。",
        "上下文组装必须同时拿到任务结果和历史结果，所以在这里等待两条路径汇合。",
        "生成回答并保存消息后，任务状态更新和记忆刷新又可以分别处理。",
        "两个后处理都结束后，本轮工作流完成。",
    ]:
        add_number(doc, step)
    add_callout(
        doc,
        "为什么叫菱形",
        "一条路径先分成两条并行路径，再重新合成一条，画出来像菱形。它的意义不是图形好看，而是减少不必要的串行等待。",
        fill="ECFDF3",
        color=GREEN,
    )

    doc.add_heading("5. LangGraph 与 Agent 的区别", level=3)
    add_comparison_table(
        doc,
        ["对象", "负责什么", "不负责什么"],
        [
            ["LangGraph", "控制业务步骤、状态流转、分支与并行", "不负责生成语言和理解知识"],
            ["大模型 Agent", "理解问题、选择工具、结合结果回答", "不应直接控制事务和权限等关键操作"],
            ["工具", "执行确定的外部能力", "不决定整条业务流程"],
        ],
        [1800, 4000, 3560],
    )
    add_body(
        doc,
        "成熟的生产系统通常是“外层确定流程，内层有限自主”：LangGraph 控制大方向，Agent 只在回答节点内决定是否调用知识库、"
        "报表或天气工具。这样既保留模型灵活性，又避免让模型随意控制数据库提交等关键操作。"
    )

    doc.add_heading("6. 节点失败时怎么办", level=3)
    add_body(
        doc,
        "并不是所有失败都应同样处理。历史语义召回失败时，可以不带远期历史继续回答；重排序器失败时，可以使用较弱的原始排名；"
        "记忆刷新失败时，可以先保留本轮消息，下次再更新。但如果核心消息无法保存，系统不应该假装成功。"
    )
    add_callout(
        doc,
        "面试表达",
        "我把 LangGraph 用作确定性的业务骨架。辅助节点支持降级，核心持久化节点失败快返；"
        "并行节点通过状态合并后进入统一回答链路。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("专题二：路由是什么，为什么项目里有两种路由", level=2)
    doc.add_heading("1. 路由的本质", level=3)
    add_body(
        doc,
        "路由就是“根据输入选择下一条处理路径”。快递分拣中心会根据地址把包裹送往不同城市；医院分诊会根据症状选择科室；"
        "Agent 系统则根据用户问题选择任务、专家或工具。"
    )
    add_callout(
        doc,
        "不要混淆",
        "路由负责选择路径，RAG 负责搜索知识。路由可能决定去调用 RAG，但它本身不是检索。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("2. 本项目的任务路由", level=3)
    add_body(
        doc,
        "任务路由回答的是：用户现在是在开启一个新问题、继续当前问题、恢复以前的问题，还是只在闲聊？"
        "它管理的是业务连续性，而不是专家分工。"
    )
    add_comparison_table(
        doc,
        ["结果", "意义", "例子"],
        [
            ["new", "新建任务，并暂停当前不相关任务", "从沙发清洁切到扫地机器人故障"],
            ["continue", "继续当前任务", "“按你说的做了，还是没用”"],
            ["resume", "恢复以前相关的任务", "“继续上次的沙发问题”"],
            ["no_task", "不绑定业务任务", "你好、谢谢、再见"],
        ],
        [1500, 3000, 4860],
    )
    add_body(
        doc,
        "任务路由通常先过滤闲聊，再检测“继续、试过了、还是”等延续表达，然后比较家具对象和意图。"
        "如果关键词结果不够确定，还可以用向量相似度寻找以前最相关的任务。"
    )

    doc.add_heading("3. 本项目的专家路由", level=3)
    add_body(
        doc,
        "专家路由回答的是：这次问题应该交给家具专家、设备专家、报表专家还是通用专家？"
        "一个问题也可能同时属于两个领域，例如“结合本月使用数据，分析扫地机器人的清洁效果”，"
        "这时可以同时选择设备专家和报表专家。"
    )
    add_comparison_table(
        doc,
        ["任务路由", "专家路由"],
        [
            ["关注对话正在解决哪件事", "关注哪种专家能力适合回答"],
            ["输出 new、continue、resume、no_task", "输出 Furniture、Device、Report、General"],
            ["主要用于上下文和状态管理", "主要用于答案生成分工"],
            ["每轮通常只对应一个当前任务", "复合问题可以选择一至两个专家"],
        ],
        [4680, 4680],
    )

    doc.add_heading("4. 为什么采用规则加 LLM 两级路由", level=3)
    add_body(
        doc,
        "规则路由速度快、成本低、结果可解释。例如问题中明确出现“衣柜”，可以直接提高家具专家分数。"
        "但用户可能说“它这个月效率怎么样”，只看关键词很难判断，需要结合上下文让 LLM 兜底。"
    )
    for step in [
        "先给家具、设备、报表和通用类别分别计算关键词分数。",
        "如果某一类别命中明确、置信度高，直接选择，省去一次模型调用。",
        "如果没有明显结果或多个类别冲突，把问题、任务和历史交给 LLM 判断。",
        "对结果去重并限制最多两个专家，控制费用、延迟和答案冲突。",
    ]:
        add_number(doc, step)
    add_body(
        doc,
        "LLM 给出的 confidence 不能直接当成真实概率。它更像一个启发式分数，阈值应通过测试集调整。"
        "路由最终应使用 precision、recall、F1 和混淆矩阵评估，而不是凭感觉说“挺准”。"
    )

    doc.add_heading("5. 路由错了会发生什么", level=3)
    add_bullet(doc, "家具问题被送到设备专家，可能检索到错误领域的资料。")
    add_bullet(doc, "闲聊被绑定到旧任务，可能污染任务状态。")
    add_bullet(doc, "一个简单问题路由过多专家，会增加费用和等待时间。")
    add_bullet(doc, "遗漏必要专家，答案可能只覆盖用户问题的一部分。")
    add_callout(
        doc,
        "面试表达",
        "项目中有两类路由：任务路由负责对话连续性，专家路由负责能力分工。"
        "专家路由采用规则优先、LLM 兜底，最多选择两个专家，避免无意义的并行扩散。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("专题三：多 Agent 到底是什么", level=2)
    doc.add_heading("1. 一个 Agent 已经能调用工具，为什么还需要多个 Agent", level=3)
    add_body(
        doc,
        "单 Agent 像一个全科客服：它可以查家具资料、设备资料和报表。多 Agent 则像一个客服团队，"
        "不同成员拥有不同提示词、工具和专业边界。多 Agent 的价值不是数量多，而是把相互干扰的职责拆开。"
    )
    add_comparison_table(
        doc,
        ["单 Agent + 多工具", "多 Agent"],
        [
            ["一个模型统一决定调用什么工具", "路由器先选择一个或多个专家"],
            ["结构简单、成本低、延迟较小", "职责清晰，可为不同领域设计独立策略"],
            ["复杂提示词可能互相干扰", "需要处理多次调用、冲突和答案合成"],
            ["适合大多数边界清晰的问题", "适合真正需要领域分工的复合问题"],
        ],
        [4680, 4680],
    )

    doc.add_heading("2. 多 Agent 系统中的角色", level=3)
    add_comparison_table(
        doc,
        ["角色", "职责", "类比"],
        [
            ["Router", "决定选择哪些专家", "客服主管分派工单"],
            ["Furniture Agent", "处理家具选购、清洁和维护", "家具售后专员"],
            ["Device Agent", "处理扫地机器人故障和维护", "设备技术支持"],
            ["Report Agent", "读取月度记录并进行说明", "数据分析专员"],
            ["General Agent", "处理问候、澄清和一般问题", "前台客服"],
            ["Composer", "整合多个专家结果并处理冲突", "负责最终回复的主管"],
        ],
        [1800, 4300, 3260],
    )

    doc.add_heading("3. 多 Agent 的一次执行", level=3)
    for step in [
        "路由器识别问题涉及哪些领域。",
        "只选择必要的专家，最多两个。",
        "两个专家可以并行处理，各自返回摘要、证据、置信度和未解决问题。",
        "答案合成器检查结果是否互补或冲突。",
        "合成器输出一个统一回答，而不是把两个答案简单拼接。",
    ]:
        add_number(doc, step)

    doc.add_heading("4. 为什么不能专家越多越好", level=3)
    add_body(
        doc,
        "每增加一个专家，就可能增加一次模型或检索调用，带来更高费用、更长尾延迟和更多冲突。"
        "如果家具专家、设备专家、报表专家和通用专家每次都执行，系统看起来复杂，却不一定更准确。"
        "是否启用多 Agent 应由评测结果决定。"
    )

    doc.add_heading("5. 专家结果冲突怎么办", level=3)
    add_body(
        doc,
        "答案合成器不能用“多数表决”或让另一个模型随便选。它应比较每个结果的领域归属、证据质量和风险等级。"
        "能够并存的信息可以分点合并；如果维修建议发生冲突，应优先可验证证据和安全方案，并明确说明不确定性。"
    )
    add_callout(
        doc,
        "当前项目的真实状态",
        "项目已经实现多 Agent 路由、并行执行和答案合成能力，也有相应测试；但默认主链路实际使用的是单 Agent 工具调用模式。"
        "面试时必须把“具备能力”和“默认正在使用”区分开。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("专题四：Redis 缓存到底是什么", level=2)
    doc.add_heading("1. 先理解缓存", level=3)
    add_body(
        doc,
        "缓存是为了加快读取而保存的数据副本。你经常查一本厚词典，如果把最近查过的几页放在桌边，下次就不必重新翻整本词典。"
        "桌边的几页是缓存，完整词典是事实源。"
    )
    add_callout(
        doc,
        "在本项目中",
        "MySQL 是完整词典，Redis 是桌边便签。Redis 丢失后可以从 MySQL 重新加载，因此 Redis 不应成为唯一数据来源。",
    )

    doc.add_heading("2. 为什么不用 MySQL 每次都查", level=3)
    add_bullet(doc, "Redis 在内存中读取，通常比磁盘数据库更快。")
    add_bullet(doc, "最近对话每一轮都要使用，属于高频数据。")
    add_bullet(doc, "使用 TTL 可以让长期不用的会话自动释放缓存。")
    add_bullet(doc, "降低 MySQL 在高并发时的重复读取压力。")

    doc.add_heading("3. Cache-Aside 读取流程", level=3)
    for step in [
        "先根据 session_id 去 Redis 查询最近对话。",
        "如果 Redis 有数据，直接返回，这叫缓存命中。",
        "如果 Redis 没有数据，去 MySQL 查询，这叫缓存未命中。",
        "将 MySQL 查询结果写回 Redis，并设置过期时间。",
        "下一次请求可以直接从 Redis 读取。",
    ]:
        add_number(doc, step)

    doc.add_heading("4. 写入时怎样保持一致", level=3)
    add_body(
        doc,
        "最重要的原则是先确保 MySQL 事实保存成功，再刷新或删除 Redis 缓存。缓存刷新失败时，"
        "最多导致下一次读取稍慢或短时间读到旧数据；如果反过来只写 Redis 不写 MySQL，Redis 过期后数据会永久消失。"
    )
    add_comparison_table(
        doc,
        ["问题", "含义", "常见处理"],
        [
            ["缓存穿透", "不断查询数据库中根本不存在的数据", "缓存空结果、参数校验"],
            ["缓存击穿", "某个热门键过期，大量请求同时查数据库", "互斥加载、提前刷新"],
            ["缓存雪崩", "大量键同时过期或 Redis 故障", "过期时间加随机值、限流和降级"],
            ["缓存脏数据", "缓存内容落后于数据库", "写库后删除或刷新缓存"],
        ],
        [1800, 3900, 3660],
    )

    doc.add_heading("5. Redis 挂了怎么办", level=3)
    add_body(
        doc,
        "Redis 挂掉时，系统应该设置很短的连接超时，快速回退到 MySQL，而不是一直等待。"
        "但大量请求同时回退可能压垮 MySQL，所以还需要熔断、限流或本地短缓存。Redis 恢复后，数据可以按访问需要重新回填。"
    )
    add_callout(
        doc,
        "面试表达",
        "Redis 在项目中是可丢失、可重建的加速层。读取采用先缓存后数据库，核心写入以 MySQL 成功为准；"
        "缓存故障时降级到 MySQL，并防止击穿。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("专题五：request_id 幂等到底是什么", level=2)
    doc.add_heading("1. 为什么同一个请求会重复出现", level=3)
    add_body(
        doc,
        "用户只点击一次发送，不代表后端一定只收到一次。网络可能超时，前端不知道请求是否成功，于是自动重试；"
        "用户也可能连续点击；网关或消息系统同样可能重放请求。"
    )
    add_callout(
        doc,
        "没有幂等时的后果",
        "同一句话可能被保存两次、同一个任务可能创建两份、同一条 Outbox 任务可能重复产生。"
        "用户看到一条回答，数据库里却有重复副作用。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("2. request_id 是什么", level=3)
    add_body(
        doc,
        "request_id 是一次业务请求的唯一编号，例如一串 UUID。它不是 user_id，也不是 session_id。"
        "同一个用户可以有多个会话，同一个会话可以有很多次请求，而每次请求都有自己的 request_id。"
    )
    add_comparison_table(
        doc,
        ["ID", "代表什么", "一个用户通常有多少个"],
        [
            ["user_uuid", "用户身份", "通常一个稳定 ID"],
            ["session_id", "一次对话会话", "可以有多个"],
            ["request_id", "会话中的某一次请求", "每发一条消息产生一个"],
            ["task_id", "正在解决的业务问题", "可能跨多轮甚至跨会话"],
        ],
        [1800, 4200, 3360],
    )

    doc.add_heading("3. 幂等是什么意思", level=3)
    add_body(
        doc,
        "幂等不是保证请求物理上只执行一次，而是保证同一请求执行一次或重复执行多次，最终业务结果相同。"
        "例如按一次电梯楼层按钮和按十次，电梯最终只需要去一次那个楼层。"
    )

    doc.add_heading("4. 项目怎样使用 request_id", level=3)
    for step in [
        "前端发送消息时生成一个 request_id。",
        "用户消息和助手消息保存时都带上这个 request_id，并使用不同 sequence_no 区分顺序。",
        "数据库设置 session_id、request_id、sequence_no 的唯一约束。",
        "相同请求再次到达时，数据库发现记录已存在，不再重复插入。",
        "任务事件也使用 request_id 去重，避免同一轮重复更新任务状态。",
    ]:
        add_number(doc, step)
    add_body(
        doc,
        "关键点是必须有数据库唯一约束。应用代码先查询“有没有”再插入并不安全：两个并发请求可能同时查询到不存在，"
        "然后都插入。数据库唯一键才是最后的并发防线。"
    )

    doc.add_heading("5. request_id 能解决什么，不能解决什么", level=3)
    add_comparison_table(
        doc,
        ["可以解决", "不能直接解决"],
        [
            ["同一个 request_id 因网络重试重复提交", "用户用不同 request_id 连续发送相同文字"],
            ["同一轮消息和任务事件重复写入", "两个不同请求同时修改同一个任务"],
            ["Worker 重试时识别相同业务事件", "外部副作用天然只执行一次"],
        ],
        [4680, 4680],
    )
    add_callout(
        doc,
        "面试表达",
        "request_id 是业务幂等键，不是用户或会话身份。项目通过 request_id 加数据库唯一约束实现重复请求去重，"
        "保证的是最终业务效果一致，而不是请求绝对只执行一次。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("专题六：乐观锁为什么叫锁，却不提前锁住数据", level=2)
    doc.add_heading("1. 并发覆盖问题", level=3)
    add_body(
        doc,
        "假设任务当前版本是 5。请求 A 和请求 B 几乎同时读取到版本 5。A 记录“用户已经清洗滚刷”，"
        "B 记录“用户不接受拆机”。如果两者都直接保存，后保存的 B 可能覆盖 A 的修改。"
    )

    doc.add_heading("2. 悲观锁与乐观锁", level=3)
    add_comparison_table(
        doc,
        ["方式", "想法", "适用情况"],
        [
            ["悲观锁", "认为冲突很可能发生，读取时就锁住数据，别人等待", "冲突频繁、操作很短"],
            ["乐观锁", "认为冲突较少，先正常读取，更新时检查版本是否变化", "读多写少、冲突不频繁"],
        ],
        [1700, 4400, 3260],
    )
    add_body(
        doc,
        "乐观锁并不是完全没有数据库锁，而是业务层不在读取后长时间占用锁。它通过版本号检测“我读取之后有没有别人修改过”。"
    )

    doc.add_heading("3. state_version 如何工作", level=3)
    for step in [
        "请求 A 和 B 都读取任务，得到 state_version = 5。",
        "A 更新任务时要求数据库中的版本仍然等于 5。",
        "A 更新成功，并把版本加到 6。",
        "B 也尝试以版本 5 更新，但数据库当前已经是 6，所以受影响行数为 0。",
        "系统判断发生冲突，B 重新读取版本 6，再决定是否合并并重试。",
    ]:
        add_number(doc, step)
    add_callout(
        doc,
        "核心条件",
        "更新动作必须同时满足“任务 ID 正确、用户正确、数据库版本等于我读取时的版本”。"
        "只有成功更新的请求才能把版本号加一。",
    )

    doc.add_heading("4. 冲突后不能简单覆盖", level=3)
    add_body(
        doc,
        "发生冲突后，不能把 B 的旧内容强制写入，因为这会重新丢掉 A 的修改。正确做法是重新读取最新任务，"
        "判断 B 的修改与 A 是否可以合并，然后有限次数重试。如果业务冲突无法自动解决，应返回冲突或交给人工处理。"
    )

    doc.add_heading("5. 幂等和乐观锁的区别", level=3)
    add_comparison_table(
        doc,
        ["request_id 幂等", "乐观锁"],
        [
            ["处理同一请求重复到达", "处理不同请求同时修改同一数据"],
            ["问题是重复执行", "问题是并发覆盖"],
            ["常用唯一键去重", "常用版本号检测冲突"],
            ["相同 request_id 应复用结果", "版本冲突后重新读取、合并或重试"],
        ],
        [4680, 4680],
    )
    add_callout(
        doc,
        "面试表达",
        "幂等解决同一事件的重复执行，乐观锁解决不同事件的并发覆盖。项目用 request_id 去重业务事件，"
        "用 state_version 检测任务状态是否已被其他请求修改。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("专题七：Outbox 异步任务为什么存在", level=2)
    doc.add_heading("1. 同步和异步是什么意思", level=3)
    add_body(
        doc,
        "同步表示用户请求必须等待工作完成后才能结束；异步表示先记录“这件事需要做”，让后台 Worker 稍后处理。"
        "保存聊天消息是核心操作，通常应同步完成；把历史对话转换成向量比较耗时，而且短暂延迟通常可以接受，适合异步。"
    )

    doc.add_heading("2. 为什么不能保存 MySQL 后直接写 Chroma", level=3)
    add_body(
        doc,
        "MySQL 和 Chroma 是两个独立系统。假设先把消息写入 MySQL，然后进程在写 Chroma 前崩溃，"
        "用户对话已经存在，但历史向量永远缺失。如果先写 Chroma 再写 MySQL，MySQL 失败时又会留下没有原始事实的孤立向量。"
        "这叫双写一致性问题。"
    )

    doc.add_heading("3. Outbox 的核心做法", level=3)
    for step in [
        "开启一个 MySQL 事务。",
        "在事务中保存用户消息和助手消息。",
        "在同一个事务中插入一条 Outbox 任务，内容是“请为这轮对话建立向量索引”。",
        "提交事务：消息和任务要么一起成功，要么一起失败。",
        "后台 Worker 不断领取 pending 任务。",
        "Worker 写入 Chroma 成功后，把任务标记为 completed。",
        "如果失败，记录错误并安排稍后重试。",
    ]:
        add_number(doc, step)
    add_callout(
        doc,
        "生活类比",
        "餐厅服务员先把顾客订单和“待出餐”小票一起写入可靠的订单系统，再由厨房按小票做菜。"
        "服务员不需要站在厨房等待，但只要订单系统里有小票，厨房恢复后就能继续处理。",
    )

    doc.add_heading("4. Worker 怎样并发领取任务", level=3)
    add_body(
        doc,
        "多个 Worker 可以同时处理任务。领取时使用 FOR UPDATE SKIP LOCKED：某个 Worker 用行锁占用自己选择的任务；"
        "其他 Worker 看到这些记录已锁定，就跳过并领取后面的任务，而不是排队等待。"
    )
    add_callout(
        doc,
        "技术用词要准确",
        "这不是“无锁消费”。它仍然使用数据库行锁，只是 SKIP LOCKED 让其他消费者跳过已锁记录，因此更准确的说法是"
        "“基于行锁的非阻塞并发领取”。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("5. Worker 写成功后崩溃会怎样", level=3)
    add_body(
        doc,
        "一种典型情况是：Worker 已经成功把向量写进 Chroma，但还没来得及把 Outbox 标记为 completed 就崩溃。"
        "任务恢复后会再次执行。因此 Outbox 通常提供至少一次投递，而不是保证物理上只执行一次。"
    )
    add_body(
        doc,
        "为了让重复执行没有副作用，向量写入使用稳定 ID 和 upsert。相同 ID 再写一次会覆盖同一条记录，而不是生成两条。"
        "这种设计叫消费者幂等。"
    )

    doc.add_heading("6. 重试、指数退避和僵死恢复", level=3)
    add_comparison_table(
        doc,
        ["机制", "解决的问题", "具体含义"],
        [
            ["指数退避", "下游故障时避免疯狂重试", "等待时间逐步增加，如 10、20、40 秒"],
            ["随机抖动", "避免大量任务同时再次重试", "在等待时间上增加随机变化"],
            ["最大重试", "避免永久错误无限循环", "超过次数进入 dead 状态"],
            ["僵死恢复", "Worker 崩溃后任务一直 processing", "超过锁定时间后恢复为 pending"],
        ],
        [1800, 3900, 3660],
    )
    add_body(
        doc,
        "僵死恢复的超时时间必须大于正常任务可能需要的最长时间，否则仍在执行的慢任务会被误判，造成多个 Worker 同时处理。"
    )

    doc.add_heading("7. 最终一致性是什么", level=3)
    add_body(
        doc,
        "消息提交后，MySQL 和 Chroma 可能在几秒或几分钟内暂时不一致；只要 Worker 最终处理成功，两边会重新一致。"
        "这叫最终一致性。Outbox 不保证立即一致，也不天然保证 exactly-once。"
    )
    add_callout(
        doc,
        "面试表达",
        "Outbox 解决 MySQL 与 Chroma 的可靠双写问题。消息和任务在同一数据库事务提交，Worker 至少一次消费，"
        "通过稳定 ID、upsert、退避重试和僵死恢复保证最终一致。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("把七个机制串成一次真实请求", level=2)
    add_callout(
        doc,
        "用户问题",
        "“我按你上次说的清洗了滚刷，还是有异响，顺便看看这个月的使用效率。”",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    for step in [
        "前端生成 request_id，表示这是一次新的业务请求；如果网络重试，继续使用同一个 request_id。",
        "LangGraph 同时启动任务路由和历史召回。",
        "任务路由判断用户在继续扫地机器人异响任务，而不是创建新任务。",
        "历史召回找到上次建议，并回 MySQL 验证用户确实被建议清洗滚刷。",
        "上下文组装记录“滚刷已经清洗但异响仍存在”，防止回答再次只建议清洗滚刷。",
        "如果使用多 Agent 专家模式，专家路由选择 Device Agent 和 Report Agent。",
        "两个专家分别处理异响排查和月度效率数据，答案合成器统一回复。",
        "核心消息按 request_id 幂等保存；重复请求不会再次插入同一轮消息。",
        "任务状态更新时使用 state_version 乐观锁，避免另一请求同时覆盖本轮结果。",
        "Redis 缓存最新对话，后续追问“那下一步呢”可以快速读取。",
        "MySQL 事务同时写入一条 Outbox 任务。",
        "后台 Worker 将本轮问答写入 Chroma；失败则退避重试，成功后未来可以通过语义搜索召回。",
    ]:
        add_number(doc, step)

    doc.add_heading("七个机制的因果关系", level=2)
    add_comparison_table(
        doc,
        ["如果缺少", "最可能出现的问题"],
        [
            ["LangGraph", "流程散落在大量条件判断中，难以测试和降级"],
            ["任务路由", "多个问题互相污染，旧任务被错误带入当前回答"],
            ["专家路由 / 多 Agent", "复合领域问题只能由一个通用角色勉强处理"],
            ["Redis", "每轮都查询 MySQL，延迟和数据库压力增加"],
            ["request_id 幂等", "网络重试产生重复消息和重复任务"],
            ["乐观锁", "不同请求同时更新时互相覆盖"],
            ["Outbox", "消息保存成功后索引任务可能永久丢失"],
        ],
        [2700, 6660],
    )

    doc.add_heading("零基础自测：能回答这些才算真正理解", level=2)
    checks = [
        "LangGraph 和大模型分别负责什么？",
        "为什么任务路由和专家路由不是一回事？",
        "多 Agent 为什么不能每次把所有专家都调用一遍？",
        "为什么 Redis 丢失后系统仍然应该能够恢复？",
        "request_id、session_id、user_id、task_id 有什么区别？",
        "幂等为什么不代表请求绝对只执行一次？",
        "两个不同 request_id 同时修改任务时，为什么还需要乐观锁？",
        "state_version 从 5 变成 6 后，拿着版本 5 的请求为什么不能继续覆盖？",
        "为什么 MySQL 消息和 Outbox 任务必须在同一个事务？",
        "FOR UPDATE SKIP LOCKED 为什么不是无锁？",
        "Worker 已经写入 Chroma 但没有标记完成就崩溃，如何避免重复向量？",
        "Outbox 为什么是最终一致，而不是立即一致？",
    ]
    for check in checks:
        add_bullet(doc, check)


def add_contents(doc: Document) -> None:
    doc.add_heading("阅读导航", level=1)
    add_body(doc, "这份手册不要求你阅读任何代码。建议先通读前三章建立全局认识，再根据面试准备需要选择后面的专题。")
    sections = [
        "零基础重点专题：LangGraph、路由、多 Agent、Redis、幂等、乐观锁与 Outbox",
        "第一部分：项目定位与整体架构",
        "第二部分：一次用户请求的完整旅程",
        "第三部分：Agent、工具与工作流",
        "第四部分：RAG 知识库系统",
        "第五部分：记忆、历史召回与任务状态",
        "第六部分：MySQL、Redis、Chroma 与 Outbox",
        "第七部分：项目从 Demo 到工程化的演进",
        "第八部分：准确性、可靠性与降级设计",
        "第九部分：项目当前的真实边界与优化路线",
        "第十部分：面试表达模板",
        "第十一部分：场景演练与自测",
        "附录：核心术语速查",
        "专项附录：按知识点整理的面试题总汇",
    ]
    for section in sections:
        add_bullet(doc, section)
    add_callout(
        doc,
        "最重要的阅读原则",
        "不要背模块名称。每学一个概念，都问自己三个问题：它解决什么问题？没有它会怎样？它与其他模块如何配合？",
        fill="FFF8E8",
        color=GOLD,
    )


def add_zero_foundation_expansion(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("零基础补课版：把核心机制彻底讲白", level=1)
    add_body(
        doc,
        "这一部分专门写给零基础读者。你不需要先会看代码，也不需要先理解框架文档，"
        "只要先把每个机制想成现实世界里的一个角色，再去理解它在项目里的作用。"
    )
    add_callout(
        doc,
        "阅读方法",
        "每个知识点都按四个问题理解：它是什么、为什么会出现、如果没有它会怎样、面试时应该怎么用一句人话说出来。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )

    doc.add_heading("1. 先把整个系统想成一家医院", level=2)
    add_body(
        doc,
        "用户发来一句话，就像病人走进医院说“我不舒服”。医院不会让第一个接待的人直接拍脑袋下结论，"
        "而是先分诊、再检查、必要时找不同科室会诊、最后再把病历归档。你的项目也是同样的思路。"
    )
    add_comparison_table(
        doc,
        ["现实里的角色", "项目里的对应物", "它解决的问题"],
        [
            ["分诊台", "路由", "先判断是什么问题，后面才知道怎么处理"],
            ["流程单", "LangGraph", "规定步骤顺序、分支和并行"],
            ["不同科室医生", "多 Agent", "复杂问题按领域分工"],
            ["病历室", "MySQL", "保存核心事实，不能丢"],
            ["前台快捷屏", "Redis", "高频上下文快速读取"],
            ["病例资料库", "RAG / 向量库 / BM25", "给答案提供证据"],
            ["后台检验中心", "Outbox + Worker", "把不必立刻完成的事可靠地放后台"],
        ],
        [1800, 3400, 4160],
    )

    doc.add_heading("2. LangGraph：它不是模型，而是流程导演", level=2)
    add_body(
        doc,
        "LangGraph 最容易被误解成“一个更高级的大模型”，但其实它不是。"
        "它更像导演或调度台，负责安排这条请求先经过哪些节点、哪些步骤能并行、哪一步失败时可以降级。"
    )
    add_callout(
        doc,
        "一句人话",
        "LangGraph 不负责‘答得聪不聪明’，它负责‘系统做事乱不乱’。",
        fill="ECFDF3",
        color=GREEN,
    )
    add_body(
        doc,
        "当流程只有两三步时，用普通函数也能写；但一旦加入任务路由、历史召回、RAG、专家并行、答案合成、持久化和异步任务，"
        "流程就不再是一条直线，而是一张图。LangGraph 的价值是把这张图显式表达出来。"
    )
    add_comparison_table(
        doc,
        ["概念", "白话解释", "在项目里的例子"],
        [
            ["节点", "一步明确动作", "任务路由、历史召回、答案生成"],
            ["边", "上一步做完后往哪走", "从路由走向检索，或走向闲聊分支"],
            ["状态", "一路传递的资料包", "用户问题、历史上下文、检索证据、最终答案"],
        ],
        [1500, 3200, 4660],
    )

    doc.add_heading("3. 多 Agent：不是越多越高级，而是分科会诊", level=2)
    add_body(
        doc,
        "多 Agent 的核心不是“多模型一起上就更强”，而是把复杂任务拆给不同职责的专家。"
        "例如家具知识、扫地机器人故障、报告解读、通用客服话术，其实是不同类型的能力。"
    )
    add_body(
        doc,
        "如果全交给一个 Agent，它会越来越像一个什么都懂一点、但越来越难控的全科医生；"
        "多 Agent 更像必要时让不同科室会诊，然后再由一个总医生统一对用户说人话。"
    )
    add_callout(
        doc,
        "为什么项目里最多只路由 2 个 Agent",
        "因为专家越多，成本、延迟和冲突概率都会上升。限制最多 2 个，本质上是在质量、速度和复杂度之间做平衡。",
        fill="FFF8E8",
        color=GOLD,
    )

    doc.add_heading("4. 路由：它本质上是在做分流", level=2)
    add_body(
        doc,
        "路由最容易被说成一个玄乎的 AI 词，但本质很朴素：就是先判断“这是什么事”，再决定“交给谁处理”。"
    )
    add_comparison_table(
        doc,
        ["层级", "它回答的问题", "典型输出"],
        [
            ["任务路由", "这次请求属于哪类处理路径", "new / continue / resume / no_task"],
            ["专家路由", "如果需要专家，该找谁", "Furniture / Device / Report / General"],
        ],
        [1700, 4200, 3460],
    )
    add_body(
        doc,
        "项目里使用‘规则关键词打分 + LLM 置信度兜底’，原因也很工程化：规则快、便宜、可解释；LLM 擅长处理模糊表达。"
        "两者结合比只押宝其中一个更稳。"
    )

    doc.add_heading("5. Redis：不是主存储，而是加速层", level=2)
    add_body(
        doc,
        "很多初学者会把 Redis 理解成“保存聊天记录的地方”。更准确的说法是：Redis 在这里主要负责短期上下文和高频读取加速。"
        "核心事实还是以 MySQL 为准。"
    )
    add_callout(
        doc,
        "一句人话",
        "MySQL 像总账本，Redis 像桌面便签。桌面便签查得快，但总账本才是最后依据。",
        fill="ECFDF3",
        color=GREEN,
    )

    doc.add_heading("6. request_id 幂等：防的是重复，不是并发覆盖", level=2)
    add_body(
        doc,
        "同一个请求为什么会重复出现？因为真实系统里会有用户连点、前端重试、网关重试、异步消费重试。"
        "所以幂等不是锦上添花，而是你默认必须要有的保护。"
    )
    add_body(
        doc,
        "但幂等最重要的定义不是‘物理上绝对只执行一次’，而是‘重复执行多次后，业务结果不应重复或错乱’。"
    )
    add_callout(
        doc,
        "面试最该背下来的一句",
        "幂等不是保证只执行一次，而是保证重复执行后业务结果不变。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("7. 乐观锁：防的是不同请求互相覆盖", level=2)
    add_body(
        doc,
        "幂等只管同一请求重复执行，但它不管两个不同请求同时修改同一条状态。"
        "这时就需要乐观锁。乐观锁通常通过版本号工作：你读取时看到版本是 5，提交时仍要求版本是 5，"
        "如果别人已经改成 6，你这次更新就应该失败或重试，而不是直接覆盖。"
    )
    add_comparison_table(
        doc,
        ["机制", "它防什么", "典型例子"],
        [
            ["request_id 幂等", "防同一事件重复处理", "同一句话因为重试被发了两次"],
            ["乐观锁", "防不同事件并发覆盖", "两个流程同时更新同一条任务状态"],
        ],
        [2200, 3100, 4060],
    )

    doc.add_heading("8. Outbox：为什么保存成功不等于全链路成功", level=2)
    add_body(
        doc,
        "假设消息已经写进 MySQL，但还没来得及写向量库，进程就崩了。"
        "这时 MySQL 里有数据，向量库里没有，对话事实和检索索引就不一致了，这就是双写问题。"
    )
    add_body(
        doc,
        "Outbox 的核心做法，是把‘后面要异步做的事’先作为任务记录，和核心业务数据一起写进同一个事务里。"
        "这样至少能保证：主事实先被可靠记住，后续再由 Worker 补做。"
    )
    add_callout(
        doc,
        "一句人话",
        "Outbox 不是让异步任务永远不失败，而是让它失败了也不会悄悄丢掉。",
        fill="ECFDF3",
        color=GREEN,
    )

    doc.add_heading("9. 最后把六个机制串成一条最容易讲的故事线", level=2)
    for item in [
        "用户提问后，路由先判断这是什么问题。",
        "LangGraph 负责安排接下来先查历史、还是先做检索、还是先走专家。",
        "如果问题复杂，多 Agent 按需要分工处理，而不是所有问题都硬塞给一个大脑。",
        "如果需要业务知识，RAG 负责找证据，而不是让模型裸答。",
        "短期高频上下文由 Redis 加速读取，核心事实仍然落到 MySQL。",
        "同一请求重复到达时，用 request_id 做幂等去重。",
        "不同请求并发更新状态时，用乐观锁防止后写覆盖前写。",
        "消息保存后，同事务写入 Outbox，由 Worker 后台补索引和其他异步动作。",
    ]:
        add_number(doc, item)


def add_follow_up_drill_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("专项追问题库：面试官继续追问时，你怎么接得住", level=1)
    add_body(
        doc,
        "这一部分不是普通题库，而是专门补‘第一问答完后，面试官最可能继续追什么’。"
        "很多人第一问答得还行，但第二问第三问就开始虚，所以这里把追问和回答一起写出来。"
    )

    groups = [
        (
            "LangGraph 追问",
            [
                (
                    "为什么不能自己写 if-else，非要用 LangGraph？",
                    "如果流程只有两三步，自己写当然可以。真正需要 LangGraph 的时候，是流程已经出现状态传递、分支、并行、降级和恢复需求。"
                    "它的收益不是‘能不能做’，而是‘复杂以后还能不能维护、测试和定位问题’。",
                    "考点：是否理解框架带来的核心价值是可维护性和可观测性。",
                ),
                (
                    "LangGraph 会自动帮你保证业务正确吗？",
                    "不会。它只是让你更容易把成功路径、失败路径和降级路径表达清楚。"
                    "真正的业务正确性，还是要靠你自己定义哪些节点必须失败、哪些节点允许降级、状态如何回滚或继续。",
                    "考点：不要把框架能力夸大成业务担保。",
                ),
            ],
        ),
        (
            "多 Agent 追问",
            [
                (
                    "什么时候你会坚持只用单 Agent？",
                    "当问题边界清楚、上下文不复杂、一个 Agent 已经能稳定处理时，我宁愿保留单 Agent。"
                    "因为多 Agent 的代价是真实存在的：延迟更高、成本更高、合成更复杂。",
                    "考点：是否理解复杂度必须换来明确收益。",
                ),
                (
                    "多个专家意见冲突怎么办？",
                    "不能直接拼接，而要交给答案合成器做去重、术语统一和冲突判断。"
                    "如果无法安全消解，我会优先保守回答，必要时重新回到证据层核查，而不是硬凑一个答案。",
                    "考点：是否知道多 Agent 难点在合成，不只在路由。",
                ),
            ],
        ),
        (
            "路由追问",
            [
                (
                    "为什么要先任务路由，再专家路由？",
                    "因为系统要先判断这是不是业务任务、是不是上一轮的延续，再决定是否值得调用专家。"
                    "如果一上来就做专家路由，很多本该在入口就过滤掉的请求也会白白消耗专家资源。",
                    "考点：是否理解入口分流的优先级。",
                ),
                (
                    "规则加权和 LLM 兜底的边界怎么定？",
                    "高频、明确、可解释的问题优先交给规则；模糊、边缘、关键词不明显的问题再交给 LLM。"
                    "更关键的是，这个边界不该靠感觉定，而应该靠样本集和评测结果不断调整。",
                    "考点：回答里要体现评测意识。",
                ),
            ],
        ),
        (
            "Redis / 记忆追问",
            [
                (
                    "为什么不把所有历史都直接塞给模型？",
                    "因为上下文窗口有限，太长的历史会带来成本和噪声。系统需要的是相关历史，而不是全部历史。"
                    "所以才会把短期上下文、摘要记忆和历史检索拆层管理。",
                    "考点：是否理解上下文越多不等于效果越好。",
                ),
                (
                    "摘要如果摘要错了怎么办？",
                    "摘要只能当作压缩层，不能当作唯一真相。"
                    "真正需要精确核对时，还是应该回到 MySQL 原始消息或精确历史检索结果。",
                    "考点：是否理解不同记忆层的权威性不同。",
                ),
            ],
        ),
        (
            "幂等 / 乐观锁追问",
            [
                (
                    "有 request_id 幂等了，为什么还会有状态冲突？",
                    "因为状态冲突通常来自不同请求，而 request_id 幂等只识别同一请求是否重复。"
                    "同一事件重复和不同事件并发，是两类完全不同的问题。",
                    "考点：能否把重复执行和并发覆盖分开。",
                ),
                (
                    "乐观锁冲突多了怎么办？",
                    "先看是不是状态粒度太粗、热点资源太集中，或者上层有无意义重试。"
                    "不要一冲突就说换悲观锁，应该先定位冲突来源，再决定是拆状态、串行化还是换锁策略。",
                    "考点：是否具备基本排障思路。",
                ),
            ],
        ),
        (
            "Outbox 追问",
            [
                (
                    "为什么不用 MQ 直接替代 Outbox？",
                    "因为 Outbox 解决的是本地事务边界问题：业务数据和待异步动作要先在同一事务里可靠提交。"
                    "MQ 很适合跨服务传递消息，但如果数据库提交和发 MQ 之间没有原子边界，照样会丢消息。",
                    "考点：是否理解 Outbox 和 MQ 关注层次不同。",
                ),
                (
                    "Outbox 为什么通常是至少一次，而不是 exactly-once？",
                    "因为 Worker 可能在外部动作成功后、回写完成状态前崩溃，恢复后这条任务会再次执行。"
                    "所以更现实的目标是至少一次投递加消费者幂等，而不是执着于物理只执行一次。",
                    "考点：是否理解业务效果幂等比字面一次执行更重要。",
                ),
            ],
        ),
        (
            "安全 / 沙箱追问",
            [
                (
                    "你说 Agent 要有沙箱，那沙箱具体限制什么？",
                    "我会从四层讲：文件访问范围、网络访问范围、命令执行范围、身份权限范围。"
                    "也就是说，它能读写哪些目录、能访问哪些域名、能执行哪些命令、用什么身份执行，这些都要被收紧。",
                    "考点：安全回答不能停留在空泛概念层。",
                ),
                (
                    "如果用户诱导 Agent 越权怎么办？",
                    "系统边界不能被用户一句话改变。真正的权限校验必须在执行层，而不是只在提示词里写‘不要这样做’。"
                    "超出权限范围的请求要直接拒绝，并写日志留痕。",
                    "考点：是否理解硬边界必须落在执行层。",
                ),
            ],
        ),
    ]

    number = 1
    for title, questions in groups:
        doc.add_heading(title, level=2)
        for question, answer, focus in questions:
            add_interview_question(doc, number, question, answer, focus)
            number += 1


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_beginner_core_concepts(doc)
    add_zero_foundation_expansion(doc)

    doc.add_heading("第一部分：项目定位与整体架构", level=1)
    doc.add_heading("1. 项目到底是什么", level=2)
    add_body(
        doc,
        "这是一个面向家具与扫地机器人场景的智能客服系统。用户可以咨询选购、清洁、保养、故障排查和使用报告。"
        "系统不会只把问题直接交给大模型，而是先识别用户正在处理的事情、召回相关历史、补充知识库证据，"
        "最后让大模型结合这些信息生成回答。"
    )
    add_callout(
        doc,
        "一句话定义",
        "大模型负责理解和表达，RAG 负责提供专业知识，数据库负责保存事实，LangGraph 负责组织处理流程。",
    )

    doc.add_heading("2. 它解决的核心问题", level=2)
    add_bullet(doc, "大模型不知道项目私有的家具资料，需要通过知识库补充。")
    add_bullet(doc, "大模型默认不会稳定记住长期用户信息，需要独立的记忆系统。")
    add_bullet(doc, "多轮客服不是闲聊，而是在持续解决一个问题，需要任务状态。")
    add_bullet(doc, "模型、数据库和向量库都可能失败，需要可靠保存、重试和降级。")
    add_bullet(doc, "用户希望看到连续输出，需要流式接口改善体验。")

    doc.add_heading("3. 整体架构的角色分工", level=2)
    add_comparison_table(
        doc,
        ["组成部分", "主要职责", "可以类比成什么"],
        [
            ["Streamlit", "展示聊天界面，收集用户输入，显示回答和反馈按钮", "客服窗口"],
            ["FastAPI", "接收请求、校验身份、组织响应", "服务前台"],
            ["LangGraph", "规定每个处理步骤及其先后关系", "业务流程导演"],
            ["大模型", "理解语言、决定是否调用工具、生成自然语言回答", "客服人员的大脑"],
            ["RAG", "从知识库寻找可用于回答的专业资料", "内部资料库"],
            ["MySQL", "永久保存用户、消息、任务和文档版本", "正式档案室"],
            ["Redis", "缓存近期常用数据，提高读取速度", "手边便签"],
            ["Chroma", "保存文本向量，支持语义相似搜索", "语义搜索引擎"],
            ["Worker", "后台处理索引和重试任务", "后台运营人员"],
        ],
        [1700, 5100, 2560],
    )

    doc.add_heading("4. 面试中的项目介绍", level=2)
    add_body(
        doc,
        "这是一个家具与扫地机器人领域的智能客服项目。系统基于 LangGraph 编排对话流程，使用工具调用连接知识库和外部数据；"
        "RAG 采用向量检索与 BM25 混合召回，再通过重排序和证据门控降低幻觉；对话侧使用 MySQL、Redis 和 Chroma "
        "构建短期上下文、会话摘要、长期用户记忆和跨会话历史召回；同时通过 Outbox 与异步 Worker "
        "保证消息持久化和向量索引之间的最终一致性。"
    )

    doc.add_heading("第二部分：一次用户请求的完整旅程", level=1)
    doc.add_heading("1. 示例问题", level=2)
    add_callout(doc, "示例", "用户说：我家的布艺沙发有咖啡污渍，我已经用水擦过了，还是没有去掉，应该怎么办？")
    doc.add_heading("2. 系统处理步骤", level=2)
    steps = [
        "前端生成本次请求的身份信息，包括用户、会话和请求编号，并将问题发送给后端。",
        "后端验证 API Key，创建本次响应通道，然后把问题交给 Agent 工作流。",
        "任务模块判断这是一个沙发清洁问题，并识别它是新任务还是已有任务的继续。",
        "历史模块判断用户是否在引用以前的对话；如果需要，则从历史数据中召回相关内容。",
        "上下文模块组合当前任务、最近对话、长期偏好和已召回历史。",
        "查询改写模块把“它”“那个”等模糊指代改成明确对象。",
        "大模型判断需要家具专业知识，于是调用知识库工具。",
        "RAG 检索沙发清洁资料，检查证据是否覆盖布艺、咖啡污渍和清洁意图。",
        "大模型基于检索证据和用户已尝试的动作生成回答，避免重复推荐无效方案。",
        "系统在同一个数据库事务中保存用户消息、助手回答和一条待索引任务。",
        "回答后更新当前任务，例如记录“已经用水擦过”“没有效果”和下一步建议。",
        "后台 Worker 最终把这一轮对话写入历史向量索引，供未来跨会话召回。",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("3. 为什么流程看起来这么长", level=2)
    add_body(
        doc,
        "如果只是做演示，用户问题直接交给大模型就够了。但客服系统需要回答可控、历史可信、数据不丢、请求可重试。"
        "这些工程要求决定了系统必须把“生成回答”拆成多个职责明确的步骤。"
    )

    doc.add_heading("第三部分：Agent、工具与工作流", level=1)
    doc.add_heading("1. 普通聊天模型与 Agent 的区别", level=2)
    add_comparison_table(
        doc,
        ["模式", "处理方式", "优点与限制"],
        [
            ["普通模型", "用户问题直接进入模型，再直接返回回答", "简单，但无法可靠访问私有数据和实时系统"],
            ["Agent", "模型可以判断需要什么信息，调用工具后再继续回答", "能力更强，但流程、权限和失败处理更复杂"],
        ],
        [1500, 3600, 4260],
    )
    add_body(
        doc,
        "在这个项目中，Agent 可以调用知识库、月度使用报告和天气工具。模型并不是亲自访问数据库，"
        "而是通过有明确输入输出的工具间接访问。这样更容易限制权限和定位错误。"
    )

    doc.add_heading("2. LangGraph 为什么存在", level=2)
    add_body(
        doc,
        "当系统只有一次模型调用时，普通函数已经够用。随着项目加入任务识别、历史召回、上下文组装、工具调用、持久化和记忆更新，"
        "流程逐渐变成一张图。LangGraph 的作用就是把这些节点和连接关系显式表达出来。"
    )
    add_bullet(doc, "流程清晰：每个节点只处理一种职责。")
    add_bullet(doc, "可以并行：任务识别与历史召回不互相依赖，可以同时执行。")
    add_bullet(doc, "方便降级：某个辅助节点失败时，可以记录错误并继续回答。")
    add_bullet(doc, "便于测试：可以单独验证节点顺序和状态传递。")

    doc.add_heading("3. 当前真实运行模式", level=2)
    add_callout(
        doc,
        "必须说准确",
        "项目实现了多专家 Agent 的路由和并行执行能力，也有相应测试；但当前默认主链路实际启用的是一个能够调用多个工具的 Agent。"
        "因此面试中应说“具备多 Agent 方案和测试，当前默认采用单 Agent 工具调用模式”，不要说成四个专家正在生产环境并行运行。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("4. 单 Agent 与多 Agent 的取舍", level=2)
    add_comparison_table(
        doc,
        ["方案", "适用情况", "主要代价"],
        [
            ["单 Agent + 工具", "问题边界较清晰，工具数量有限，需要控制成本和延迟", "一个模型承担全部判断，复杂问题可能相互干扰"],
            ["多 Agent", "多个领域需要独立策略，复杂问题需要并行分工", "调用次数、成本、延迟和结果冲突都会增加"],
        ],
        [1800, 4200, 3360],
    )

    doc.add_heading("第四部分：RAG 知识库系统", level=1)
    doc.add_heading("1. 为什么需要 RAG", level=2)
    add_body(
        doc,
        "大模型拥有通用知识，但它不知道项目自己的家具资料，也可能把不确定内容说得像真的。RAG 的目标不是让模型变聪明，"
        "而是在回答前把可信资料放到模型面前，并要求模型只依据资料回答。"
    )
    add_callout(doc, "核心流程", "用户问题 → 检索相关资料 → 检查证据 → 把证据交给模型 → 生成回答。")

    doc.add_heading("2. 知识库如何建立", level=2)
    for text in [
        "读取 TXT 或 PDF 文件。",
        "根据 FAQ、编号章节或普通文本结构进行切分。",
        "为每个小片段提取分类、意图、页码和关键词等元数据。",
        "把文字转换成向量并写入 Chroma。",
        "在 MySQL 中记录文档、版本、切片和索引任务状态。",
        "确认新版本写入完整后再将它设为当前有效版本。",
    ]:
        add_number(doc, text)

    doc.add_heading("3. 为什么必须切分文本", level=2)
    add_bullet(doc, "整份文档包含多个主题，直接检索会降低精度。")
    add_bullet(doc, "大段文本会占用大量模型上下文。")
    add_bullet(doc, "小片段更容易判断究竟是哪段资料支持了回答。")
    add_bullet(doc, "切分时保留少量重叠，可以避免关键信息刚好被切断。")

    doc.add_heading("4. 为什么同时使用向量检索和 BM25", level=2)
    add_comparison_table(
        doc,
        ["检索方式", "擅长什么", "典型例子"],
        [
            ["向量检索", "语义相近但字面不同的问题", "“沙发怎么去污”与“布艺坐具清洁方法”"],
            ["BM25", "精确关键词、型号、数字和错误码", "E12、具体温度、产品型号"],
        ],
        [1700, 3900, 3760],
    )
    add_body(
        doc,
        "系统把两种结果进行排名融合，再使用 BGE Reranker 对候选资料进行更精细的相关性判断。可以把它理解为："
        "前两种检索负责尽量不漏掉，重排序负责把真正相关的内容放到前面。"
    )

    doc.add_heading("5. 什么是证据门控", level=2)
    add_body(
        doc,
        "检索到内容不代表可以回答。系统还会检查对象、意图、材料、数字、型号和错误码是否得到覆盖，"
        "并检查文档版本是否仍然有效。证据不足时返回固定拒答，而不是让大模型凭常识补全。"
    )
    add_callout(
        doc,
        "面试关键词",
        "混合召回提高召回率，Reranker 提高排序精度，证据门控控制是否允许生成。",
        fill="ECFDF3",
        color=GREEN,
    )

    doc.add_heading("6. 为什么要做文档版本管理", level=2)
    add_body(
        doc,
        "旧方案只记录文件是否处理过，难以处理更新、删除和回滚。新版为同一文档保存多个版本，"
        "只有一个版本处于有效状态。这样可以明确回答“当前使用的是哪份资料”，也便于清理旧向量和审计索引状态。"
    )

    doc.add_heading("第五部分：记忆、历史召回与任务状态", level=1)
    doc.add_heading("1. 四种不同层次的记忆", level=2)
    add_comparison_table(
        doc,
        ["记忆类型", "保存内容", "主要用途"],
        [
            ["最近对话", "最近几轮原始问答", "处理“它、那个、继续”等紧邻追问"],
            ["会话摘要", "目标、事实、尝试、问题和约束的压缩结果", "避免长对话无限占用上下文"],
            ["长期用户记忆", "预算、家庭环境、偏好、避讳和拥有的物品", "跨会话提供个性化"],
            ["历史向量记忆", "过去每轮问答的语义索引", "寻找以前讨论过的相似问题"],
        ],
        [1850, 4200, 3310],
    )

    doc.add_heading("2. 记忆为什么不能只保存聊天记录", level=2)
    add_body(
        doc,
        "原始聊天记录完整但冗长，模型每次读取全部历史会越来越慢、越来越贵。摘要和结构化记忆的价值在于压缩信息，"
        "同时把目标、约束和已尝试动作这些重要内容突出出来。"
    )

    doc.add_heading("3. 精确历史与语义历史的区别", level=2)
    add_bullet(doc, "“我第一次问了什么”属于精确问题，必须直接查询 MySQL 原始消息。")
    add_bullet(doc, "“以前是不是聊过沙发清洁”属于语义问题，可以先用向量搜索寻找候选。")
    add_bullet(doc, "向量搜索得到的结果仍要回 MySQL 验证，避免把错误索引当成用户事实。")

    doc.add_heading("4. 什么是任务状态", level=2)
    add_body(
        doc,
        "聊天历史记录用户和助手说了什么；任务状态记录系统正在解决什么。一个任务可以包含主题、目标、约束、"
        "用户已经尝试的动作、产生的结果、拒绝的方案和下一步行动。"
    )
    add_callout(
        doc,
        "例子",
        "“沙发有污渍”是任务主题；“不使用刺激性清洁剂”是约束；“已经用清水擦过”是尝试；"
        "“仍然有印记”是结果；“检查面料护理标签”是下一步。",
    )

    doc.add_heading("5. 任务路由", level=2)
    add_comparison_table(
        doc,
        ["动作", "含义", "例子"],
        [
            ["new", "创建一个新问题任务", "从沙发清洁切换到扫地机器人故障"],
            ["continue", "继续当前正在处理的任务", "“我试过了，还是没用”"],
            ["resume", "恢复以前暂停或结束的相关任务", "“继续上次那个沙发问题”"],
            ["no_task", "不产生业务任务", "你好、谢谢、再见"],
        ],
        [1500, 3000, 4860],
    )

    doc.add_heading("第六部分：MySQL、Redis、Chroma 与 Outbox", level=1)
    doc.add_heading("1. 三种存储的边界", level=2)
    add_callout(
        doc,
        "必须记住",
        "MySQL 是事实源，Redis 是缓存，Chroma 是搜索索引。Redis 和 Chroma 可以根据 MySQL 重建，MySQL 中的业务事实不能随意丢失。",
        fill="ECFDF3",
        color=GREEN,
    )
    add_comparison_table(
        doc,
        ["存储", "保存什么", "为什么用它"],
        [
            ["MySQL", "用户、会话、消息、任务、记忆、文档版本、异步任务", "事务可靠、关系清晰、可审计"],
            ["Redis", "最近对话与常用记忆缓存", "速度快、支持过期时间"],
            ["Chroma", "知识片段和历史对话的向量", "支持语义相似搜索"],
        ],
        [1700, 4300, 3360],
    )

    doc.add_heading("2. 什么是事务", level=2)
    add_body(
        doc,
        "事务表示一组数据库操作要么全部成功，要么全部失败。例如保存一轮对话时，用户消息、助手消息和 Outbox 任务"
        "必须一起成功。如果中间任何一步失败，数据库会回滚，避免出现只有半轮对话的状态。"
    )

    doc.add_heading("3. 什么是幂等", level=2)
    add_body(
        doc,
        "网络不稳定时，前端可能重复发送同一个请求。幂等意味着同一请求执行一次或执行多次，最终结果相同。"
        "项目通过 request_id 和数据库唯一约束识别重复请求，避免重复保存消息或重复创建任务。"
    )

    doc.add_heading("4. 为什么需要 Outbox", level=2)
    add_body(
        doc,
        "MySQL 和 Chroma 是两个独立系统，无法简单使用同一个数据库事务。如果消息已经写入 MySQL，"
        "但写向量时 Chroma 暂时不可用，就会产生数据不一致。"
    )
    for text in [
        "在保存消息的同一个 MySQL 事务中，额外保存一条“待写入向量库”的任务。",
        "事务提交后，后台 Worker 领取任务。",
        "Worker 成功写入 Chroma 后，将任务标为完成。",
        "失败时按照逐渐增加的间隔重试。",
        "多次失败后进入 dead 状态，等待人工排查。",
    ]:
        add_number(doc, text)

    doc.add_heading("5. 最终一致性", level=2)
    add_body(
        doc,
        "Outbox 不保证 MySQL 和 Chroma 在每一毫秒都一致，而是保证在故障恢复后最终达到一致。"
        "这叫最终一致性，适合搜索索引、通知、日志和异步计算等场景。"
    )

    doc.add_heading("6. Worker 如何避免重复领取任务", level=2)
    add_body(
        doc,
        "多个 Worker 可以并行工作。领取任务时数据库会锁定选中的记录，同时跳过已被其他 Worker 锁定的记录。"
        "因此不同 Worker 可以安全地领取不同批次，提高处理能力。"
    )

    doc.add_heading("第七部分：项目从 Demo 到工程化的演进", level=1)
    doc.add_heading("1. 第一阶段：可运行的功能 Demo", level=2)
    add_body(
        doc,
        "最初版本的重点是证明功能可行：用户能够聊天，Agent 能调用工具，知识库能够检索，"
        "Redis 和 MySQL 能保存部分历史，Streamlit 能显示流式效果。"
    )
    add_bullet(doc, "优点：结构简单，能够快速验证想法。")
    add_bullet(doc, "问题：职责集中、写入可靠性不足、缺少幂等、知识库版本不可管理。")

    doc.add_heading("2. 第二阶段：数据访问分层", level=2)
    add_body(
        doc,
        "数据库操作被拆到 Repository 层。业务模块只表达“保存一轮对话”“读取用户记忆”，"
        "不再关心具体 SQL。这降低了耦合，也方便测试时使用模拟仓储。"
    )

    doc.add_heading("3. 第三阶段：持久化可靠性", level=2)
    add_body(
        doc,
        "旧版使用后台线程异步写 MySQL，进程退出时可能丢数据。新版将一轮消息同步放入事务，"
        "并加入 request_id 幂等和 Outbox，确保核心数据先可靠保存。"
    )

    doc.add_heading("4. 第四阶段：分层记忆与任务状态", level=2)
    add_body(
        doc,
        "项目从“保存最近聊天”升级为“理解用户正在解决什么”。新增会话摘要、长期用户记忆、历史向量召回和任务状态，"
        "使系统能够处理跨轮和跨会话问题。"
    )

    doc.add_heading("5. 第五阶段：RAG 工程化", level=2)
    add_body(
        doc,
        "知识库从简单向量搜索升级为文档版本管理、混合召回、重排序、证据检查和异步索引。"
        "目标从“能找到资料”转向“找到的资料可追踪、可更新、可判断是否足够回答”。"
    )

    doc.add_heading("6. 第六阶段：显式工作流与测试", level=2)
    add_body(
        doc,
        "LangGraph 将处理过程拆成清晰节点。测试覆盖工作流顺序、查询改写、任务恢复、历史验证、"
        "Outbox 事务、幂等和 RAG 证据判断。当前测试集合共包含 47 个测试。"
    )

    doc.add_heading("第八部分：准确性、可靠性与降级设计", level=1)
    doc.add_heading("1. 如何降低幻觉", level=2)
    add_bullet(doc, "需要领域知识时优先调用 RAG，而不是直接使用模型常识。")
    add_bullet(doc, "RAG 提示词要求模型只能依据检索资料回答。")
    add_bullet(doc, "型号、材料、错误码和数字必须在证据中出现。")
    add_bullet(doc, "证据不足时固定拒答。")
    add_bullet(doc, "历史语义结果必须回 MySQL 验证。")
    add_bullet(doc, "用户最新明确陈述的优先级高于旧记忆。")

    doc.add_heading("2. 如何处理服务故障", level=2)
    add_comparison_table(
        doc,
        ["故障", "当前思路", "影响"],
        [
            ["Redis 不可用", "记录警告，回退到 MySQL", "速度下降，但历史仍可读取"],
            ["Reranker 不可用", "退回融合排序或 BM25 判断", "排序质量可能下降"],
            ["向量检索不可用", "尝试使用 BM25", "语义召回能力下降"],
            ["记忆更新失败", "记录错误，不影响已保存的回答", "个性化暂时不完整"],
            ["Chroma 写入失败", "Outbox Worker 重试", "索引短时间滞后"],
            ["核心消息保存失败", "整轮请求返回错误", "避免向用户确认未保存的结果"],
        ],
        [1950, 4300, 3110],
    )

    doc.add_heading("3. 为什么有些错误可以降级，有些必须失败", level=2)
    add_body(
        doc,
        "判断标准是它是否影响核心事实。历史向量检索和记忆刷新属于增强功能，失败时可以先回答；"
        "消息持久化属于核心事实，如果回答已经告诉用户成功却没有保存，会产生更严重的不一致，因此应该明确失败。"
    )

    doc.add_heading("第九部分：项目当前的真实边界与优化路线", level=1)
    add_callout(
        doc,
        "面试态度",
        "优秀的项目介绍不是假装系统完美，而是能够准确说明当前状态、设计取舍和下一步优先级。",
        fill="FFF8E8",
        color=GOLD,
    )
    doc.add_heading("1. 当前主要问题", level=2)
    issues = [
        "当前流式输出是完整答案生成后再逐字发送，视觉上流式，但首字延迟没有真正降低。",
        "记忆刷新会重复处理最近多条用户消息，可能产生不必要的大模型调用。",
        "BM25 索引在每次请求中重新构建，知识库规模扩大后会成为性能瓶颈。",
        "健康检查会更新数据库状态，不符合健康检查通常应为只读操作的原则。",
        "知识索引 Worker 在特定中断时可能残留旧版本向量，需要加强幂等清理。",
        "Streamlit 直接连接 MySQL，界面层和数据库耦合过深。",
        "Docker 部署配置目前仍需修正和完整验证。",
        "多 Agent 路径已经实现并测试，但默认生产路径尚未启用。",
    ]
    for issue in issues:
        add_bullet(doc, issue)

    doc.add_heading("2. 推荐优化顺序", level=2)
    priorities = [
        "先修复 Docker 和环境变量，让项目能够稳定、一键复现。",
        "实现真正的模型 token 流式输出，降低用户感知延迟。",
        "把记忆提取移到异步 Worker，并只处理本轮新增消息。",
        "按知识库版本缓存 BM25 索引，避免每次全量重建。",
        "把健康检查改为纯读取，并补充延迟、失败率和待处理任务数监控。",
        "增强向量版本清理的幂等性，保证 Worker 重试不会留下脏数据。",
        "最后再评估是否真正需要默认启用多 Agent，避免为了架构复杂而复杂。",
    ]
    for item in priorities:
        add_number(doc, item)

    doc.add_heading("第十部分：面试表达模板", level=1)
    doc.add_heading("1. 30 秒项目介绍", level=2)
    add_callout(
        doc,
        "推荐表达",
        "我做的是一个家具与扫地机器人领域的智能客服系统。它不是简单调用大模型，而是通过 LangGraph 组织任务识别、"
        "历史召回、上下文组装、工具调用和持久化。知识侧采用向量检索与 BM25 的混合 RAG，并使用重排序和证据门控降低幻觉；"
        "对话侧使用 MySQL、Redis 和 Chroma 构建分层记忆，同时通过 Outbox 和后台 Worker 保证异步索引的最终一致性。",
    )

    doc.add_heading("2. 两分钟项目介绍", level=2)
    add_body(
        doc,
        "项目的业务目标是帮助用户处理家具选购、清洁、保养、售后以及扫地机器人故障。"
        "一次请求进入后，系统会并行判断当前任务和是否需要召回历史，再将任务状态、最近对话、长期用户偏好和历史证据组装为上下文。"
        "如果问题有模糊指代，系统先做查询改写；随后大模型可以自主调用知识库、使用报告或天气工具。"
    )
    add_body(
        doc,
        "知识库采用混合检索：向量搜索负责语义相似，BM25 负责错误码和型号等精确匹配，之后进行融合和重排序。"
        "系统还会检查证据是否真正覆盖对象、意图和关键术语，不足时拒答。回答完成后，用户消息、助手消息和 Outbox 任务在同一个 MySQL "
        "事务中提交；后台 Worker 再将对话和文档写入 Chroma。这样即使向量库暂时故障，核心对话也不会丢失。"
    )
    add_body(
        doc,
        "当前项目已经实现并测试了多 Agent 专家路由，但默认使用的是单 Agent 工具调用路径。下一步重点是实现真正的 token 流式、"
        "异步化记忆提取、缓存 BM25 索引，并修复部署配置。"
    )

    doc.add_heading("3. 表达提醒", level=2)
    add_comparison_table(
        doc,
        ["不准确说法", "更好的说法"],
        [
            ["项目当前是四个 Agent 并行运行", "项目实现了多 Agent 路径，但默认运行单 Agent 工具调用路径"],
            ["系统完全不会产生幻觉", "系统通过证据门控降低幻觉，但无法保证绝对为零"],
            ["Redis 用来永久保存聊天记录", "MySQL 永久保存，Redis 只缓存近期数据"],
            ["Outbox 保证立即一致", "Outbox 保证异步操作最终一致"],
            ["所有代码都是我独立写的", "AI 辅助完成初版，我正在通过审查、测试和亲自优化完成项目接管"],
        ],
        [4300, 5060],
    )

    doc.add_heading("第十一部分：场景演练与自测", level=1)
    doc.add_heading("1. 场景一：用户说“它还是不工作”", level=2)
    add_body(
        doc,
        "系统先判断“它”需要结合上下文解释。如果当前任务是扫地机器人回充失败，就将问题改写为明确的扫地机器人问题；"
        "任务状态会提供用户之前已经尝试过的动作；RAG 再检索对应故障资料。回答时不能原样重复用户已经失败的方案。"
    )

    doc.add_heading("2. 场景二：用户问“我第一次问了什么”", level=2)
    add_body(
        doc,
        "这是精确历史问题，系统不应依赖语义相似搜索，而是直接查询当前会话中第一条用户消息。"
        "这是因为精确问题要求确定性，向量搜索只能提供近似结果。"
    )

    doc.add_heading("3. 场景三：Chroma 暂时不可用", level=2)
    add_body(
        doc,
        "对话消息仍然应先可靠写入 MySQL，待索引任务保存在 Outbox。Worker 后续重试写入 Chroma。"
        "在知识检索阶段，如果向量检索不可用，可以尝试 BM25；但如果最终证据仍不足，则拒答。"
    )

    doc.add_heading("4. 场景四：相同请求被发送两次", level=2)
    add_body(
        doc,
        "两次请求携带相同 request_id。数据库唯一约束会识别第二次操作，系统复用已经保存的消息或任务，"
        "而不是创建重复记录。这就是幂等设计。"
    )

    doc.add_heading("5. 自测问题", level=2)
    self_checks = [
        "你能否不用技术名词，向非技术人员解释这个项目的业务价值？",
        "你能否完整描述一次请求从前端到数据库的旅程？",
        "你能否说明 MySQL、Redis 和 Chroma 的职责边界？",
        "你能否解释 RAG 为什么需要混合召回、重排序和证据门控？",
        "你能否解释任务状态与聊天历史的区别？",
        "你能否说明 Outbox 为什么比后台临时线程可靠？",
        "你能否准确说出当前真正启用的是单 Agent 还是多 Agent？",
        "你能否主动说出项目三个真实问题和对应优化方案？",
    ]
    for item in self_checks:
        add_bullet(doc, item)

    doc.add_heading("6. 建议学习顺序", level=2)
    plan = [
        "第一天：掌握项目一句话定义、整体架构和一次请求流程。",
        "第二天：掌握 Agent、工具调用、LangGraph 和当前运行模式。",
        "第三天：掌握 RAG、Embedding、BM25、Reranker 和证据门控。",
        "第四天：掌握四层记忆、任务状态、精确历史与语义历史。",
        "第五天：掌握 MySQL、Redis、Chroma、事务、幂等、Outbox 和 Worker。",
        "第六天：复述项目演进路线，并理解每次升级解决的问题。",
        "第七天：练习 30 秒和两分钟介绍，随机回答高频问题。",
        "之后：亲自完成至少一个优化，让项目经验真正变成你的实践经验。",
    ]
    for item in plan:
        add_number(doc, item)

    doc.add_heading("附录：核心术语速查", level=1)
    glossary = [
        ["LLM", "大语言模型，负责理解和生成自然语言。"],
        ["Agent", "能够根据目标选择和调用工具的大模型应用。"],
        ["Tool Calling", "模型以结构化方式请求系统执行知识库、数据库或其他能力。"],
        ["RAG", "先检索外部知识，再让模型基于检索结果生成回答。"],
        ["Embedding", "将文本转换为能够计算相似度的数字向量。"],
        ["Vector Store", "保存和搜索向量的系统，本项目使用 Chroma。"],
        ["BM25", "基于词频和稀有程度的关键词检索算法。"],
        ["Reranker", "对初次召回的少量候选进行更精细的相关性排序。"],
        ["Chunk", "文档切分后用于检索的小片段。"],
        ["Metadata", "片段附带的分类、页码、版本和意图等描述信息。"],
        ["LangGraph", "用于描述和执行 Agent 状态流程图的框架。"],
        ["SSE", "服务器通过同一个 HTTP 连接持续向客户端发送事件。"],
        ["Transaction", "一组数据库操作要么全部成功，要么全部失败。"],
        ["Idempotency", "相同请求执行多次与执行一次具有相同最终效果。"],
        ["Optimistic Lock", "使用版本号检测并发修改冲突。"],
        ["Outbox", "与业务数据一起提交的待处理任务记录。"],
        ["Worker", "后台不断领取并执行异步任务的进程。"],
        ["Final Consistency", "不同系统允许短暂不一致，但故障恢复后最终一致。"],
        ["Cache", "为了提高读取速度而保存的可重建数据副本。"],
        ["Source of Truth", "被视为最终权威的数据来源，本项目主要是 MySQL。"],
        ["Sandbox", "给 Agent 划定的受限运行空间，用来限制文件、网络、命令和权限范围。"],
        ["Least Privilege", "最小权限原则，即每个 Agent 只拿到完成当前任务所需的最少能力。"],
        ["Degradation", "部分能力故障时使用较弱方案继续提供服务。"],
        ["Hallucination", "模型生成看似合理但缺乏事实依据的内容。"],
    ]
    add_comparison_table(doc, ["术语", "通俗解释"], glossary, [2200, 7160])

    doc.add_heading("结语：怎样把 AI 生成的项目真正变成自己的项目", level=1)
    add_body(
        doc,
        "面试官真正关心的通常不是每一行代码是不是你亲手敲出来，而是你是否理解系统、是否能发现问题、"
        "是否能做出合理取舍，以及出现故障时是否知道从哪里排查。"
    )
    add_body(
        doc,
        "你不需要假装这个项目完全由自己从零编写。更可信的表达是：AI 辅助完成了初版实现，"
        "你通过架构梳理、测试验证、问题审查和后续优化逐步接管项目。只要你能独立解释设计原因，"
        "并亲自完成几个关键优化，这段经历就会逐渐成为真正属于你的项目经验。"
    )
    add_callout(
        doc,
        "最终目标",
        "当面试官提出“为什么这样设计”“如果某个组件挂了怎么办”“下一步如何优化”时，"
        "你能够从业务目标、数据一致性、准确性、成本和延迟几个角度作出判断，而不是背诵框架名称。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_follow_up_drill_appendix(doc)
    add_interview_bank(doc)
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
