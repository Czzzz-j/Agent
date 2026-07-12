from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DESKTOP_DIR = Path(r"C:\Users\陈梓杰\Desktop\陈梓杰简历")
OUTPUT = DESKTOP_DIR / "家具智能客服项目理解与面试手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN_FILL = "ECFDF3"
GOLD_FILL = "FFF8E8"
RED_FILL = "FFF1F2"
GREEN = "1F7A4D"
GOLD = "7A5A00"
RED = "9B1C1C"

TABLE_WIDTH = 9360
TABLE_INDENT = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths: list[int], indent=TABLE_INDENT):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def setup_page_furniture(doc: Document):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = "家具智能客服项目接收手册"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.runs[0], 9, MUTED)
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("第 ")
    set_font(r, 9, MUTED)
    page_field(fp)
    r = fp.add_run(" 页")
    set_font(r, 9, MUTED)


def add_body(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        set_font(run)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    for run in p.runs:
        set_font(run)
    return p


def add_callout(doc: Document, label: str, text: str, fill=CALLOUT, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    r = p.add_run(f"{label}: ")
    set_font(r, 10.5, color, True)
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = ""
        cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, 10.5, DARK_BLUE, True)
        set_paragraph_spacing(p, after=2)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = ""
            if idx == 0:
                cell_shading(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            r = p.add_run(text)
            set_font(r, 10, INK, bold=(idx == 0))
            set_paragraph_spacing(p, after=2, line=1.15)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_title_page(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("家具智能客服项目\n接收与面试通关手册")
    set_font(r, 28, INK, True)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("写给零基础接手者：先看懂项目，再讲赢面试")
    set_font(r, 14, DARK_BLUE, True)

    doc.add_paragraph()
    add_callout(
        doc,
        "使用目标",
        "这不是代码说明书，也不是名词背诵本。它的目标是让一个没做过 Agent 项目的人，能从业务画面、请求链路、技术因果、工程取舍和面试表达五个层面接住这个项目。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_matrix(
        doc,
        ["读者状态", "读完后应该能做到"],
        [
            ["完全不懂 Agent", "能用生活类比解释 Agent、路由、RAG、记忆和 Outbox"],
            ["没看过代码", "能讲出一条用户请求从入口到后台索引的完整过程"],
            ["准备面试", "能回答第一问、追问和继续追问，不只背概念"],
        ],
        [2300, 7060],
    )
    add_body(doc, f"版本说明：小白接收版，生成日期 {datetime.now().strftime('%Y-%m-%d')}。")
    doc.add_page_break()


def add_navigation(doc: Document):
    doc.add_heading("阅读路线", level=1)
    add_body(
        doc,
        "建议不要从题库开始背。小白最容易失败的地方，是知道很多名词，但不知道这些名词为什么必须连在一起。"
        "正确路线是先建立项目画面，再理解一条请求，再把技术点放进链路，最后才背面试表达。"
    )
    add_matrix(
        doc,
        ["阶段", "重点", "你要形成的能力"],
        [
            ["第 0-1 章", "项目价值和业务画面", "知道这个项目到底卖什么"],
            ["第 2 章", "一条请求的完整故事", "能把系统从用户提问讲到后台索引"],
            ["第 3-4 章", "技术因果链和架构关系", "知道每个技术为什么出现、连着谁"],
            ["第 5 章", "从 Demo 到工程化", "能讲项目演进和设计取舍"],
            ["第 6-8 章", "面试表达、追问、速记", "能把理解转换成面试回答"],
        ],
        [1700, 3400, 4260],
    )


def add_chapter_0(doc: Document):
    doc.add_heading("第 0 章：这个项目到底卖什么", level=1)
    add_callout(
        doc,
        "先记住",
        "这个项目卖的不是“我接了一个大模型 API”，而是“我把智能客服做成了一条可检索、可记忆、可恢复、可解释的业务链路”。",
        fill=GREEN_FILL,
        color=GREEN,
    )
    add_body(
        doc,
        "面试官看一个 Agent 项目，真正想知道的不是页面能不能聊天，而是你有没有把真实业务里的复杂问题处理掉。"
        "真实客服系统会遇到知识不准、上下文断片、重复请求、并发覆盖、后台任务丢失、服务故障降级等问题。"
        "你的项目价值就在于：它不是让模型裸答，而是给模型配了一整套工程护栏。"
    )
    add_matrix(
        doc,
        ["普通 Demo", "你的项目应该卖的点"],
        [
            ["用户问一句，模型答一句", "先识别任务，再找证据，再结合记忆回答"],
            ["知识靠模型自己编", "用混合 RAG 找企业知识，证据不足就拒答"],
            ["历史靠上下文硬塞", "用短期缓存、会话摘要、长期画像和历史召回分层处理"],
            ["保存和索引顺手写", "用事务、幂等、乐观锁、Outbox 和 Worker 做可靠链路"],
            ["出错就报错", "辅助节点可降级，核心事实不能假装成功"],
        ],
        [4200, 5160],
    )
    doc.add_heading("面试官为什么会觉得它有价值", level=2)
    add_bullet(doc, "它覆盖了 Agent 应用的主链路：路由、工具、RAG、记忆、回答、持久化、异步任务。")
    add_bullet(doc, "它不只讲效果，还讲可靠性：幂等、乐观锁、Outbox、重试、降级。")
    add_bullet(doc, "它能体现工程判断：哪些交给模型，哪些必须交给规则、数据库和事务。")
    add_callout(
        doc,
        "小白表达模板",
        "我这个项目不是简单聊天机器人，而是一个面向家具客服场景的 Agent 系统。它把大模型、知识库、历史记忆和后台可靠任务组合起来，让客服回答既有证据，也能在多轮对话和系统故障下保持稳定。",
        fill=CALLOUT,
        color=DARK_BLUE,
    )


def add_chapter_1(doc: Document):
    doc.add_heading("第 1 章：先建立业务画面", level=1)
    add_body(
        doc,
        "先不要急着理解 LangGraph。先想一个用户真实来问客服的场景：他家沙发被咖啡弄脏了，之前已经用水擦过，但没效果。"
        "一个好客服不能只说“用清洁剂试试”，因为他要知道用户问的是哪类产品、以前尝试过什么、知识库有没有支持、有没有风险提醒。"
    )
    add_matrix(
        doc,
        ["系统角色", "生活类比", "在项目里负责什么"],
        [
            ["用户", "来咨询的顾客", "提出问题，可能说得模糊，也可能延续上文"],
            ["Agent", "会查资料的客服", "理解问题，决定是否调用工具和知识库"],
            ["路由", "分诊台", "判断这是闲聊、延续问题、售前咨询还是故障排查"],
            ["RAG", "内部资料库", "把产品资料、保养方法、故障说明查出来"],
            ["记忆系统", "客服的服务记录", "记住当前会话和跨会话的重要信息"],
            ["MySQL", "正式档案库", "保存核心消息、任务状态、文档版本等事实"],
            ["Redis", "桌面便签", "快速读取近期上下文和临时状态"],
            ["Chroma", "按意思找资料的索引", "支持语义检索和历史语义召回"],
            ["Outbox + Worker", "后台工单队列", "保证异步索引任务失败后还能重试"],
        ],
        [1700, 2300, 5360],
    )
    add_callout(
        doc,
        "小白要跨过的第一道坎",
        "不要把项目理解成“一个模型”。它更像一个客服团队：前台接待、分诊、资料库、病历、后台工单都在协作。模型只是团队里的核心表达者，不是全部系统。",
        fill=GOLD_FILL,
        color=GOLD,
    )
    doc.add_heading("项目里的核心矛盾", level=2)
    add_body(
        doc,
        "大模型擅长理解语言和组织回答，但它不擅长保证事实一定来自企业知识，也不擅长保证数据库一致性。"
        "所以系统设计的核心矛盾是：既要利用模型的语言能力，又不能把所有关键决定都交给模型自由发挥。"
    )
    add_bullet(doc, "事实问题交给 RAG 和证据门控。")
    add_bullet(doc, "流程问题交给 LangGraph。")
    add_bullet(doc, "重复请求交给 request_id 幂等。")
    add_bullet(doc, "并发覆盖交给乐观锁。")
    add_bullet(doc, "异步可靠性交给 Outbox 和 Worker。")


def add_chapter_2(doc: Document):
    doc.add_heading("第 2 章：一条用户请求的完整故事线", level=1)
    add_callout(
        doc,
        "本章目标",
        "读完这一章，你要能把系统从“用户发一句话”讲到“后台完成索引”这一整条链路。面试官最容易通过这条链路判断你是真懂还是背词。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    doc.add_heading("示例问题", level=2)
    add_body(doc, "用户说：我家的布艺沙发有咖啡污渍，我已经用水擦过了，还是没有去掉，应该怎么办？")
    doc.add_heading("系统怎么处理", level=2)
    steps = [
        "前端把用户问题、session_id、user_uuid 和 request_id 发给后端。",
        "任务路由先判断：这不是闲聊，而是一个清洁/维护类业务问题。",
        "历史召回检查用户之前有没有提过同一张沙发、材质、已尝试的方法。",
        "上下文组装把当前问题、历史信息、用户偏好和任务状态整理成模型可用的材料。",
        "查询改写把口语化问题变成更适合检索的表达，例如“布艺沙发咖啡污渍清洁方法”。",
        "RAG 同时走向量检索和 BM25 检索，再融合、重排，筛出证据。",
        "证据门控判断资料是否真的覆盖“布艺沙发”“咖啡污渍”“已用水擦过仍无效”。",
        "如果证据足够，Agent 基于证据和历史生成回答；如果证据不足，系统拒答或追问。",
        "回答生成后，用户消息和助手消息按 request_id 幂等写入 MySQL。",
        "任务状态用乐观锁更新，防止另一个请求同时覆盖当前状态。",
        "同一个 MySQL 事务里写入 Outbox 任务，记录“这轮对话要补向量索引”。",
        "Worker 后台领取 Outbox 任务，把这轮对话写入 Chroma，供以后语义召回使用。",
    ]
    for step in steps:
        add_number(doc, step)
    add_heading = doc.add_heading
    add_heading("为什么这条链路看起来这么长", level=2)
    add_body(
        doc,
        "因为真实业务不是“答一句话”这么简单。前半段要保证回答准确，后半段要保证数据可靠。"
        "如果只追求短链路，系统可能会答得快，但会出现幻觉、丢历史、重复写入或索引不一致。"
    )
    add_matrix(
        doc,
        ["链路阶段", "目标", "保护机制"],
        [
            ["识别阶段", "知道问题属于哪类任务", "任务路由、延续意图识别"],
            ["查证阶段", "找到可引用的知识证据", "混合 RAG、重排序、证据门控"],
            ["回答阶段", "组织自然语言回复", "Agent、工具调用、多 Agent 合成"],
            ["保存阶段", "核心消息不能丢、不能重复", "MySQL 事务、request_id 幂等"],
            ["异步阶段", "后台索引迟早要补齐", "Outbox、Worker、重试、僵死恢复"],
        ],
        [1800, 3300, 4260],
    )


def mechanism_block(
    doc: Document,
    title: str,
    analogy: str,
    position: str,
    absence: str,
    relation: str,
    interview: str,
):
    doc.add_heading(title, level=2)
    add_matrix(
        doc,
        ["理解角度", "内容"],
        [
            ["生活类比", analogy],
            ["项目位置", position],
            ["没有它会怎样", absence],
            ["上下游关系", relation],
            ["面试怎么说", interview],
        ],
        [1900, 7460],
    )


def add_chapter_3(doc: Document):
    doc.add_heading("第 3 章：核心机制为什么被迫出现", level=1)
    add_body(
        doc,
        "这一章按因果链讲，不孤立背名词。你要形成的感觉是：系统每变复杂一步，就会被迫引入一个机制来解决新的麻烦。"
    )
    mechanism_block(
        doc,
        "1. LangGraph：流程变复杂后，需要一个导演",
        "像医院流程单，决定挂号后去分诊，检查后去医生，医生后去缴费。",
        "它在请求主链路里负责组织节点：路由、历史召回、上下文组装、回答、持久化、记忆刷新。",
        "流程会散落在大量 if-else 里，出错时不知道卡在哪一步，也不容易做降级和测试。",
        "上游接收用户请求，下游连接路由、RAG、Agent、持久化和记忆刷新。",
        "LangGraph 不提升模型智商，它提升的是复杂流程的可控性、可观测性和可恢复性。",
    )
    mechanism_block(
        doc,
        "2. 路由：问题进来后，先决定走哪条路",
        "像医院分诊台，先判断去皮肤科、骨科还是普通咨询。",
        "任务路由判断 new、continue、resume、no_task；专家路由决定是否调用家具、设备、报表或通用专家。",
        "所有问题都会被同一种方式处理，可能找错专家、查错资料、浪费模型调用。",
        "上游接用户问题和历史状态，下游影响 RAG 范围、专家选择和上下文组装。",
        "路由不是简单分类器，而是系统资源调度入口。",
    )
    mechanism_block(
        doc,
        "3. 多 Agent：问题跨领域后，需要分工",
        "像多科室会诊，不是一个医生硬扛全部问题。",
        "当问题同时涉及家具知识、设备故障、报表数据时，专家 Agent 分别处理，再由答案合成器统一表达。",
        "单 Agent 提示词会越来越杂，容易上下文混乱，复杂问题回答不稳定。",
        "上游依赖专家路由，下游依赖答案合成器处理去重、冲突和统一话术。",
        "多 Agent 的价值是职责拆分，不是越多越高级；简单问题保留单 Agent 更稳。",
    )
    mechanism_block(
        doc,
        "4. RAG：模型不知道企业私有知识，必须先查证据",
        "像客服翻内部手册，而不是凭记忆瞎说。",
        "它在回答前检索家具资料、扫地机器人故障文档和维护保养知识。",
        "模型可能编出不存在的材质、保修规则或故障处理方法。",
        "上游接查询改写后的问题，下游把证据交给 Agent，并由证据门控决定能否回答。",
        "RAG 不是为了显得高级，而是让答案有企业知识依据，证据不足时宁可拒答。",
    )
    mechanism_block(
        doc,
        "5. 记忆系统：多轮对话不能只靠当前一句话",
        "像客服看服务记录，知道用户之前问过什么、试过什么。",
        "短期 Redis 支撑当前会话，会话摘要压缩长对话，长期画像跨会话复用用户偏好，历史召回解决指代。",
        "用户说“它还是不行”时，系统不知道“它”是谁，也不知道哪些方案已经试过。",
        "上游接用户身份和会话，下游给上下文组装、任务路由和回答节点提供历史材料。",
        "记忆不是把所有聊天记录塞给模型，而是把相关历史分层取出来。",
    )
    mechanism_block(
        doc,
        "6. Redis：高频短期信息要快，但不能当事实源",
        "像桌面便签，查得快，但正式档案仍在档案柜。",
        "它负责缓存近期上下文和短期状态，提升多轮对话读取速度。",
        "每次都查 MySQL 会更慢；如果只写 Redis 不写 MySQL，Redis 过期后事实会丢。",
        "上游来自会话读写，下游服务历史召回和上下文组装；核心事实仍以 MySQL 为准。",
        "Redis 是可丢失、可重建的加速层，不是永久事实库。",
    )
    mechanism_block(
        doc,
        "7. request_id 幂等：同一请求重复来，结果不能重复",
        "像快递单号，同一个单号重复扫描，不应该生成两件快递。",
        "保存用户消息、助手消息、任务事件时用 request_id 和唯一约束去重。",
        "用户连点、网络重试或网关重试会导致同一轮消息被保存两次，甚至产生重复 Outbox 任务。",
        "上游来自前端/网关请求，下游保护 MySQL 消息、任务事件和 Outbox 投递。",
        "幂等不是保证物理只执行一次，而是保证重复执行后的业务结果一致。",
    )
    mechanism_block(
        doc,
        "8. 乐观锁：不同请求同时改状态，不能互相覆盖",
        "像多人编辑同一份文档，保存前要确认版本没被别人改过。",
        "任务状态更新时检查 state_version，版本一致才更新，冲突后重新读取再处理。",
        "两个不同 request_id 同时更新同一任务，后写可能把前写覆盖掉。",
        "上游来自任务路由和回答结果，下游保护任务状态、事实提取和下一步行动建议。",
        "幂等防重复，乐观锁防覆盖。两者解决的不是同一个问题。",
    )
    mechanism_block(
        doc,
        "9. Outbox：MySQL 成功后，外部索引仍可能失败",
        "像前台收下工单后，先把工单登记进系统，再交给后台慢慢处理。",
        "消息保存时同事务插入 Outbox 任务，Worker 后台把对话或知识写入 Chroma。",
        "如果保存 MySQL 后直接写 Chroma，中间崩溃会造成 MySQL 有消息、Chroma 没索引。",
        "上游接 MySQL 事务，下游连接 Worker、Chroma、重试和僵死任务恢复。",
        "Outbox 解决双写可靠性，通常是至少一次投递，所以消费者必须幂等。",
    )
    mechanism_block(
        doc,
        "10. Agent 沙箱：模型能做事后，必须限制它能做什么",
        "像给实习客服一个工牌：可以查资料、填工单，但不能随便退款、删库或外发隐私。",
        "它处在 Agent 调工具和外部系统之间，负责控制可用工具、参数范围、文件/网络/数据库权限和高风险操作审批。",
        "模型一旦误判、被提示词攻击或拿到过大权限，可能调用危险工具、泄露数据、越权修改业务状态。",
        "上游接 Agent 的工具调用意图，下游保护数据库、文件系统、外部 API、用户隐私和后台任务入口。",
        "Agent 沙箱不是限制模型能力，而是把模型能力关进可控边界：低风险自动执行，高风险校验或人工确认。",
    )


def add_chapter_4(doc: Document):
    doc.add_heading("第 4 章：把技术点串成系统架构", level=1)
    add_callout(
        doc,
        "本章一句话",
        "系统不是一堆技术名词，而是三条线一起工作：前台回答线、事实保存线、后台索引线。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_matrix(
        doc,
        ["链路", "包含模块", "解决的问题"],
        [
            ["前台回答线", "FastAPI / LangGraph / 路由 / RAG / Agent / 答案合成", "让用户尽快得到有依据的回答"],
            ["事实保存线", "MySQL / request_id / 事务 / 乐观锁", "保证核心消息和任务状态可靠、不重复、不互相覆盖"],
            ["后台索引线", "Outbox / Worker / Chroma / 重试恢复", "保证异步索引最终补齐，前台不被慢任务卡住"],
        ],
        [1900, 4200, 3260],
    )
    doc.add_heading("模块之间谁保护谁", level=2)
    add_bullet(doc, "LangGraph 保护流程清晰，避免节点散落。")
    add_bullet(doc, "路由保护资源使用，避免什么问题都查一遍、问一遍。")
    add_bullet(doc, "RAG 和证据门控保护事实准确性，避免模型硬编。")
    add_bullet(doc, "记忆系统保护多轮连续性，避免用户说“它”时系统断片。")
    add_bullet(doc, "MySQL 保护核心事实，Redis 和 Chroma 都可以围绕它重建。")
    add_bullet(doc, "request_id 保护重复请求，乐观锁保护并发更新。")
    add_bullet(doc, "Outbox 保护跨存储双写，Worker 负责最终补齐。")
    add_bullet(doc, "Agent 沙箱保护权限边界，避免模型把“会调用工具”变成“能随便操作系统”。")
    doc.add_heading("小白最容易混的概念对照", level=2)
    add_matrix(
        doc,
        ["概念 A", "概念 B", "区别"],
        [
            ["LangChain", "LangGraph", "LangChain 偏模型、工具、检索的接线；LangGraph 偏有状态流程编排。"],
            ["Agent", "多 Agent", "Agent 是能调工具的执行单元；多 Agent 是多个角色分工协作。"],
            ["任务路由", "专家路由", "任务路由决定这是什么事；专家路由决定交给谁处理。"],
            ["RAG", "微调", "RAG 解决知识更新和可追溯；微调解决稳定行为、格式和风格。"],
            ["Redis", "MySQL", "Redis 是加速层，可重建；MySQL 是事实源，不能随便丢。"],
            ["幂等", "乐观锁", "幂等防同一请求重复；乐观锁防不同请求并发覆盖。"],
            ["Outbox", "后台线程", "Outbox 有持久任务和重试恢复；临时线程进程崩溃就可能丢。"],
            ["Outbox", "MQ", "Outbox 先解决本地事务边界；MQ 更偏跨服务消息传递。"],
            ["Agent 沙箱", "普通工具调用", "普通工具调用只关注能不能调；沙箱关注能调哪些、参数是否合法、风险是否需要拦截。"],
            ["RAG", "模型训练", "RAG 把外部知识查出来给模型用；训练/微调更偏改变模型行为、风格和任务习惯。"],
        ],
        [1700, 1700, 5960],
    )


def add_chapter_5(doc: Document):
    doc.add_heading("第 5 章：项目从 Demo 到工程化的演进", level=1)
    add_body(
        doc,
        "这一章是面试里很加分的部分。不要只说“我用了什么技术”，要讲“原来有什么问题，所以我怎么改，改完带来什么收益”。"
    )
    add_matrix(
        doc,
        ["阶段", "旧问题", "新设计", "收益"],
        [
            ["功能 Demo", "能聊天，但职责混在一起", "先跑通前端、后端、模型、工具、知识库", "证明业务想法可行"],
            ["数据访问分层", "业务逻辑直接关心 SQL 和存储细节", "引入 Repository / Service 分层", "降低耦合，便于测试"],
            ["持久化可靠性", "后台线程写库可能丢数据", "MySQL 事务 + request_id 幂等 + Outbox", "核心事实先可靠保存"],
            ["记忆与任务状态", "多轮指代和跨会话信息容易断", "短期缓存、摘要、长期画像、任务状态", "能处理“它还是不行”这类问题"],
            ["RAG 工程化", "只做向量检索容易错召回", "向量 + BM25 + 融合 + 重排 + 证据门控", "减少幻觉，提高可解释性"],
            ["工作流显式化", "流程散在函数和条件判断中", "LangGraph 节点化编排", "并行、降级和测试更清楚"],
            ["权限边界", "Agent 能调工具后风险变大", "工具白名单、参数校验、权限分级、人工确认", "降低越权调用和提示词攻击风险"],
        ],
        [1450, 2500, 2700, 2710],
    )
    add_callout(
        doc,
        "面试表达",
        "项目演进不是为了堆框架，而是每一步都在解决真实问题：先让它能跑，再让它可靠，再让它可解释，再让它可维护。",
        fill=GREEN_FILL,
        color=GREEN,
    )


def add_chapter_6(doc: Document):
    doc.add_heading("第 6 章：面试表达训练", level=1)
    add_heading = doc.add_heading
    add_heading("30 秒版本", level=2)
    add_callout(
        doc,
        "直接背这一段",
        "我做的是一个面向家具和扫地机器人场景的智能客服 Agent 系统。它不是简单调用大模型，而是通过 LangGraph 组织任务路由、历史召回、RAG 检索、回答生成、消息持久化和后台索引。系统用混合 RAG 降低幻觉，用分层记忆解决多轮指代，用 request_id 幂等、乐观锁和 Outbox 保证重复请求、并发更新和异步索引的可靠性，同时通过工具权限边界避免 Agent 越权操作。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_heading("1 分钟版本", level=2)
    add_body(
        doc,
        "这个项目面向家具客服、导购、保养和扫地机器人售后场景。用户的问题可能很模糊，也可能延续上一轮，所以系统先通过任务路由判断当前请求属于新任务、延续任务还是闲聊。"
        "然后系统会召回近期对话、会话摘要、长期画像和历史语义记录，把这些信息组装成上下文。需要专业知识时，RAG 会用向量检索和 BM25 双路召回，再经过融合、重排和证据判断，证据不足时拒答。"
        "回答生成后，核心消息进入 MySQL，同事务写入 Outbox，由后台 Worker 补齐向量索引。这个设计的重点是让 Agent 不只会答，还能在真实业务里稳定运行。"
    )
    add_heading("3 分钟深挖版本", level=2)
    add_body(
        doc,
        "我会从三条线介绍。第一条是前台回答线：请求进入后，LangGraph 同时组织任务路由和历史召回，再进行上下文组装、查询改写、RAG 检索和 Agent 回答。"
        "第二条是事实保存线：用户消息、助手消息、任务状态都以 MySQL 为核心事实源，通过 request_id 避免重复写入，通过 state_version 乐观锁避免并发覆盖。"
        "第三条是后台索引线：因为 MySQL 和 Chroma 无法共享一个本地事务，所以我用事务性 Outbox 先记录待处理任务，再由 Worker 异步消费，失败时指数退避重试，卡住时做僵死任务恢复。"
        "所以这个项目的重点不是某一个模型效果，而是把 Agent、RAG、记忆和可靠性机制组合成可维护的业务链路。"
    )
    add_heading("表达提醒", level=2)
    add_matrix(
        doc,
        ["不要这样说", "应该这样说"],
        [
            ["我做了一个大模型客服", "我做了一个带路由、RAG、记忆和异步一致性的客服 Agent 系统"],
            ["LangGraph 让模型更聪明", "LangGraph 让复杂流程更可控、可测试、可降级"],
            ["Redis 保存聊天记录", "MySQL 是事实源，Redis 是短期缓存和加速层"],
            ["Outbox 保证 exactly-once", "Outbox 通常是至少一次投递，靠消费者幂等保证业务结果正确"],
            ["多 Agent 一定更好", "多 Agent 适合复杂跨领域问题，简单问题保留单 Agent 更稳"],
        ],
        [3900, 5460],
    )


def add_question_chain(
    doc: Document,
    title: str,
    interviewer: str,
    normal: str,
    strong: str,
    followups: list[tuple[str, str]],
):
    doc.add_heading(title, level=2)
    add_body(doc, f"面试官想考什么：{interviewer}", bold_lead="面试官想考什么：")
    add_matrix(
        doc,
        ["回答类型", "内容"],
        [
            ["普通回答", normal],
            ["高分回答", strong],
        ],
        [1800, 7560],
    )
    for idx, (q, a) in enumerate(followups, start=1):
        add_callout(doc, f"追问 {idx}", q, fill=GOLD_FILL, color=GOLD)
        add_body(doc, f"回答：{a}", bold_lead="回答：")


def add_question_bank_section(doc: Document, title: str, rows: list[tuple[str, str, str]]):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["题目", "面试官想考什么", "高分回答要点"]
    hdr = table.rows[0]
    repeat_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = ""
        cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, 9.5, DARK_BLUE, True)
        set_paragraph_spacing(p, after=2, line=1.1)
    for question, intent, answer in rows:
        cells = table.add_row().cells
        for idx, text in enumerate([question, intent, answer]):
            cells[idx].text = ""
            if idx == 0:
                cell_shading(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            r = p.add_run(text)
            set_font(r, 9.2, INK, bold=(idx == 0))
            set_paragraph_spacing(p, after=1, line=1.08)
    table_geometry(table, [2700, 2600, 4060])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_comprehensive_question_bank(doc: Document):
    doc.add_heading("B. 按知识点集中总题库（面试前按这个刷）", level=2)
    add_callout(
        doc,
        "刷题方法",
        "先看题目，自己用 20 秒说一遍；说不出来再看高分回答要点。不要逐字背，要背“因果链”：为什么需要、解决什么、和上下游怎么连、有什么边界。",
        fill=GREEN_FILL,
        color=GREEN,
    )
    sections: list[tuple[str, list[tuple[str, str, str]]]] = [
        (
            "1. 项目总览与真实性表达",
            [
                ("请你整体介绍一下这个项目。", "是否能从业务价值讲到系统链路。", "先讲家具客服场景，再讲路由、RAG、记忆、持久化、Outbox，最后强调不是裸调模型，而是工程化 Agent 链路。"),
                ("这个项目和普通聊天机器人有什么区别？", "是否知道项目卖点。", "普通聊天机器人偏问答；本项目强调证据检索、多轮记忆、任务状态、并发安全和异步一致性。"),
                ("你在项目里最核心的贡献是什么？", "是否能讲自己做了什么，而不是泛泛而谈。", "围绕主链路说：工作流编排、路由策略、RAG 证据门控、记忆、幂等和 Outbox 可靠链路。"),
                ("如果面试官质疑这是 AI 帮你做的，你怎么回答？", "是否诚实且能证明理解。", "承认使用 AI 辅助，但强调自己能解释业务链路、技术取舍、故障场景和优化方向，能独立维护和二次开发。"),
                ("这个项目最能体现工程能力的地方是什么？", "是否能跳出模型效果。", "不是页面聊天，而是把不确定的大模型包进确定性的流程、事务、幂等、锁、重试和降级机制里。"),
                ("项目当前最大的不足是什么？", "是否有边界感。", "可观测性、自动化评测集、部署稳定性、权限隔离、流式体验仍可增强，不把项目吹成生产满分系统。"),
            ],
        ),
        (
            "2. 技术选型题",
            [
                ("为什么前端选择 React？", "是否能从项目需要解释选型。", "客服系统需要组件化管理对话区、任务卡片、状态提示和历史列表；React 生态成熟，适合快速搭建交互型后台。"),
                ("为什么用 Python / FastAPI？", "是否知道后端和 AI 生态关系。", "Python 的 LangChain、LangGraph、向量检索、模型 SDK 生态更完整；FastAPI 异步支持好，适合 IO 密集的模型和检索调用。"),
                ("为什么选择 LangChain？", "是否知道 LangChain 的边界。", "LangChain 适合接模型、Prompt、工具、检索器和输出解析；但复杂状态流程不靠它硬写，而交给 LangGraph。"),
                ("为什么选择 LangGraph？", "是否能区分框架能力。", "当流程有状态、分支、并行、失败降级时，LangGraph 比手写 if-else 更清晰，便于定位节点和恢复。"),
                ("为什么要引入多 Agent？", "是否知道多 Agent 不是炫技。", "复杂问题跨家具、设备、报表等领域时，用专家分工降低单 Prompt 混乱；简单问题仍走单 Agent，避免过度设计。"),
                ("为什么用 Redis？", "是否知道缓存价值。", "近期上下文、短期状态、热点会话需要快速读取；Redis 是加速层，MySQL 才是事实源。"),
                ("为什么用 Chroma / 向量库？", "是否理解语义检索。", "用户表达不一定命中文档原词，向量库可以按语义找相近内容，用于知识 RAG 和历史语义召回。"),
                ("为什么不直接全用一个大模型解决？", "是否知道模型边界。", "模型擅长语言理解，不擅长事实一致性、事务、权限、并发控制；工程系统必须把事实和流程兜住。"),
            ],
        ),
        (
            "3. LangGraph 与工作流编排",
            [
                ("LangGraph 在你的项目里负责什么？", "是否知道它的位置。", "负责把路由、历史召回、上下文组装、RAG、Agent 回答、持久化、记忆刷新串成显式节点。"),
                ("LangGraph 和 LangChain 区别是什么？", "是否能区分工具链和流程图。", "LangChain 偏模型/工具/检索组件；LangGraph 偏有状态工作流、节点、边、条件分支和恢复。"),
                ("不用 LangGraph 行不行？", "是否有取舍意识。", "简单 Demo 可以不用；但节点多、状态多、分支多时，手写流程可维护性会下降。"),
                ("菱形图工作流是什么？", "是否理解架构描述。", "入口收敛到路由，路由后分支并行或选择专家，最后再汇聚到答案合成和持久化，形状上像发散再收敛。"),
                ("节点失败后怎么处理？", "是否能讲容错。", "辅助节点可降级，如历史召回失败就少用历史；核心持久化失败要失败快返，不能假装成功。"),
                ("LangGraph 会提升模型效果吗？", "是否避免错误吹法。", "不直接提升模型智商，它提升流程可控性；答案质量主要靠 RAG、Prompt、模型和评测。"),
            ],
        ),
        (
            "4. 多 Agent 设计",
            [
                ("什么是 Agent？", "是否掌握基础概念。", "Agent 是能基于目标决定是否调用工具、读取知识、执行步骤并生成结果的执行单元。"),
                ("什么是多 Agent？", "是否理解分工。", "多个角色 Agent 按职责分工处理问题，再由合成器融合结果，像多科室会诊。"),
                ("多 Agent 有什么风险？", "是否知道复杂度成本。", "调用成本更高、延迟更大、结果可能冲突、调试更难，所以必须有路由、上限和合成策略。"),
                ("为什么最多路由 2 个 Agent？", "是否理解限制设计。", "防止每个问题都广播给所有专家，控制成本、延迟和冲突概率。"),
                ("答案合成器做什么？", "是否知道汇聚阶段。", "把多个专家答案去重、消歧、处理冲突，生成统一口径的最终回复。"),
                ("专家 Agent 冲突怎么办？", "是否能处理复杂场景。", "优先看证据可信度、业务优先级和是否需要追问；不能硬合并成自相矛盾的回答。"),
            ],
        ),
        (
            "5. 路由与意图识别",
            [
                ("任务路由和专家路由有什么区别？", "是否能讲清两级路由。", "任务路由判断这是什么事；专家路由决定交给哪个能力模块或专家 Agent。"),
                ("为什么规则引擎 + LLM 兜底？", "是否知道成本与稳定性。", "明确关键词用规则便宜稳定，模糊表达再让 LLM 判断，兼顾成本、速度和召回。"),
                ("关键词加权打分怎么理解？", "是否理解规则路由。", "不同关键词对不同专家有不同权重，累计分数达到阈值才路由，避免单个词误判。"),
                ("复合词检测为什么重要？", "是否知道中文场景坑点。", "比如“扫地机器人”和“机器人”不能重复计分，否则会放大某个方向的分数。"),
                ("延续意图怎么判断？", "是否理解多轮任务。", "看当前问题是否依赖上文，如“它还是不行”，结合 session、任务状态、历史摘要判断。"),
                ("路由错了有什么后果？", "是否知道路由是入口。", "会影响检索范围、专家选择、上下文材料，后面 RAG 和 Agent 只能部分补救。"),
            ],
        ),
        (
            "6. RAG / 检索 / 重排序 / 拒答",
            [
                ("RAG 是什么？", "是否掌握核心定义。", "先检索企业知识，再让模型基于证据回答，降低模型凭空编造。"),
                ("为什么要查询归一化？", "是否理解检索前处理。", "把口语化、错别字、指代和冗余表达整理成更适合检索的查询。"),
                ("向量检索和 BM25 分别解决什么？", "是否理解混合检索。", "向量抓语义相似，BM25 抓型号、术语、错误码等精确词。"),
                ("RRF 融合是什么思路？", "是否知道融合价值。", "不直接相信单一路召回，而是按排名融合多路结果，提高稳健性。"),
                ("CrossEncoder 重排序为什么比向量召回更准？", "是否理解召回与精排。", "召回快但粗，CrossEncoder 同时看 query 和文档，判断相关性更细，但成本更高。"),
                ("证据可信度怎么判断？", "是否知道拒答门控。", "看证据是否覆盖对象、意图、关键约束、来源和时效；不足就追问或拒答。"),
                ("为什么要拒答？", "是否有可信度意识。", "客服场景宁可说证据不足，也不能编保修规则、清洁方法或故障处理。"),
                ("RAG 和微调怎么选？", "是否理解知识更新。", "频繁变化的企业知识优先 RAG；格式、风格、分类习惯可考虑微调。"),
            ],
        ),
        (
            "7. Redis / 记忆 / 多轮对话",
            [
                ("三层记忆分别是什么？", "是否理解记忆分层。", "短期 Redis、会话摘要、长期用户画像，分别解决近期上下文、长对话压缩和跨会话偏好。"),
                ("为什么不能把所有历史都塞给模型？", "是否知道上下文成本。", "token 成本高、噪声大、容易干扰当前问题，所以要摘要和相关召回。"),
                ("精确召回和语义召回区别是什么？", "是否理解历史检索。", "精确召回按 session、task、用户等字段找；语义召回按意思找相似历史。"),
                ("Redis 和 MySQL 的边界是什么？", "是否知道事实源。", "Redis 是快但可丢的缓存，MySQL 是核心事实源。"),
                ("用户说“它还是不行”，系统怎么理解？", "是否能落到场景。", "结合当前会话、任务状态、历史摘要和语义召回，定位“它”指哪个商品/问题。"),
                ("记忆错误会带来什么风险？", "是否知道上下文污染。", "错误记忆会误导路由和回答，所以画像和摘要要可更新、可覆盖、可追溯。"),
            ],
        ),
        (
            "8. request_id 幂等 / 乐观锁 / 并发安全",
            [
                ("request_id 是什么？", "是否理解幂等入口。", "每次业务请求的唯一标识，用来识别网络重试、用户连点导致的重复请求。"),
                ("幂等解决什么问题？", "是否区分重复执行和结果一致。", "不保证物理只执行一次，而是保证同一请求重复到达时业务结果不重复。"),
                ("幂等一般怎么实现？", "是否知道工程手段。", "request_id 唯一约束、幂等表、状态机检查、重复请求返回已有结果。"),
                ("乐观锁解决什么问题？", "是否区分并发覆盖。", "不同请求同时更新同一任务时，用 version 检查防止后写覆盖前写。"),
                ("幂等和乐观锁最大区别是什么？", "是否能一句话讲清。", "幂等防同一请求重复，乐观锁防不同请求并发覆盖。"),
                ("乐观锁冲突后怎么办？", "是否理解冲突处理。", "重新读最新状态，再决定重试、合并、放弃或提示用户，而不是强行覆盖。"),
            ],
        ),
        (
            "9. Outbox / Worker / 最终一致性",
            [
                ("Outbox 解决什么问题？", "是否理解双写。", "解决 MySQL 成功但 Chroma/MQ/外部索引失败导致的数据不一致。"),
                ("为什么 Outbox 要和业务数据同事务？", "是否理解事务边界。", "业务消息和待办任务一起提交，保证只要事实保存成功，就一定有后台补偿任务。"),
                ("FOR UPDATE SKIP LOCKED 有什么用？", "是否理解并发消费。", "多个 Worker 抢任务时跳过已锁行，提高并发且避免重复领取同一行。"),
                ("指数退避重试为什么需要？", "是否知道故障恢复。", "外部服务短暂故障时不要疯狂重试，按间隔逐渐拉长，减少雪崩。"),
                ("僵死任务怎么恢复？", "是否理解 Worker 崩溃场景。", "任务被领取后 Worker 崩溃，需要扫描超时 processing 任务重新置回可执行。"),
                ("Outbox 是 exactly-once 吗？", "是否避免错误说法。", "通常是至少一次投递，靠消费者幂等 upsert 保证业务结果正确。"),
                ("Outbox 和 MQ 有什么区别？", "是否理解边界。", "MQ 负责跨服务消息传递，Outbox 先解决本地数据库提交和消息发送的原子性问题。"),
            ],
        ),
        (
            "10. Agent 沙箱 / 权限控制 / 安全",
            [
                ("为什么 Agent 需要沙箱？", "是否知道行动风险。", "Agent 能调工具后，风险从答错变成做错，所以必须限制工具、参数、数据和高危动作。"),
                ("工具白名单是什么？", "是否理解最小权限。", "只开放当前角色需要的工具，不把数据库、文件、网络能力全部暴露给模型。"),
                ("参数校验为什么重要？", "是否知道 Prompt 攻击风险。", "模型可能生成危险参数，系统要校验范围、类型、权限和业务规则。"),
                ("哪些操作需要人工确认？", "是否能落到场景。", "删除数据、导出隐私、退款、改订单、批量更新索引等高风险操作需要确认。"),
                ("如何防止提示词注入？", "是否理解安全边界。", "外部文档内容不能直接改变系统规则；工具调用前做权限校验、内容隔离和审计。"),
                ("权限控制和模型 Prompt 哪个更可靠？", "是否知道系统控制优先。", "Prompt 是软约束，权限系统是硬约束；生产系统不能只靠模型自觉。"),
            ],
        ),
        (
            "11. 大模型基础 / 微调 / 训练 / 推理 / 评测",
            [
                ("大模型为什么会幻觉？", "是否理解概率生成。", "模型按概率生成看似合理的文本，不天然知道企业事实是否存在，所以需要 RAG 和证据门控。"),
                ("Prompt 工程解决什么？", "是否知道低成本调优。", "通过角色、任务、约束、示例和输出格式提升稳定性，是改代码/微调前的低成本方案。"),
                ("微调适合解决什么？", "是否区分能力边界。", "适合固定格式、领域话术、分类习惯、工具调用风格，不适合频繁更新事实知识。"),
                ("LoRA 是什么思路？", "是否掌握微调基础。", "冻结大部分模型参数，只训练小规模低秩适配参数，成本比全量微调低。"),
                ("训练、微调、RAG 区别是什么？", "是否能宏观区分。", "训练塑造基础能力，微调塑造任务行为，RAG 动态提供外部知识。"),
                ("推理温度 temperature 有什么影响？", "是否理解生成参数。", "温度越高越发散，客服场景通常要低温保证稳定和一致。"),
                ("如何评测一个客服 Agent？", "是否知道效果闭环。", "看意图识别准确率、检索命中率、答案事实性、拒答合理性、延迟、成本和用户满意度。"),
                ("怎么构建评测集？", "是否能工程化落地。", "从真实问题、边界问题、错误案例、拒答案例、多轮指代案例抽样，维护标准答案和证据。"),
            ],
        ),
        (
            "12. 大厂高频场景题与 HR 追问",
            [
                ("如果模型服务超时怎么办？", "是否有降级设计。", "设置超时、重试、降级回复、排队提示；核心写库不能依赖模型成功。"),
                ("如果 RAG 检索结果互相矛盾怎么办？", "是否能处理证据冲突。", "按来源可信度、时效、业务优先级判断；无法确定就追问或拒答，不硬合并。"),
                ("如果成本太高怎么优化？", "是否知道成本控制。", "缓存、规则优先、限制多 Agent 数量、小模型做路由、大模型只处理复杂环节。"),
                ("如果并发量上来，瓶颈在哪里？", "是否能系统分析。", "模型调用、向量检索、数据库连接、Worker 吞吐、Redis 热点都可能是瓶颈，要用监控定位。"),
                ("如果让你上线生产，还缺什么？", "是否知道生产化要求。", "鉴权、权限、日志、监控告警、灰度、评测集、数据备份、限流和安全审计。"),
                ("你怎么证明 RAG 有效果？", "是否有评测意识。", "用固定问题集对比裸模型、单向量检索、混合检索，统计事实准确率、命中率和拒答合理性。"),
                ("你遇到最难的问题是什么？", "是否能讲项目经历。", "建议讲多轮上下文和可靠双写：它们不是调 Prompt 就能解决，必须靠记忆、幂等、锁和 Outbox。"),
                ("你简历写得很高级，实际做到哪一步？", "是否诚实可靠。", "按已实现、实验中、可优化三层讲，不夸大；强调自己能解释设计、复现主链路和继续迭代。"),
            ],
        ),
    ]
    total = sum(len(rows) for _, rows in sections)
    add_callout(
        doc,
        "题量说明",
        f"本总题库共 {total} 道，覆盖项目总览、选型、LangGraph、多 Agent、路由、RAG、记忆、并发、Outbox、沙箱、微调训练和大厂场景题。前面的深挖题练表达，这里的总题库练覆盖面。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    for title, rows in sections:
        add_question_bank_section(doc, title, rows)


def add_chapter_7(doc: Document):
    doc.add_heading("第 7 章：追问题库", level=1)
    add_body(
        doc,
        "这一章不是让你背所有题，而是训练你接住连续追问。每个题都按“面试官想考什么、普通回答、高分回答、继续追问”来准备。"
    )
    doc.add_heading("A. 高频深挖追问链（先练这部分）", level=2)
    add_question_chain(
        doc,
        "1. 你为什么用 LangGraph？",
        "看你是否知道 LangGraph 解决的是流程复杂度，而不是模型能力。",
        "因为项目流程比较复杂，所以用了 LangGraph。",
        "我用 LangGraph 不是为了让模型更聪明，而是因为项目已经不是一次模型调用。它有任务路由、历史召回、上下文组装、RAG、回答、持久化和记忆刷新。LangGraph 能把这些节点、状态和分支显式化，方便并行、降级和测试。",
        [
            ("不用 LangGraph，自己写函数行不行？", "简单流程可以。但复杂后会出现状态散落、分支混乱、故障难定位的问题。我的判断标准不是步骤数量，而是流程是否有状态、分支、并行和恢复需求。"),
            ("LangGraph 提升答案质量了吗？", "它不直接提升答案质量。答案质量主要靠 RAG、提示词和评测。LangGraph 提升的是流程可控性和工程可维护性。"),
            ("节点失败怎么办？", "辅助节点可以降级，比如历史召回失败就少用历史；核心持久化失败必须失败快返，不能假装成功。"),
        ],
    )
    add_question_chain(
        doc,
        "2. 为什么需要两级路由？",
        "看你是否能区分任务分类和专家选择。",
        "因为问题不同，所以要路由给不同 Agent。",
        "我把路由拆成任务路由和专家路由。任务路由先判断这是新任务、延续任务、恢复旧任务还是闲聊；专家路由再判断是否需要家具、设备、报表或通用专家。这样入口分流和能力分发分开，逻辑更清晰，也更容易评测和扩展。",
        [
            ("为什么不用一个 LLM 一次性判断完？", "可以做，但成本、稳定性和可解释性会差一些。高频明确场景用规则更便宜、更稳定，模糊表达再交给 LLM 兜底。"),
            ("路由错了后面 RAG 能不能救回来？", "只能部分缓解。路由错会导致上下文、专家和检索范围都偏掉，所以路由是整条链路的资源调度入口。"),
        ],
    )
    add_question_chain(
        doc,
        "3. RAG 为什么要做混合检索和证据门控？",
        "看你是否知道 RAG 的重点是证据质量，不是只要检索到文本。",
        "为了让模型回答更准确，所以用了 RAG。",
        "项目里 RAG 的目标是让回答有企业知识依据。向量检索擅长语义相似，BM25 擅长型号、术语、错误码等关键词命中，所以用双路召回再融合重排。最后还要做证据门控，判断证据是否覆盖对象、意图和关键术语，不够就拒答。",
        [
            ("为什么不只用向量检索？", "向量检索可能语义相似但漏掉精确型号或关键词；BM25 正好补这个短板。"),
            ("拒答会不会影响体验？", "短期看保守，长期看能保护可信度。更好的方向是提高证据覆盖和澄清能力，而不是没证据也硬答。"),
        ],
    )
    add_question_chain(
        doc,
        "4. Redis、MySQL、Chroma 分别干什么？",
        "看你是否能讲清存储边界。",
        "MySQL 存数据，Redis 做缓存，Chroma 做向量库。",
        "MySQL 是核心事实源，保存消息、任务、文档版本和 Outbox；Redis 是短期缓存和加速层，适合近期上下文；Chroma 是语义搜索索引，用来做知识检索和历史语义召回。Redis 和 Chroma 都可以围绕 MySQL 重建，MySQL 里的事实不能随便丢。",
        [
            ("Redis 挂了系统还能用吗？", "应该能降级，只是近期上下文读取变慢或能力变弱。核心事实仍在 MySQL。"),
            ("Chroma 挂了怎么办？", "写入侧通过 Outbox 等待恢复；读取侧可以降级到 BM25，但证据不足仍要拒答。"),
        ],
    )
    add_question_chain(
        doc,
        "5. request_id 幂等和乐观锁有什么区别？",
        "看你能否区分重复请求和并发覆盖。",
        "幂等和锁都是为了避免数据出错。",
        "request_id 幂等解决同一请求重复到达，比如用户连点或网络重试；乐观锁解决不同请求同时更新同一任务状态。前者防重复，后者防覆盖。两者都需要，因为真实系统里这两类问题都会发生。",
        [
            ("幂等是不是保证请求只执行一次？", "不是。它保证的是重复执行后业务结果一致，不是物理层绝对只跑一次。"),
            ("乐观锁冲突后怎么办？", "重新读取最新状态再判断是否重试、合并或放弃，不能检测到冲突后还强行覆盖。"),
        ],
    )
    add_question_chain(
        doc,
        "6. 为什么需要 Outbox？",
        "看你是否理解跨存储双写问题。",
        "因为后台任务要异步执行，所以用了 Outbox。",
        "Outbox 的核心是解决 MySQL 和 Chroma 之间的可靠双写。用户消息和 Outbox 任务在同一个 MySQL 事务里提交，先保证核心事实不丢。之后 Worker 至少一次消费任务，把索引写入 Chroma。因为可能重复消费，所以消费者要用稳定 ID 做幂等 upsert。",
        [
            ("为什么不用后台线程？", "线程没有可靠持久化任务记录，进程崩溃就可能丢。Outbox 至少能恢复和重试。"),
            ("为什么不用 MQ 直接替代？", "MQ 适合跨服务传递，但数据库提交和发 MQ 之间仍有原子性问题。Outbox 先解决本地事务边界。"),
            ("Outbox 能保证 exactly-once 吗？", "通常不能。它更现实的目标是至少一次投递加消费者幂等，保证业务效果正确。"),
        ],
    )
    add_question_chain(
        doc,
        "7. 这个项目有哪些真实不足？",
        "看你是否诚实，以及是否有优化判断。",
        "项目还有一些地方可以优化。",
        "我会诚实说，当前项目主链路和工程化机制已经具备，但还有优化空间：比如真正 token 级流式体验、BM25 索引缓存、评测集持续扩充、线上可观测性、权限隔离和部署稳定性都可以继续加强。我的态度不是把项目吹成完美系统，而是能说清当前边界和下一步优先级。",
        [
            ("如果让你只优化一项，你先做什么？", "我会先看面试或业务目标。如果偏演示体验，先做 token 流式；如果偏生产可靠性，先补可观测性和部署验证；如果偏效果，先补评测集和 RAG 样本闭环。"),
        ],
    )
    add_question_chain(
        doc,
        "8. Agent 为什么需要沙箱和权限边界？",
        "看你是否知道 Agent 会调用工具后，风险不只在回答错误，还在错误行动。",
        "因为 Agent 可能会乱调用工具，所以要限制权限。",
        "我理解 Agent 沙箱的核心是把模型的行动能力限制在业务允许的范围内。比如客服 Agent 可以查知识库、生成建议、写普通工单，但不能直接删除数据、导出隐私或执行高风险退款。实现上可以用工具白名单、参数校验、用户/角色权限、只读优先、危险操作二次确认和审计日志。这样既保留 Agent 自动化能力，也防止提示词攻击或模型误判造成越权操作。",
        [
            ("如果模型说必须调用一个高权限工具怎么办？", "不能因为模型说必须就执行。高风险工具要走规则校验和人工确认，模型只能提出建议，最终权限由系统控制。"),
            ("沙箱会不会降低 Agent 能力？", "会限制一部分自由度，但这是生产系统必须付出的代价。工程上追求的是可控的能力，而不是无限权限。"),
            ("家具客服项目里哪些动作需要权限边界？", "查询知识库和读取公开产品信息风险较低；修改订单、导出用户历史、删除索引、改任务状态都需要更严格校验。"),
        ],
    )
    add_question_chain(
        doc,
        "9. 为什么用 RAG，而不是直接微调或重新训练模型？",
        "看你是否能区分知识注入、行为训练和成本边界。",
        "RAG 更方便更新知识，微调比较麻烦。",
        "我会把 RAG 和微调分开看。家具客服最核心的问题是企业知识经常变化，比如产品资料、保养规则、故障处理文档，这类知识更适合用 RAG，因为更新知识库就能生效，而且答案可以追溯到证据。微调更适合让模型学会稳定格式、固定话术、特定分类习惯或工具调用风格；如果为了更新几条产品知识就微调，成本高、周期长，也很难保证可追溯。",
        [
            ("那微调在这个项目完全没价值吗？", "不是。微调可以用于稳定意图分类、客服话术风格、结构化输出格式，但不应该替代知识库检索。"),
            ("重新训练大模型可行吗？", "对个人项目和中小业务基本不现实。训练需要大量数据、算力和评测体系，项目里更现实的是 RAG、提示词、少量微调或 LoRA。"),
            ("如果 RAG 检索不到，微调能不能救？", "不能可靠救。微调可能记住部分知识，但知识是否最新、是否可追溯都不好保证；检索不到证据时更应该拒答或追问。"),
        ],
    )
    add_comprehensive_question_bank(doc)


def add_chapter_8(doc: Document):
    doc.add_heading("第 8 章：面试前速记卡片", level=1)
    add_callout(
        doc,
        "项目一句话",
        "这是一个面向家具和扫地机器人场景的智能客服 Agent 系统，用 LangGraph 编排流程，用混合 RAG 提供证据，用分层记忆解决多轮上下文，用幂等、乐观锁和 Outbox 保证工程可靠性。",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
    )
    add_matrix(
        doc,
        ["关键词", "一句话记忆"],
        [
            ["LangGraph", "流程导演，负责节点、状态、分支和降级。"],
            ["路由", "分诊台，先判断是什么事，再决定找谁。"],
            ["多 Agent", "分科会诊，复杂问题分工处理。"],
            ["RAG", "先查证据，再让模型回答。"],
            ["证据门控", "证据不够就拒答，别让模型硬编。"],
            ["Redis", "桌面便签，短期缓存和加速。"],
            ["MySQL", "正式档案库，核心事实源。"],
            ["Chroma", "语义索引，按意思找知识和历史。"],
            ["request_id", "同一请求的身份证，用来防重复。"],
            ["乐观锁", "版本检查，防止并发覆盖。"],
            ["Outbox", "可靠后台工单，防止异步任务悄悄丢。"],
            ["Agent 沙箱", "给模型行动能力加护栏：工具白名单、参数校验、权限分级。"],
            ["微调/训练", "适合稳定行为和格式，不适合频繁更新企业知识；知识更新优先 RAG。"],
        ],
        [2100, 7260],
    )
    doc.add_heading("最后三句话", level=2)
    add_bullet(doc, "我不会把所有问题都交给模型，事实、流程、一致性和权限边界都需要工程机制兜住。")
    add_bullet(doc, "这个项目的核心价值不是能聊天，而是把智能客服做成可检索、可记忆、可恢复、可解释的链路。")
    add_bullet(doc, "面试时先讲业务链路，再讲技术因果，最后讲取舍和边界。")


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    setup_page_furniture(doc)
    add_title_page(doc)
    add_navigation(doc)
    add_chapter_0(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_chapter_7(doc)
    add_chapter_8(doc)
    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    if OUTPUT.exists():
        backup = OUTPUT.with_name(f"{OUTPUT.stem}-旧版备份-{datetime.now().strftime('%Y%m%d-%H%M%S')}{OUTPUT.suffix}")
        backup.write_bytes(OUTPUT.read_bytes())
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
