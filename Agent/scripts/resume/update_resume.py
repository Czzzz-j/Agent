from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\陈梓杰\Desktop\陈梓杰简历\陈梓杰-简历.docx")
OUTPUT = Path(r"C:\Users\陈梓杰\Desktop\陈梓杰简历\陈梓杰-简历-更新版.docx")


def _copy_run_properties(source_run, target_run) -> None:
    source_rpr = source_run._r.rPr
    if source_rpr is None:
        return
    target_rpr = target_run._r.get_or_add_rPr()
    target_rpr.getparent().remove(target_rpr)
    target_run._r.insert(0, deepcopy(source_rpr))


def _run_template(paragraph, *, bold: bool | None = None, fallback_index: int = 0):
    for run in paragraph.runs:
        if bold is None or bool(run.bold) is bold:
            return run
    return paragraph.runs[min(fallback_index, len(paragraph.runs) - 1)]


def replace_paragraph_runs(paragraph, chunks: list[tuple[str, object]]) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    for text, template_run in chunks:
        run = paragraph.add_run(text)
        _copy_run_properties(template_run, run)


def replace_label_body(paragraph, label: str, body: str) -> None:
    label_template = _run_template(paragraph, bold=True)
    body_template = _run_template(paragraph, bold=False, fallback_index=0)
    replace_paragraph_runs(
        paragraph,
        [
            (label, label_template),
            (body, body_template),
        ],
    )


def replace_plain(paragraph, text: str) -> None:
    template = paragraph.runs[0]
    replace_paragraph_runs(paragraph, [(text, template)])


def update_resume() -> None:
    doc = Document(SOURCE)
    paragraphs = doc.paragraphs

    # Contact information: fix the existing email typo only.
    replace_plain(paragraphs[1], "177-2829-6516  | 965064261@qq.com")

    # Preserve the existing project-title formatting and hyperlink color.
    project_title = paragraphs[10]
    project_templates = list(project_title.runs)
    replace_paragraph_runs(
        project_title,
        [
            ("全屋家具智能管家 Agent", project_templates[0]),
            (" |", project_templates[1]),
            (" Agent 应用开发 | ", project_templates[2]),
            ("https://github.com/Czzzz-j/Agent.git", project_templates[3]),
            ("\t", project_templates[4]),
            ("2026.03 - 2026.05", project_templates[5]),
        ],
    )

    replace_label_body(
        paragraphs[11],
        "项目介绍：",
        "面向家具与扫地机器人场景的智能客服系统，围绕多轮任务跟踪、知识库问答和跨会话记忆进行设计，"
        "支持 Agent 工具调用、混合 RAG、任务路由及异步索引。",
    )
    replace_label_body(
        paragraphs[12],
        "技术栈：",
        "Python、LangGraph、LangChain、FastAPI、ChromaDB、Redis、MySQL、BM25、BGE Reranker、Streamlit",
    )
    replace_label_body(
        paragraphs[13],
        "工作流编排：",
        "使用 LangGraph 组织任务路由、历史召回、上下文组装、回答生成和持久化；"
        "兼容 Legacy 与 Agentic 两种执行路径，对路由、召回和记忆等辅助节点设置降级。",
    )
    replace_label_body(
        paragraphs[14],
        "多 Agent 路由：",
        "在专家模式下先用规则对家具、设备、报表和通用问题打分，低置信度时由 LLM 兜底；"
        "复合问题最多选择 2 个专家并行处理，再统一合成结果。",
    )
    replace_label_body(
        paragraphs[15],
        "混合 RAG：",
        "对查询提取对象、意图、型号和错误码等信息，结合向量检索与 BM25 召回，"
        "经 RRF 融合和 BGE CrossEncoder 重排后进行证据判断；证据不足时直接拒答。",
    )
    replace_label_body(
        paragraphs[16],
        "任务与上下文：",
        "通过闲聊过滤、延续意图、关键词和语义相似度识别新任务或恢复旧任务；"
        "使用 request_id 防止重复写入，并通过版本号处理并发更新冲突。",
    )
    replace_label_body(
        paragraphs[17],
        "分层记忆：",
        "Redis 缓存近期对话，MySQL 保存会话摘要与用户画像，Chroma 建立跨会话历史索引；"
        "精确问题查询原始消息，模糊历史问题采用语义召回并回库验证。",
    )
    replace_label_body(
        paragraphs[18],
        "异步索引：",
        "消息和 Outbox 任务在同一 MySQL 事务中提交，Worker 并发领取任务并写入向量库；"
        "失败任务支持指数退避、超时恢复和幂等重试。",
    )

    # Replace the second project with an honest internship placeholder.
    internship_title = paragraphs[19]
    internship_templates = list(internship_title.runs)
    replace_paragraph_runs(
        internship_title,
        [
            ("【公司名称】", internship_templates[0]),
            (" |", internship_templates[1]),
            (" Agent 应用开发实习生", internship_templates[2]),
            ("\t", internship_templates[3]),
            ("2026.06 - 至今", internship_templates[4]),
        ],
    )
    replace_label_body(
        paragraphs[20],
        "实习方向：",
        "参与【业务 / 项目名称】相关的 Agent 应用开发。",
    )
    replace_label_body(
        paragraphs[21],
        "技术环境：",
        "【按实际填写，例如 Python、FastAPI、LangGraph、MySQL、Redis】",
    )
    replace_label_body(
        paragraphs[22],
        "当前工作：",
        "参与【需求理解、功能开发、接口联调、测试或问题排查，根据实际保留】。",
    )
    replace_label_body(
        paragraphs[23],
        "负责内容：",
        "【填写自己独立负责的模块，以及与上下游的协作方式】",
    )
    replace_label_body(
        paragraphs[24],
        "阶段成果：",
        "【填写已上线功能、修复的问题、效率提升或其他可量化结果】",
    )
    replace_label_body(
        paragraphs[25],
        "后续补充：",
        "实习进行中，结束前根据实际工作内容统一更新。",
    )
    # Keep the existing spacing paragraph but remove the old fabricated metric.
    for run in list(paragraphs[26].runs):
        run._element.getparent().remove(run._element)

    # Update the section title without changing the banner/table structure.
    section_heading = doc.tables[1].cell(0, 0).paragraphs[0]
    replace_plain(section_heading, "项目及实习经历")

    # Rewrite skills and strengths in a more credible, interview-friendly tone.
    replace_label_body(
        paragraphs[29],
        "编程与后端：",
        "熟悉 Python、FastAPI、MySQL、Redis，能够完成接口开发、数据持久化、异常处理与基础测试；了解 Java、C++。",
    )
    replace_label_body(
        paragraphs[30],
        "Agent 应用开发：",
        "使用 LangChain、LangGraph 搭建过工具调用、状态编排、多轮对话、任务路由与上下文组装，"
        "能够独立完成 Agent 业务链路开发和问题排查。",
    )
    replace_label_body(
        paragraphs[31],
        "RAG 与检索：",
        "熟悉文档切分、Embedding、向量检索、BM25、RRF 融合、CrossEncoder 重排和拒答策略；"
        "使用过 ChromaDB、Qdrant。",
    )
    replace_label_body(
        paragraphs[32],
        "数据与工程化：",
        "了解 MySQL 事务、request_id 幂等、乐观锁、Redis 缓存和 Outbox 异步任务，"
        "能够通过日志与测试定位常见问题。",
    )
    replace_label_body(
        paragraphs[33],
        "模型与 Prompt：",
        "理解工具调用、结构化输出、查询改写、上下文管理与幻觉控制，能够根据测试结果迭代 Prompt 和检索策略。",
    )
    replace_label_body(
        paragraphs[34],
        "学习与协作：",
        "习惯使用 Git、Codex、Claude Code、Cursor 辅助开发，能够阅读现有项目、拆解问题并通过测试验证修改。",
    )

    # Update the banner text while preserving its original style and geometry.
    skill_heading = doc.tables[2].cell(0, 0).paragraphs[0]
    replace_plain(skill_heading, "专业技能与个人优势")

    doc.save(OUTPUT)


if __name__ == "__main__":
    update_resume()
    print(OUTPUT)
